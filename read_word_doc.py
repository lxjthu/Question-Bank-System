"""
Python脚本用于读取Word文档内容和基本格式信息
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
import os

def read_word_document(file_path):
    """
    读取Word文档内容和基本格式信息
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return
    
    try:
        doc = Document(file_path)
        print(f"正在读取文档: {file_path}")
        print("="*50)
        
        # 读取段落内容和格式
        print("段落内容:")
        print("-" * 30)
        for i, paragraph in enumerate(doc.paragraphs):
            print(f"段落 {i+1}:")
            print(f"  文本: {paragraph.text}")
            
            # 检查段落格式
            para_format = paragraph.paragraph_format
            print(f"  对齐方式: {para_format.alignment if para_format.alignment is not None else '默认'}")
            print(f"  行间距: {para_format.line_spacing if para_format.line_spacing is not None else '默认'}")
            
            # 读取段落中的格式化文本
            if paragraph.runs:
                print("  格式化文本:")
                for j, run in enumerate(paragraph.runs):
                    print(f"    Run {j+1}: '{run.text}'")
                    print(f"      粗体: {run.bold}")
                    print(f"      斜体: {run.italic}")
                    print(f"      下划线: {run.underline}")
                    if run.font.size:
                        print(f"      字体大小: {run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size}")
                    if run.font.name:
                        print(f"      字体名称: {run.font.name}")
                    print()
        
        # 读取表格内容
        print("\n表格内容:")
        print("-" * 30)
        for i, table in enumerate(doc.tables):
            print(f"表格 {i+1}:")
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                print(f"  行 {row_idx+1}: {row_data}")
        
        # 读取文档属性
        print("\n文档属性:")
        print("-" * 30)
        core_props = doc.core_properties
        print(f"标题: {core_props.title if core_props.title else '无'}")
        print(f"作者: {core_props.author if core_props.author else '无'}")
        print(f"主题: {core_props.subject if core_props.subject else '无'}")
        print(f"创建时间: {core_props.created if core_props.created else '无'}")
        print(f"修改时间: {core_props.modified if core_props.modified else '无'}")
        
    except Exception as e:
        print(f"读取文档时发生错误: {str(e)}")

def extract_text_only(file_path):
    """
    仅提取文档文本内容
    """
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"提取文本时发生错误: {str(e)}")
        return ""

if __name__ == "__main__":
    # 请将此处路径替换为您的Word文档路径
    file_path = r"D:\code\试卷\林业经济学（双语）试卷A-打印.docx"
    
    # 读取完整文档信息
    read_word_document(file_path)
    
    print("\n" + "="*50)
    print("仅文本内容:")
    print("-" * 30)
    text_content = extract_text_only(file_path)
    print(text_content)