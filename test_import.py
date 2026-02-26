from docx import Document
import os

# Create a test Word document with questions in the expected format
doc = Document()

# Add a single choice question
doc.add_paragraph('[单选][D]')
doc.add_paragraph('湖北省的省会城市是（）？')
doc.add_paragraph('[A]长沙')
doc.add_paragraph('[B]宜昌')
doc.add_paragraph('[C]荆州')
doc.add_paragraph('[D]武汉')
doc.add_paragraph('<解析>')
doc.add_paragraph('这是一段解析')
doc.add_paragraph('武汉是湖北省省会城市')
doc.add_paragraph('<解析>')

# Add a multiple choice question
doc.add_paragraph('[多选][ABD]')
doc.add_paragraph('请问下列哪些书籍是中国四大小说名著？')
doc.add_paragraph('[A]三国演义')
doc.add_paragraph('[B]红楼梦')
doc.add_paragraph('[C]老残游记')
doc.add_paragraph('[D]水浒传')
doc.add_paragraph('<解析>')
doc.add_paragraph('老残游记不属于四大名著，其他都属于')
doc.add_paragraph('<解析>')

# Add an essay question
doc.add_paragraph('[简答]')
doc.add_paragraph('你对今天这节课有什么评价？')
doc.add_paragraph('<参考答案>')
doc.add_paragraph('这是一段参考答案')
doc.add_paragraph('课程非常的生动')
doc.add_paragraph('<参考答案>')

# Save the test document
test_file_path = 'temp/test_questions.docx'
os.makedirs('temp', exist_ok=True)
doc.save(test_file_path)

print(f"Test Word document created at: {test_file_path}")