import os
import json
import uuid
from html.parser import HTMLParser
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import csv
import re

# Directory for storing question images on disk
_HERE = os.path.dirname(os.path.abspath(__file__))
_IMAGES_DIR = os.path.join(os.path.dirname(_HERE), 'uploads', 'images')


def allowed_file(filename):
    """Check if the uploaded file is allowed"""
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'doc', 'docx', 'csv'}
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_image_file(image_bytes: bytes, content_type: str,
                    question_id=None, field: str = 'content') -> str:
    """Save an image to disk and create a QuestionImageModel record.

    Returns the image_id string (e.g. 'img_abc12345').
    The DB record is flushed but NOT committed — caller must commit.
    """
    from app.db_models import db, QuestionImageModel

    # Determine extension from content_type
    ext_map = {
        'image/png': '.png', 'image/jpeg': '.jpg', 'image/jpg': '.jpg',
        'image/gif': '.gif', 'image/bmp': '.bmp', 'image/webp': '.webp',
    }
    ext = ext_map.get((content_type or '').lower(), '.png')

    image_id = 'img_' + uuid.uuid4().hex[:8]
    filename = image_id + ext

    os.makedirs(_IMAGES_DIR, exist_ok=True)
    filepath = os.path.join(_IMAGES_DIR, filename)
    with open(filepath, 'wb') as f:
        f.write(image_bytes)

    record = QuestionImageModel(
        image_id=image_id,
        question_id=question_id,
        field=field,
        filename=filename,
        content_type=content_type or 'image/png',
        file_size=len(image_bytes),
    )
    db.session.add(record)
    db.session.flush()
    return image_id


def delete_question_images(question_id: str) -> None:
    """Delete all image files and DB records associated with question_id.

    Caller must commit after this call.
    """
    from app.db_models import db, QuestionImageModel

    records = QuestionImageModel.query.filter_by(question_id=question_id).all()
    for rec in records:
        try:
            os.remove(os.path.join(_IMAGES_DIR, rec.filename))
        except OSError:
            pass
        db.session.delete(rec)


def _associate_images_in_html(html_content: str, question_id: str) -> None:
    """Find all image_ids in HTML and update their question_id to question_id.

    Call after a question is committed to associate orphan images (uploaded
    before the question existed) with the question.
    Does NOT commit — caller must commit.
    """
    from app.db_models import db, QuestionImageModel

    if not html_content:
        return
    img_ids = re.findall(r'/api/images/(img_[a-f0-9]+)', html_content)
    if img_ids:
        QuestionImageModel.query.filter(
            QuestionImageModel.image_id.in_(img_ids),
            QuestionImageModel.question_id.is_(None)
        ).update({'question_id': question_id}, synchronize_session=False)




# ── HTML → Word helpers ──────────────────────────────────────────────────────

def _apply_para_fmt(p, indent_cm=0, space_before_pt=2, space_after_pt=2):
    """Apply standard paragraph formatting."""
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.line_spacing = 1.15
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)


class _HtmlToDocxParser(HTMLParser):
    """Stream Quill / docx_importer HTML into python-docx paragraphs/tables/images."""

    def __init__(self, doc, size_pt=10.5, indent_cm=0,
                 space_before_pt=2, space_after_pt=2, prefix_text=''):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.size_pt = size_pt
        self.indent_cm = indent_cm
        self.space_before_pt = space_before_pt
        self.space_after_pt = space_after_pt

        self._para = None
        self._first_para = True
        self._prefix_text = prefix_text

        self._bold = False
        self._italic = False
        self._underline = False

        # Table accumulation
        self._in_table = False
        self._table_rows = []   # list[list[str]]
        self._cur_row = None    # list[str] | None
        self._cur_cell = None   # str | None

    def _ensure_para(self):
        if self._para is None:
            self._para = self.doc.add_paragraph()
            _apply_para_fmt(self._para, self.indent_cm,
                            self.space_before_pt, self.space_after_pt)
            if self._first_para and self._prefix_text:
                run = self._para.add_run(self._prefix_text)
                _set_run_font(run, self.size_pt)
            self._first_para = False
        return self._para

    def _close_para(self):
        self._para = None

    def handle_starttag(self, tag, attrs):
        attrs_d = dict(attrs)
        if tag == 'p':
            self._close_para()
        elif tag == 'br':
            if self._para is not None:
                run = self._para.add_run('\n')
                _set_run_font(run, self.size_pt)
        elif tag in ('strong', 'b'):
            self._bold = True
        elif tag in ('em', 'i'):
            self._italic = True
        elif tag == 'u':
            self._underline = True
        elif tag == 'img':
            src = attrs_d.get('src', '')
            m = re.search(r'/api/images/(img_[a-f0-9]+)', src)
            if m:
                self._close_para()
                self._insert_image(m.group(1))
                self._close_para()
        elif tag == 'table':
            self._close_para()
            self._in_table = True
            self._table_rows = []
        elif tag == 'tr':
            self._cur_row = []
        elif tag in ('td', 'th'):
            self._cur_cell = ''

    def handle_endtag(self, tag):
        if tag == 'p':
            self._close_para()
        elif tag in ('strong', 'b'):
            self._bold = False
        elif tag in ('em', 'i'):
            self._italic = False
        elif tag == 'u':
            self._underline = False
        elif tag == 'tr':
            if self._cur_row is not None:
                self._table_rows.append(self._cur_row)
            self._cur_row = None
        elif tag in ('td', 'th'):
            if self._cur_row is not None and self._cur_cell is not None:
                self._cur_row.append(self._cur_cell)
            self._cur_cell = None
        elif tag == 'table':
            self._in_table = False
            self._render_table()

    def handle_data(self, data):
        if self._in_table:
            if self._cur_cell is not None:
                self._cur_cell += data
            return
        # Split on newlines: each \n separates logical paragraphs
        segments = data.split('\n')
        for idx, segment in enumerate(segments):
            stripped = segment.strip()
            if stripped:
                para = self._ensure_para()
                run = para.add_run(stripped)
                run.bold = self._bold
                run.italic = self._italic
                run.underline = self._underline
                _set_run_font(run, self.size_pt)
            # Close paragraph between segments (not after the last)
            if idx < len(segments) - 1:
                self._close_para()

    def _insert_image(self, image_id):
        from app.db_models import QuestionImageModel
        rec = QuestionImageModel.query.filter_by(image_id=image_id).first()
        if rec is None:
            return
        img_path = os.path.join(_IMAGES_DIR, rec.filename)
        if not os.path.exists(img_path):
            return
        para = self._ensure_para()
        run = para.add_run()
        try:
            run.add_picture(img_path, width=Inches(3.5))
        except Exception:
            pass

    def _render_table(self):
        if not self._table_rows:
            return
        cols = max((len(r) for r in self._table_rows), default=0)
        if cols == 0:
            return
        tbl = self.doc.add_table(rows=len(self._table_rows), cols=cols)
        try:
            tbl.style = 'Table Grid'
        except Exception:
            pass
        for r_idx, row_data in enumerate(self._table_rows):
            for c_idx in range(cols):
                cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
                cell = tbl.rows[r_idx].cells[c_idx]
                cell.text = cell_text
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, self.size_pt - 1)


def _add_html_to_doc(doc, html_str: str, size_pt: float = 10.5,
                     indent_cm: float = 0,
                     space_before_pt: float = 2, space_after_pt: float = 2,
                     prefix_text: str = ''):
    """Append html_str as Word content to doc.

    Handles Quill-generated HTML (<p>, <strong>, <img>, etc.) and
    docx_importer output (plain text + <img> + <table> joined by newlines).
    Falls back to plain-text fast path if no '<' found.
    prefix_text is prepended to the very first run (e.g. '1. ').
    """
    if not html_str and not prefix_text:
        return

    if not html_str:
        p = doc.add_paragraph()
        _apply_para_fmt(p, indent_cm, space_before_pt, space_after_pt)
        run = p.add_run(prefix_text)
        _set_run_font(run, size_pt)
        return

    if '<' not in html_str:
        # Plain-text fast path
        first = True
        for line in html_str.split('\n'):
            text = line.strip()
            if not text and not first:
                continue
            p = doc.add_paragraph()
            _apply_para_fmt(p, indent_cm, space_before_pt, space_after_pt)
            run = p.add_run((prefix_text if first else '') + text)
            _set_run_font(run, size_pt)
            first = False
        return

    # HTML path
    parser = _HtmlToDocxParser(doc, size_pt=size_pt, indent_cm=indent_cm,
                                space_before_pt=space_before_pt,
                                space_after_pt=space_after_pt,
                                prefix_text=prefix_text)
    parser.feed(html_str)


def _add_labeled_html_field(doc, label: str, content: str,
                             size_pt: float = 10.5, indent_cm: float = 0.5):
    """Add a bold label followed by content that may contain HTML.

    For plain text: label + content on the same paragraph.
    For HTML: bold label paragraph, then HTML content paragraphs.
    """
    if not content:
        return
    if '<' not in content:
        p = doc.add_paragraph()
        _apply_para_fmt(p, indent_cm, 2, 0)
        run_label = p.add_run(label)
        _set_run_font(run_label, size_pt, bold=True)
        run_text = p.add_run(content)
        _set_run_font(run_text, size_pt)
    else:
        p_label = doc.add_paragraph()
        _apply_para_fmt(p_label, indent_cm, 2, 0)
        run_label = p_label.add_run(label)
        _set_run_font(run_label, size_pt, bold=True)
        _add_html_to_doc(doc, content, size_pt=size_pt, indent_cm=indent_cm,
                         space_before_pt=0, space_after_pt=2)


def generate_word_template():
    """Generate a standard Word template for question bank import.

    Reads question types from the database so that newly added custom types
    automatically appear in the template.
    """
    from app.db_models import QuestionTypeModel

    doc = Document()

    # Add title
    title = doc.add_heading('试题库导入模板', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add instructions
    doc.add_paragraph('使用说明：')
    doc.add_paragraph('1. 每道题以 [题型] 开头，题型标识符用方括号包围')
    doc.add_paragraph('2. 选择题答案紧跟在题型标识后，用方括号包围，如 [单选][A]、[多选][ABD]')
    doc.add_paragraph('3. 是非题答案只能是 [正确] 或 [错误]，无需列选项')
    doc.add_paragraph('4. 选择题选项格式：[A]选项内容（字母大写，紧跟内容，无空格）')
    doc.add_paragraph('5. 解析可选，用 <解析>...</解析> 包围')
    doc.add_paragraph('6. 参考答案用 <参考答案>...</参考答案> 包围')
    doc.add_paragraph('7. 双语出题：在 <解析> 区段内写 "英文题目:English question?" 提供英文题干')
    doc.add_paragraph('8. 双语选项：先列全部中文选项 [A][B][C][D]，再在其后统一列英文选项 [A_en][B_en][C_en][D_en]')
    doc.add_paragraph('9. 解析区段内可包含元数据行：知识点:xxx / 标签:xxx / 英文题目:xxx / 难度:easy|medium|hard')

    # Query all question types from DB
    question_types = QuestionTypeModel.query.order_by(QuestionTypeModel.id).all()

    for qt in question_types:
        doc.add_heading(f'{qt.label}示例：', level=1)

        # 是非题 special case: no options, answer is [正确] or [错误]
        if qt.name == '是非':
            example_text = (
                f"[{qt.name}][正确]\n"
                f"示例陈述句，此句表述正确。\n"
                f"<解析>\n这是一段解析\n知识点:学科知识\n标签:基础,概念\n"
                f"英文题目:Example statement that is true.\n难度:medium\n</解析>"
            )
        elif qt.has_options:
            # 多选题 shows multi-letter answer; 单选题 shows single letter
            sample_answer = 'ABD' if qt.name == '多选' else 'A'
            example_text = (
                f"[{qt.name}][{sample_answer}]\n示例题目内容？\n"
                f"[A]选项A\n[B]选项B\n[C]选项C\n[D]选项D\n"
                f"[A_en]Option A\n[B_en]Option B\n[C_en]Option C\n[D_en]Option D\n"
                f"<解析>\n这是一段解析\n知识点:学科知识\n标签:基础,概念\n"
                f"英文题目:Example question content?\n难度:medium\n</解析>"
            )
        else:
            example_text = (
                f"[{qt.name}]\n示例题目内容？\n<参考答案>\n这是一段参考答案\n</参考答案>\n"
                f"<解析>\n这是一段解析\n知识点:学科知识\n标签:基础,概念\n"
                f"英文题目:Example question content?\n难度:medium\n</解析>"
            )
        doc.add_paragraph(example_text)

    # Save to in-memory stream to avoid file-lock conflicts on Windows
    import io as _io
    stream = _io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


CHINESE_NUMBERS = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                   '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']


def _set_run_font(run, size_pt, bold=False, cn_font='宋体', en_font='Times New Roman'):
    """Set font properties on a run: size, bold, Chinese font, English font."""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = en_font
    run.element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)


def _add_styled_paragraph(doc, text, size_pt, bold=False, alignment=None,
                          cn_font='宋体', en_font='Times New Roman',
                          space_before=None, space_after=None, line_spacing=None):
    """Add a paragraph with consistent font styling."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    _set_run_font(run, size_pt, bold=bold, cn_font=cn_font, en_font=en_font)
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    return p


def _add_page_number_footer(doc):
    """Add '第 X 页（共 Y 页）' footer using PAGE/NUMPAGES field codes."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run_pre = p.add_run('第 ')
    _set_run_font(run_pre, 9)

    # PAGE field
    fld_page = parse_xml(
        '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'.format(nsdecls('w'))
    )
    p._element.append(fld_page)

    run_mid = p.add_run(' 页（共 ')
    _set_run_font(run_mid, 9)

    # NUMPAGES field
    fld_total = parse_xml(
        '<w:fldSimple {} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'.format(nsdecls('w'))
    )
    p._element.append(fld_total)

    run_end = p.add_run(' 页）')
    _set_run_font(run_end, 9)


def export_exam_to_word(exam, filepath: str, mode: str = 'zh', show_answer: bool = True):
    """Export an exam to a Word document matching formal exam paper formatting.

    Accepts an ExamModel instance (SQLAlchemy model).
    mode: 'zh' (Chinese only), 'en' (English only), 'both' (bilingual, English first)
    show_answer: True = include answers/explanations, False = questions only
    """
    from app.db_models import CourseSettingsModel

    doc = Document()

    # --- Page setup: A4, standard margins ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # --- Set default document font ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.15

    # --- Load course settings ---
    course_settings = CourseSettingsModel.query.first()
    cs = course_settings  # shorthand

    # --- Header section ---
    # Line 1: institution + semester (centered, 18pt bold)
    header_line1_parts = []
    if cs and cs.institution_name:
        header_line1_parts.append(cs.institution_name)
    if cs and cs.semester_info:
        header_line1_parts.append(cs.semester_info)
    if header_line1_parts:
        _add_styled_paragraph(doc, ' '.join(header_line1_parts), 18, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              cn_font='黑体', space_after=0)

    # Line 2: exam title (centered, 18pt bold)
    exam_title = (cs.exam_title if cs and cs.exam_title else '期末考试试卷')
    _add_styled_paragraph(doc, exam_title, 18, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          cn_font='黑体', space_before=0, space_after=6)

    # Empty line
    _add_styled_paragraph(doc, '', 10.5, space_before=0, space_after=0)

    # Course info lines (14pt)
    course_name = (cs.course_name if cs else '') or ''
    paper_label = (cs.paper_label if cs else 'A') or 'A'
    if course_name:
        _add_styled_paragraph(doc, f'课程名称：《{course_name}》（{paper_label}）卷', 14,
                              space_before=0, space_after=2)
    course_code = (cs.course_code if cs else '') or ''
    if course_code:
        _add_styled_paragraph(doc, f'课程代号：{course_code}', 14,
                              space_before=0, space_after=2)

    # Exam format + method on same line
    fmt_parts = []
    if cs and cs.exam_format:
        fmt_parts.append(f'考试形式：{cs.exam_format}')
    if cs and cs.exam_method:
        fmt_parts.append(f'考试方式：{cs.exam_method}')
    if fmt_parts:
        _add_styled_paragraph(doc, '    '.join(fmt_parts), 14,
                              space_before=0, space_after=2)

    if cs and cs.target_audience:
        _add_styled_paragraph(doc, f'使用对象：{cs.target_audience}', 14,
                              space_before=0, space_after=2)

    # Total score line
    total_score = exam.calculate_total_score()
    if total_score > 0:
        _add_styled_paragraph(doc, f'满分：{total_score} 分', 14,
                              space_before=0, space_after=2)

    # Separator line
    _add_styled_paragraph(doc, '━' * 50, 12,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=6, space_after=6)

    # --- Questions ---
    questions = exam.get_ordered_questions()
    config = json.loads(exam.config) if exam.config else {}

    # Group questions by type (preserving order)
    question_groups = {}
    group_order = []
    for question in questions:
        if question.question_type not in question_groups:
            question_groups[question.question_type] = []
            group_order.append(question.question_type)
        question_groups[question.question_type].append(question)

    # Render each question type section
    for section_idx, q_type in enumerate(group_order):
        type_questions = question_groups[q_type]
        cn_num = CHINESE_NUMBERS[section_idx] if section_idx < len(CHINESE_NUMBERS) else str(section_idx + 1)

        # Build section title with stats
        count = len(type_questions)
        type_config = config.get(q_type, {})
        points = type_config.get('points', 0)
        # Display only the sub-type for hierarchical types (e.g. '简答>材料分析' → '材料分析')
        type_label = q_type.split('>')[-1] if '>' in q_type else q_type
        if points > 0:
            section_title = f'{cn_num}、{type_label}（共{count}题，每题{points}分，共计{count * points}分）'
        else:
            section_title = f'{cn_num}、{type_label}（共{count}题）'

        _add_styled_paragraph(doc, section_title, 12, bold=True,
                              space_before=12, space_after=6)

        for i, question in enumerate(type_questions, 1):
            options_zh = json.loads(question.options) if question.options else []
            options_en = json.loads(question.options_en) if question.options_en else []
            content_en = question.content_en or ''

            # --- Question content ---
            if mode == 'both' and content_en:
                # English first (always plain text), then Chinese (may be HTML)
                p_en = doc.add_paragraph()
                p_en.paragraph_format.space_before = Pt(3)
                p_en.paragraph_format.space_after = Pt(1)
                p_en.paragraph_format.line_spacing = 1.15
                run_en = p_en.add_run(f'{i}. {content_en}')
                _set_run_font(run_en, 10.5)
                _add_html_to_doc(doc, question.content, size_pt=10.5,
                                 space_before_pt=0, space_after_pt=3,
                                 prefix_text='   ')
            elif mode == 'en':
                if content_en:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.15
                    run = p.add_run(f'{i}. {content_en}')
                    _set_run_font(run, 10.5)
                else:
                    _add_html_to_doc(doc, question.content, size_pt=10.5,
                                     space_before_pt=3, space_after_pt=3,
                                     prefix_text=f'{i}. ')
            else:  # zh or both without English content
                _add_html_to_doc(doc, question.content, size_pt=10.5,
                                 space_before_pt=3, space_after_pt=3,
                                 prefix_text=f'{i}. ')

            # --- Options ---
            if mode == 'both' and options_en:
                # English options first
                for j, opt in enumerate(options_en):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run(f'[{chr(65+j)}] {opt}')
                    _set_run_font(run, 10.5)
                # Chinese options
                for j, opt in enumerate(options_zh):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run(f'[{chr(65+j)}] {opt}')
                    _set_run_font(run, 10.5)
            elif mode == 'en':
                display_opts = options_en if options_en else options_zh
                for j, opt in enumerate(display_opts):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run(f'[{chr(65+j)}] {opt}')
                    _set_run_font(run, 10.5)
            else:  # zh
                for j, opt in enumerate(options_zh):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run(f'[{chr(65+j)}] {opt}')
                    _set_run_font(run, 10.5)

            # --- Answers ---
            if show_answer:
                if question.answer:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run(f'答案：{question.answer}')
                    _set_run_font(run, 10.5, bold=True)

                if question.reference_answer:
                    _add_labeled_html_field(doc, '参考答案：', question.reference_answer)

                if question.explanation:
                    _add_labeled_html_field(doc, '解析：', question.explanation)

    # --- Page footer with page numbers ---
    _add_page_number_footer(doc)

    # Save the document
    doc.save(filepath)


def parse_question_template(content: str) -> list:
    """Parse questions from the template format"""
    questions = []
    
    # Split content by question markers
    # This is a simplified parser - a full implementation would be more complex
    lines = content.split('\n')
    
    current_question = None
    current_section = None  # 'question', 'options', 'answer', 'explanation', 'reference_answer'
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('[') and ']' in line[:20] and not (len(line) >= 3 and line[1].isalpha() and line[2] == ']') and not re.match(r'^\[[A-Z]_en\]', line):
            # Found a new question
            if current_question:
                questions.append(current_question)
            
            # Parse question type and answer
            bracket_end = line.find(']')
            question_info = line[1:bracket_end]
            
            # Use the full identifier as the type (e.g. '简答>材料分析', not just '简答')
            q_type = question_info

            # Determine if there's an answer in the bracket
            answer = None
            if len(line) > bracket_end + 1:
                # Look for answer after the bracket
                remaining = line[bracket_end + 1:].strip()
                if remaining.startswith('[') and ']' in remaining:
                    answer_end = remaining.find(']')
                    answer = remaining[1:answer_end]

            current_question = {
                'type': q_type,
                'content': '',
                'options': [],
                'options_en': [],
                'content_en': '',
                'subject': '',
                'knowledge_point': '',
                'tags': '',
                'difficulty': '',
                'answer': answer,
                'reference_answer': '',
                'explanation': ''
            }
            
            # The next line should be the question content.
            # After the type bracket, there may be an answer bracket like [B] or [ABD].
            # We must skip that bracket and only capture actual text content.
            rest_of_line = line[bracket_end + 1:].strip()
            # If rest_of_line is solely an answer bracket (e.g. "[B]", "[ABD]", "[正确]", "[错误]"),
            # the actual question content is on the next line.
            _answer_only = re.match(r'^\[([A-Z]{1,4}|正确|错误)\]$', rest_of_line)
            _answer_prefix = re.match(r'^\[([A-Z]{1,4}|正确|错误)\]', rest_of_line)
            if rest_of_line and _answer_only:
                # Only an answer bracket remains — read content from the next line
                i += 1
                if i < len(lines):
                    current_question['content'] = lines[i].strip()
            elif rest_of_line and _answer_prefix:
                # Answer bracket followed by inline content — strip the bracket
                current_question['content'] = re.sub(r'^\[([A-Z]{1,4}|正确|错误)\]\s*', '', rest_of_line)
            elif rest_of_line:
                current_question['content'] = rest_of_line
            else:
                i += 1
                if i < len(lines):
                    current_question['content'] = lines[i].strip()
        
        elif current_question and re.match(r'^\[[A-Z]_en\]', line):
            # English option line like [A_en] option text
            option_text = re.sub(r'^\[[A-Z]_en\]\s*', '', line)
            current_question['options_en'].append(option_text)

        elif current_question and line.startswith('[') and len(line) >= 3 and line[1].isalpha() and line[2] == ']':
            # Option line like [A] option text
            option_letter = line[1]
            option_text = line[3:].strip()
            current_question['options'].append(option_text)
        
        elif line.startswith('<参考答案>'):
            current_section = 'reference_answer'
            current_question['reference_answer'] = ''
            # Check if the content is on the same line
            content_start = line.find('<参考答案>') + len('<参考答案>')
            if content_start < len(line):
                current_question['reference_answer'] = line[content_start:]
        
        elif line.startswith('<解析>'):
            current_section = 'explanation'
            current_question['explanation'] = ''
            # Check if the content is on the same line
            content_start = line.find('<解析>') + len('<解析>')
            if content_start < len(line):
                current_question['explanation'] = line[content_start:]
        
        elif line.startswith('</参考答案>') or line.startswith('</解析>'):
            current_section = None
        
        elif current_section == 'reference_answer':
            current_question['reference_answer'] += '\n' + line
        
        elif current_section == 'explanation':
            if re.match(r'^科目[:：]', line):
                current_question['subject'] = re.sub(r'^科目[:：]\s*', '', line)
            elif re.match(r'^知识点[:：]', line):
                current_question['knowledge_point'] = re.sub(r'^知识点[:：]\s*', '', line)
            elif re.match(r'^标签[:：]', line):
                current_question['tags'] = re.sub(r'^标签[:：]\s*', '', line)
            elif re.match(r'^英文题目[:：]', line):
                current_question['content_en'] = re.sub(r'^英文题目[:：]\s*', '', line)
            elif re.match(r'^难度[:：]', line):
                current_question['difficulty'] = re.sub(r'^难度[:：]\s*', '', line)
            else:
                current_question['explanation'] += '\n' + line
        
        elif current_question and current_section is None and not line.startswith('[') and not line.startswith('<') and not line.startswith('('):
            # Additional content for the question
            if not current_question['content']:
                current_question['content'] = line
            else:
                current_question['content'] += '\n' + line
        
        i += 1
    
    # Add the last question if it exists
    if current_question:
        questions.append(current_question)
    
    return questions