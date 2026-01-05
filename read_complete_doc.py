"""
更全面的Word文档读取脚本
用于提取林业经济学试卷的完整内容
"""

from docx import Document
from docx.table import Table
import os

def clean_text(text):
    """
    清理文本中的特殊字符
    """
    if text is None:
        return ""
    # 移除不可见字符和特殊字符
    cleaned = text.replace('\x00', '').replace('\ufffd', '').strip()
    return cleaned

def read_complete_document(file_path):
    """
    完整读取文档内容，包括所有段落和表格
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return
    
    try:
        doc = Document(file_path)
        print(f"正在完整读取文档: {file_path}")
        print("="*60)
        
        # 文档属性
        print("文档属性:")
        print("-" * 30)
        core_props = doc.core_properties
        print(f"标题: {clean_text(core_props.title) if core_props.title else '无'}")
        print(f"作者: {clean_text(core_props.author) if core_props.author else '无'}")
        print(f"主题: {clean_text(core_props.subject) if core_props.subject else '无'}")
        print(f"创建时间: {core_props.created if core_props.created else '无'}")
        print(f"修改时间: {core_props.modified if core_props.modified else '无'}")
        
        print("\n" + "="*60)
        
        # 读取所有段落
        print("文档内容:")
        print("-" * 30)
        for i, paragraph in enumerate(doc.paragraphs):
            text = clean_text(paragraph.text)
            if text:  # 只显示非空段落
                # 获取对齐方式
                alignment = paragraph.paragraph_format.alignment
                alignment_text = ""
                if alignment == 0: alignment_text = "左对齐"
                elif alignment == 1: alignment_text = "居中"
                elif alignment == 2: alignment_text = "右对齐"
                elif alignment == 3: alignment_text = "两端对齐"
                
                print(f"段落 {i+1} ({alignment_text}): {text}")
                
                # 检查段落中的格式化文本
                if paragraph.runs:
                    for j, run in enumerate(paragraph.runs):
                        run_text = clean_text(run.text)
                        if run_text:
                            formatting = []
                            if run.bold: formatting.append("粗体")
                            if run.italic: formatting.append("斜体")
                            if run.underline: formatting.append("下划线")
                            if run.font.name: formatting.append(f"字体:{run.font.name}")
                            if run.font.size: 
                                size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                                formatting.append(f"大小:{size}")
                            
                            if formatting:
                                print(f"  - Run {j+1} [{' '.join(formatting)}]: '{run_text}'")
        
        print("\n" + "="*60)
        
        # 读取所有表格内容
        print("表格内容:")
        print("-" * 30)
        if doc.tables:
            for i, table in enumerate(doc.tables):
                print(f"表格 {i+1}:")
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = clean_text(cell.text)
                        row_data.append(cell_text)
                    # 只显示包含内容的行
                    if any(cell_text.strip() for cell_text in row_data):
                        print(f"  行 {row_idx+1}: {row_data}")
        else:
            print("文档中没有找到表格。")
        
        print("\n" + "="*60)
        
        # 统计信息
        print("文档统计:")
        print("-" * 30)
        print(f"总段落数: {len(doc.paragraphs)}")
        print(f"总表格数: {len(doc.tables)}")
        
        # 计算总字符数
        total_chars = 0
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            total_chars += len(text)
        print(f"总字符数(不含格式): {total_chars}")
        
    except Exception as e:
        print(f"读取文档时发生错误: {str(e)}")
        import traceback
        traceback.print_exc()

def extract_pages_info(file_path):
    """
    提取页面信息，识别试卷的各个部分
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return

    try:
        doc = Document(file_path)
        page_info = {}
        current_page = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = clean_text(paragraph.text)
            if "第" in text and "页(共" in text and "页)" in text:
                # 检测页码
                try:
                    page_num = int([c for c in text if c.isdigit()][0])
                    current_page = page_num
                    page_info[current_page] = []
                except:
                    continue
            elif current_page > 0:
                if current_page not in page_info:
                    page_info[current_page] = []
                if text.strip():
                    page_info[current_page].append(text)
        
        print("\n按页整理的内容:")
        print("-" * 30)
        for page_num in sorted(page_info.keys()):
            print(f"第 {page_num} 页:")
            for j, content in enumerate(page_info[page_num]):
                print(f"  {j+1}. {content}")
            print()

    except Exception as e:
        print(f"提取页面信息时发生错误: {str(e)}")

if __name__ == "__main__":
    file_path = r"D:\code\试卷\林业经济学（双语）试卷A-打印.docx"
    
    # 完整读取文档
    read_complete_document(file_path)
    
    # 提取页面信息
    extract_pages_info(file_path)