"""
深度读取Word文档内容的脚本
特别针对包含表格的试卷文档
"""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import os

def clean_text(text):
    """
    清理文本中的特殊字符
    """
    if text is None:
        return ""
    # 移除不可见字符和特殊字符
    cleaned = text.replace('\x00', '').replace('\ufffd', '').strip()
    return cleaned

def iterate_all_elements(doc):
    """
    遍历文档中的所有元素，包括段落和表格
    """
    elements = []
    
    # 首先处理所有段落
    for i, paragraph in enumerate(doc.paragraphs):
        text = clean_text(paragraph.text)
        if text:
            elements.append(('paragraph', i, paragraph))
    
    # 然后处理所有表格中的内容
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 获取单元格中的段落
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    text = clean_text(paragraph.text)
                    if text:
                        elements.append(('table_cell', (table_idx, row_idx, col_idx, para_idx), paragraph))
    
    return elements

def read_document_detailed(file_path):
    """
    详细读取文档内容，包括表格中的内容
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return
    
    try:
        doc = Document(file_path)
        print(f"正在深度读取文档: {file_path}")
        print("="*60)
        
        # 文档属性
        print("文档属性:")
        print("-" * 30)
        core_props = doc.core_properties
        print(f"标题: {clean_text(core_props.title) if core_props.title else '无'}")
        print(f"作者: {clean_text(core_props.author) if core_props.author else '无'}")
        print(f"主题: {clean_text(core_props.subject) if core_props.subject else '无'}")
        print(f"创建时间: {core_props.created if core_props.created else '无'}")
        print(f"修改时间: {core_props.modified if core_props.modified else '无'}")
        
        print("\n文档结构:")
        print("-" * 30)
        print(f"总段落数: {len(doc.paragraphs)}")
        print(f"总表格数: {len(doc.tables)}")
        
        # 统计表格中的内容
        total_cells = 0
        for table in doc.tables:
            total_cells += len(table.rows) * len(table.columns) if table.rows else 0
        print(f"总单元格数: {total_cells}")
        
        print("\n文档内容:")
        print("-" * 30)
        
        # 获取所有元素
        elements = iterate_all_elements(doc)
        
        # 按类型分类处理
        paragraphs = []
        table_cells = []
        
        for elem_type, idx, element in elements:
            if elem_type == 'paragraph':
                paragraphs.append((idx, element))
            elif elem_type == 'table_cell':
                table_cells.append((idx, element))
        
        # 输出段落内容
        print(f"段落数: {len(paragraphs)}")
        for i, (idx, paragraph) in enumerate(paragraphs):
            text = clean_text(paragraph.text)
            alignment = paragraph.paragraph_format.alignment
            alignment_text = {0: "左对齐", 1: "居中", 2: "右对齐", 3: "两端对齐"}.get(alignment, "未知")
            
            print(f"  段落 {idx+1} ({alignment_text}): {text}")
            
            # 检查段落中的格式化文本
            if paragraph.runs:
                for j, run in enumerate(paragraph.runs):
                    run_text = clean_text(run.text)
                    if run_text:
                        formatting = []
                        if run.bold: formatting.append("粗体")
                        if run.italic: formatting.append("斜体")
                        if run.underline: formatting.append("下划线")
                        if run.font.name: formatting.append(f"字体:{run.font.name}")
                        if run.font.size: 
                            size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                            formatting.append(f"大小:{size}")
                        
                        if formatting:
                            print(f"    - Run {j+1} [{' '.join(formatting)}]: '{run_text}'")
        
        # 输出表格内容
        print(f"\n表格单元格数: {len(table_cells)}")
        for i, ((table_idx, row_idx, col_idx, para_idx), paragraph) in enumerate(table_cells):
            text = clean_text(paragraph.text)
            alignment = paragraph.paragraph_format.alignment
            alignment_text = {0: "左对齐", 1: "居中", 2: "右对齐", 3: "两端对齐"}.get(alignment, "未知")
            
            print(f"  表格{table_idx+1}第{row_idx+1}行第{col_idx+1}列段落{para_idx+1} ({alignment_text}): {text}")
            
            # 检查表格段落中的格式化文本
            if paragraph.runs:
                for j, run in enumerate(paragraph.runs):
                    run_text = clean_text(run.text)
                    if run_text:
                        formatting = []
                        if run.bold: formatting.append("粗体")
                        if run.italic: formatting.append("斜体")
                        if run.underline: formatting.append("下划线")
                        if run.font.name: formatting.append(f"字体:{run.font.name}")
                        if run.font.size: 
                            size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                            formatting.append(f"大小:{size}")
                        
                        if formatting:
                            print(f"    - Run {j+1} [{' '.join(formatting)}]: '{run_text}'")
        
        print("\n表格结构:")
        print("-" * 30)
        for i, table in enumerate(doc.tables):
            print(f"表格 {i+1}: {len(table.rows)} 行 x {len(table.columns) if table.rows else 0} 列")
            for j, row in enumerate(table.rows):
                row_content = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    if cell_text:
                        # 只显示前50个字符，避免输出过长
                        display_text = cell_text[:50] + "..." if len(cell_text) > 50 else cell_text
                        row_content.append(display_text)
                if any(row_content):  # 只显示包含内容的行
                    print(f"  行 {j+1}: {row_content}")
        
    except Exception as e:
        print(f"读取文档时发生错误: {str(e)}")
        import traceback
        traceback.print_exc()

def extract_document_text_only(file_path):
    """
    提取文档的纯文本内容
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return ""
    
    try:
        doc = Document(file_path)
        content = []
        
        # 添加段落内容
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if text:
                content.append(text)
        
        # 添加表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    if cell_text:
                        content.append(cell_text)
        
        return "\n".join(content)
    
    except Exception as e:
        print(f"提取文本时发生错误: {str(e)}")
        return ""

if __name__ == "__main__":
    file_path = r"D:\code\试卷\林业经济学（双语）试卷A-打印.docx"
    
    # 详细读取文档
    read_document_detailed(file_path)
    
    print("\n" + "="*60)
    print("纯文本内容:")
    print("-" * 30)
    
    # 提取纯文本内容
    text_content = extract_document_text_only(file_path)
    print(text_content)