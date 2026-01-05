"""
林业经济学（双语）试题生成器 - GUI版本
包含可视化界面、题库管理、试卷生成等功能
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import json
import os
from datetime import datetime
import random
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import csv

class ExamGeneratorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("林业经济学（双语）试题生成器")
        self.root.geometry("1200x800")
        
        # 初始化题库
        self.question_database = {
            "single_choice_questions": [],
            "true_false_questions": [],
            "essay_questions": [],
            "calculation_questions": []
        }
        
        # 题目使用历史
        self.question_usage = {}
        
        # 创建界面
        self.create_widgets()
        
    def create_widgets(self):
        # 创建主框架
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建笔记本控件（Tab）
        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        
        # 题库管理标签页
        self.create_question_bank_tab(notebook)
        
        # 试卷生成标签页
        self.create_exam_generation_tab(notebook)
        
        # 模板下载标签页
        self.create_template_download_tab(notebook)
        
    def create_question_bank_tab(self, parent):
        # 创建题库管理界面
        question_bank_frame = ttk.Frame(parent)
        parent.add(question_bank_frame, text="题库管理")
        
        # 控制框架
        control_frame = ttk.Frame(question_bank_frame)
        control_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 模板下载按钮
        download_template_btn = ttk.Button(control_frame, text="下载试题模板", command=self.download_template)
        download_template_btn.pack(side=tk.LEFT, padx=5)
        
        # 导入试题按钮
        import_questions_btn = ttk.Button(control_frame, text="导入试题", command=self.import_questions)
        import_questions_btn.pack(side=tk.LEFT, padx=5)
        
        # 导出题库按钮
        export_questions_btn = ttk.Button(control_frame, text="导出题库", command=self.export_questions)
        export_questions_btn.pack(side=tk.LEFT, padx=5)
        
        # 题库显示框架
        tree_frame = ttk.Frame(question_bank_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 创建树形视图
        columns = ("ID", "题型", "难度", "知识点", "使用次数", "最后使用")
        self.question_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
        
        # 设置列标题
        for col in columns:
            self.question_tree.heading(col, text=col)
            self.question_tree.column(col, width=100)
        
        # 添加滚动条
        tree_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.question_tree.yview)
        self.question_tree.configure(yscrollcommand=tree_scroll.set)
        
        self.question_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 按钮框架
        button_frame = ttk.Frame(question_bank_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 添加题目按钮
        add_btn = ttk.Button(button_frame, text="添加题目", command=self.add_question_dialog)
        add_btn.pack(side=tk.LEFT, padx=5)
        
        # 编辑题目按钮
        edit_btn = ttk.Button(button_frame, text="编辑题目", command=self.edit_question_dialog)
        edit_btn.pack(side=tk.LEFT, padx=5)
        
        # 删除题目按钮
        delete_btn = ttk.Button(button_frame, text="删除题目", command=self.delete_question)
        delete_btn.pack(side=tk.LEFT, padx=5)
        
        # 刷新按钮
        refresh_btn = ttk.Button(button_frame, text="刷新", command=self.refresh_question_list)
        refresh_btn.pack(side=tk.LEFT, padx=5)
        
    def create_exam_generation_tab(self, parent):
        # 创建试卷生成界面
        exam_gen_frame = ttk.Frame(parent)
        parent.add(exam_gen_frame, text="试卷生成")
        
        # 配置框架
        config_frame = ttk.LabelFrame(exam_gen_frame, text="试卷配置")
        config_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 基本信息
        basic_info_frame = ttk.Frame(config_frame)
        basic_info_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(basic_info_frame, text="学校:").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.school_var = tk.StringVar(value="中南财经政法大学")
        ttk.Entry(basic_info_frame, textvariable=self.school_var, width=20).grid(row=0, column=1, sticky=tk.W, padx=5)
        
        ttk.Label(basic_info_frame, text="学年:").grid(row=0, column=2, sticky=tk.W, padx=5)
        self.academic_year_var = tk.StringVar(value="2025–2026学年第1学期")
        ttk.Entry(basic_info_frame, textvariable=self.academic_year_var, width=20).grid(row=0, column=3, sticky=tk.W, padx=5)
        
        ttk.Label(basic_info_frame, text="课程名称:").grid(row=1, column=0, sticky=tk.W, padx=5)
        self.course_name_var = tk.StringVar(value="林业经济学")
        ttk.Entry(basic_info_frame, textvariable=self.course_name_var, width=20).grid(row=1, column=1, sticky=tk.W, padx=5)
        
        ttk.Label(basic_info_frame, text="试卷类型:").grid(row=1, column=2, sticky=tk.W, padx=5)
        self.paper_type_var = tk.StringVar(value="A")
        ttk.Combobox(basic_info_frame, textvariable=self.paper_type_var, values=["A", "B", "C", "D"], width=18).grid(row=1, column=3, sticky=tk.W, padx=5)
        
        # 题型配置
        question_config_frame = ttk.LabelFrame(config_frame, text="题型配置")
        question_config_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 单选题配置
        ttk.Label(question_config_frame, text="单选题:").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.sc_count_var = tk.StringVar(value="5")
        self.sc_score_var = tk.StringVar(value="2")
        ttk.Entry(question_config_frame, textvariable=self.sc_count_var, width=5).grid(row=0, column=1, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="题 ×").grid(row=0, column=2, sticky=tk.W, padx=2)
        ttk.Entry(question_config_frame, textvariable=self.sc_score_var, width=5).grid(row=0, column=3, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="分/题").grid(row=0, column=4, sticky=tk.W, padx=5)
        
        # 是非题配置
        ttk.Label(question_config_frame, text="是非题:").grid(row=1, column=0, sticky=tk.W, padx=5)
        self.tf_count_var = tk.StringVar(value="5")
        self.tf_score_var = tk.StringVar(value="2")
        ttk.Entry(question_config_frame, textvariable=self.tf_count_var, width=5).grid(row=1, column=1, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="题 ×").grid(row=1, column=2, sticky=tk.W, padx=2)
        ttk.Entry(question_config_frame, textvariable=self.tf_score_var, width=5).grid(row=1, column=3, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="分/题").grid(row=1, column=4, sticky=tk.W, padx=5)
        
        # 论述题配置
        ttk.Label(question_config_frame, text="论述题:").grid(row=2, column=0, sticky=tk.W, padx=5)
        self.es_count_var = tk.StringVar(value="2")
        self.es_score_var = tk.StringVar(value="15")
        ttk.Entry(question_config_frame, textvariable=self.es_count_var, width=5).grid(row=2, column=1, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="题 ×").grid(row=2, column=2, sticky=tk.W, padx=2)
        ttk.Entry(question_config_frame, textvariable=self.es_score_var, width=5).grid(row=2, column=3, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="分/题").grid(row=2, column=4, sticky=tk.W, padx=5)
        
        # 计算题配置
        ttk.Label(question_config_frame, text="计算题:").grid(row=3, column=0, sticky=tk.W, padx=5)
        self.calc_count_var = tk.StringVar(value="1")
        self.calc_score_var = tk.StringVar(value="20")
        ttk.Entry(question_config_frame, textvariable=self.calc_count_var, width=5).grid(row=3, column=1, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="题 ×").grid(row=3, column=2, sticky=tk.W, padx=2)
        ttk.Entry(question_config_frame, textvariable=self.calc_score_var, width=5).grid(row=3, column=3, sticky=tk.W, padx=2)
        ttk.Label(question_config_frame, text="分/题").grid(row=3, column=4, sticky=tk.W, padx=5)
        
        # 预览和生成按钮
        button_frame = ttk.Frame(exam_gen_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=5)
        
        preview_btn = ttk.Button(button_frame, text="预览试卷", command=self.preview_exam)
        preview_btn.pack(side=tk.LEFT, padx=5)
        
        generate_btn = ttk.Button(button_frame, text="生成试卷", command=self.generate_exam)
        generate_btn.pack(side=tk.LEFT, padx=5)
        
        # 预览文本框
        preview_frame = ttk.LabelFrame(exam_gen_frame, text="试卷预览")
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.preview_text = scrolledtext.ScrolledText(preview_frame, height=20)
        self.preview_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
    def create_template_download_tab(self, parent):
        # 创建模板下载界面
        template_frame = ttk.Frame(parent)
        parent.add(template_frame, text="模板下载")
        
        # 说明文本
        instruction_text = """
        使用说明：
        1. 点击下方按钮下载试题模板
        2. 在模板中填写您的题目内容
        3. 保存为CSV格式
        4. 返回"题库管理"页签，点击"导入试题"
        5. 在弹出的对话框中选择您保存的CSV文件
        6. 系统会自动将试题添加到题库中
        
        模板格式说明：
        单选题模板包含列：ID, type, difficulty, chinese_stem, english_stem, option_A_ch, option_A_en, option_B_ch, option_B_en, option_C_ch, option_C_en, option_D_ch, option_D_en, correct_answer, knowledge_point, tags
        是非题模板包含列：ID, type, chinese_stem, english_stem, correct_answer, explanation_ch, explanation_en, knowledge_point, tags
        论述题模板包含列：ID, type, chinese_stem, english_stem, scoring_guide_ch, scoring_guide_en, knowledge_point, tags
        计算题模板包含列：ID, type, context_ch, context_en, alternatives_ch, alternatives_en, parameters, requirements_ch, requirements_en, knowledge_point, tags
        """
        
        instruction_label = ttk.Label(template_frame, text=instruction_text, justify=tk.LEFT)
        instruction_label.pack(padx=10, pady=10)
        
        # 下载按钮框架
        download_btn_frame = ttk.Frame(template_frame)
        download_btn_frame.pack(pady=20)
        
        single_choice_btn = ttk.Button(download_btn_frame, text="下载单选题模板", command=lambda: self.download_question_template("single_choice"))
        single_choice_btn.pack(side=tk.LEFT, padx=10)
        
        true_false_btn = ttk.Button(download_btn_frame, text="下载是非题模板", command=lambda: self.download_question_template("true_false"))
        true_false_btn.pack(side=tk.LEFT, padx=10)
        
        essay_btn = ttk.Button(download_btn_frame, text="下载论述题模板", command=lambda: self.download_question_template("essay"))
        essay_btn.pack(side=tk.LEFT, padx=10)
        
        calculation_btn = ttk.Button(download_btn_frame, text="下载计算题模板", command=lambda: self.download_question_template("calculation"))
        calculation_btn.pack(side=tk.LEFT, padx=10)
        
    def download_template(self):
        """下载通用试题模板"""
        # 为每种题型创建CSV模板
        templates = {
            "single_choice_template.csv": [
                ["ID", "type", "difficulty", "chinese_stem", "english_stem", 
                 "option_A_ch", "option_A_en", "option_B_ch", "option_B_en", 
                 "option_C_ch", "option_C_en", "option_D_ch", "option_D_en", 
                 "correct_answer", "knowledge_point", "tags"],
                ["sc_001", "single_choice", "medium", "题目中文内容", "Question English Content",
                 "选项A中文", "Option A English", "选项B中文", "Option B English",
                 "选项C中文", "Option C English", "选项D中文", "Option D English",
                 "A", "知识点名称", "tag1,tag2"]
            ],
            "true_false_template.csv": [
                ["ID", "type", "chinese_stem", "english_stem", "correct_answer", 
                 "explanation_ch", "explanation_en", "knowledge_point", "tags"],
                ["tf_001", "true_false", "题目中文内容", "Question English Content", 
                 "T", "解释中文", "Explanation English", "知识点名称", "tag1,tag2"]
            ],
            "essay_template.csv": [
                ["ID", "type", "chinese_stem", "english_stem", "scoring_guide_ch", 
                 "scoring_guide_en", "knowledge_point", "tags"],
                ["es_001", "essay", "题目中文内容", "Question English Content", 
                 "评分标准中文", "Scoring Guide English", "知识点名称", "tag1,tag2"]
            ],
            "calculation_template.csv": [
                ["ID", "type", "context_ch", "context_en", "alternatives_ch", 
                 "alternatives_en", "parameters", "requirements_ch", "requirements_en", 
                 "knowledge_point", "tags"],
                ["calc_001", "calculation", "背景中文", "Background English", 
                 "选项中文", "Options English", "参数", "要求中文", "Requirements English", 
                 "知识点名称", "tag1,tag2"]
            ]
        }
        
        # 询问用户保存位置
        messagebox.showinfo("下载模板", "请选择保存模板文件的目录")
        directory = filedialog.askdirectory()
        if not directory:
            return
            
        for filename, content in templates.items():
            filepath = os.path.join(directory, filename)
            with open(filepath, 'w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                writer.writerows(content)
        
        messagebox.showinfo("成功", "模板文件已下载完成！")
    
    def download_question_template(self, question_type):
        """下载特定题型的模板"""
        templates = {
            "single_choice": [
                ["ID", "type", "difficulty", "chinese_stem", "english_stem", 
                 "option_A_ch", "option_A_en", "option_B_ch", "option_B_en", 
                 "option_C_ch", "option_C_en", "option_D_ch", "option_D_en", 
                 "correct_answer", "knowledge_point", "tags"],
                ["sc_001", "single_choice", "medium", "Merit goods在林业经济学中指的是什么？（）", "What does \"Merit goods\" refer to in forest economics? ()",
                 "社会价值低于私人消费者价值的商品", "Merchandise with social value lower than private consumer value", 
                 "社会价值超过私人消费者价值的商品", "Merchandise with social value exceeding private consumer value",
                 "仅具有私人价值的商品", "Merchandise with private value only", 
                 "仅具有社会价值的商品", "Merchandise with social value only",
                 "B", "Merit goods definition", "basic,definition,public goods"]
            ],
            "true_false": [
                ["ID", "type", "chinese_stem", "english_stem", "correct_answer", 
                 "explanation_ch", "explanation_en", "knowledge_point", "tags"],
                ["tf_001", "true_false", "\"Contingent valuation method\"是一种了解消费者偏好的显示型（Revealed preference)方法。", 
                 "The Contingent Valuation Method (CVM) is a revealed preference approach for eliciting consumer preferences.", 
                 "F", "CVM是陈述型（Stated preference）方法，而非显示型（Revealed preference）方法。", 
                 "CVM is a stated preference approach, not a revealed preference approach.", 
                 "CVM methodology", "valuation,method,preference"]
            ],
            "essay": [
                ["ID", "type", "chinese_stem", "english_stem", "scoring_guide_ch", 
                 "scoring_guide_en", "knowledge_point", "tags"],
                ["es_001", "essay", "有哪些不同类型的\"外部性\"？请列举出至少三种，并举例分析每种外部性分别是如何产生的？有何应对措施？", 
                 "What are the different types of externalities? Please identify at least three distinct categories and provide illustrative examples demonstrating how each type of externality arises. What policy interventions or remedial measures can be implemented to address them?",
                 "识别三种以上外部性类型 (5分);每种类型提供具体例子 (5分);分析产生机制 (3分);提出应对措施 (2分)",
                 "Identify three or more types of externalities (5 points);Provide specific examples for each type (5 points);Analyze generation mechanisms (3 points);Propose countermeasures (2 points)",
                 "Types of externalities", "externality,classification,policy"]
            ],
            "calculation": [
                ["ID", "type", "context_ch", "context_en", "alternatives_ch", 
                 "alternatives_en", "parameters", "requirements_ch", "requirements_en", 
                 "knowledge_point", "tags"],
                ["calc_001", "calculation", "某村庄有一片荒地，如果自然更新，每30年可以产出价值10.9万元的木材，采伐成本为1万元。村委会现在要从以下投资决策中做出选择：", 
                 "A village possesses a tract of wasteland. Under natural regeneration, it would yield timber valued at ¥109,000 every 30 years, with harvesting costs of ¥10,000. The village committee must now choose among the following investment alternatives:",
                 "投资20万元种植松树，每20年产出木材价值70万元的木材，采伐成本为6万元，每年的管理成本为1000元。#将荒地直接承包给外来经营者，对方承诺经过5年投产期后，每年向村集体交付1万元租金#投资170万元，将荒地改造为生态旅游景区，1年后开始获得稳定回报，预期每年能带来10万元净回报",
                 "Invest ¥200,000 to establish a pine plantation, generating timber valued at ¥700,000 every 20 years, with harvesting costs of ¥60,000 and annual management costs of ¥1,000.#Lease the wasteland directly to an external operator, who commits to paying annual rent of ¥10,000 to the village collective after a 5-year establishment period.#Invest ¥1,700,000 to transform the wasteland into an eco-tourism destination, which would begin generating stable returns after one year, with anticipated annual net returns of ¥100,000.",
                 "折现率:0.05;时间系数:5年=1.3,10年=1.6,20年=2.6,30年=4.3;公式:F=A*[(1+i)^n-1]/i",
                 "请计算不同投资决策的净现值，并根据净现值考虑选择哪种投资决策？#假如你是村委会的决策人，在投资决策时，除了净现值以外，你还将考虑哪些因素？分析在考虑这些因素后，你会如何改变决策？",
                 "Calculate the net present value (NPV) of each investment alternative and determine which option should be selected based on NPV analysis.#If you were the decision-maker for the village committee, what additional factors beyond NPV would you consider in the investment decision? Analyze how consideration of these factors might alter your decision.",
                 "NPV calculation", "NPV,investment,decision"]
            ]
        }
        
        messagebox.showinfo("下载模板", f"请选择保存{question_type}模板文件的位置")
        filepath = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            initialfile=f"{question_type}_template.csv"
        )
        
        if filepath:
            with open(filepath, 'w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                writer.writerows(templates[question_type])
            messagebox.showinfo("成功", f"{question_type}模板已下载完成！")
    
    def import_questions(self):
        """导入试题"""
        filepath = filedialog.askopenfilename(
            title="选择试题文件",
            filetypes=[("Word files", "*.docx"), ("CSV files", "*.csv"), ("JSON files", "*.json"), ("All files", "*.*")]
        )

        if not filepath:
            return

        try:
            if filepath.endswith('.docx'):
                self.import_from_word(filepath)
            elif filepath.endswith('.csv'):
                self.import_from_csv(filepath)
            elif filepath.endswith('.json'):
                self.import_from_json(filepath)

            messagebox.showinfo("成功", "试题导入完成！")
            self.refresh_question_list()
        except Exception as e:
            messagebox.showerror("错误", f"导入试题失败: {str(e)}")
    
    def import_from_csv(self, filepath):
        """从CSV文件导入试题"""
        with open(filepath, 'r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            for row in reader:
                question_type = row.get('type', '').lower()

                if question_type == 'single_choice':
                    question = {
                        "id": row.get('id', f"sc_{len(self.question_database['single_choice_questions']) + 1}"),
                        "type": "single_choice",
                        "difficulty": row.get('difficulty', 'medium'),
                        "chinese_stem": row.get('chinese_stem', ''),
                        "english_stem": row.get('english_stem', ''),
                        "options": {
                            "A": row.get('option_A_ch', ''),
                            "A_en": row.get('option_A_en', ''),
                            "B": row.get('option_B_ch', ''),
                            "B_en": row.get('option_B_en', ''),
                            "C": row.get('option_C_ch', ''),
                            "C_en": row.get('option_C_en', ''),
                            "D": row.get('option_D_ch', ''),
                            "D_en": row.get('option_D_en', '')
                        },
                        "correct_answer": row.get('correct_answer', ''),
                        "knowledge_point": row.get('knowledge_point', ''),
                        "tags": [tag.strip() for tag in row.get('tags', '').split(',') if tag.strip()]
                    }
                    self.question_database["single_choice_questions"].append(question)

                elif question_type == 'true_false':
                    question = {
                        "id": row.get('id', f"tf_{len(self.question_database['true_false_questions']) + 1}"),
                        "type": "true_false",
                        "chinese_stem": row.get('chinese_stem', ''),
                        "english_stem": row.get('english_stem', ''),
                        "correct_answer": row.get('correct_answer', ''),
                        "explanation": row.get('explanation', ''),
                        "explanation_en": row.get('explanation_en', ''),
                        "knowledge_point": row.get('knowledge_point', ''),
                        "tags": [tag.strip() for tag in row.get('tags', '').split(',') if tag.strip()]
                    }
                    self.question_database["true_false_questions"].append(question)

                elif question_type == 'essay':
                    scoring_guide = row.get('scoring_guide', '')
                    scoring_guide_en = row.get('scoring_guide_en', '')
                    question = {
                        "id": row.get('id', f"es_{len(self.question_database['essay_questions']) + 1}"),
                        "type": "essay",
                        "chinese_stem": row.get('chinese_stem', ''),
                        "english_stem": row.get('english_stem', ''),
                        "scoring_guide": scoring_guide.split(';') if scoring_guide else [],
                        "scoring_guide_en": scoring_guide_en.split(';') if scoring_guide_en else [],
                        "knowledge_point": row.get('knowledge_point', ''),
                        "tags": [tag.strip() for tag in row.get('tags', '').split(',') if tag.strip()]
                    }
                    self.question_database["essay_questions"].append(question)

                elif question_type == 'calculation':
                    # 解析context
                    context_ch = row.get('context_chinese', '')
                    context_en = row.get('context_english', '')

                    # 解析parameters
                    params_str = row.get('parameters', '')
                    parameters = {}
                    if params_str:
                        for param in params_str.split(';'):
                            if ':' in param:
                                key, value = param.split(':', 1)
                                parameters[key.strip()] = value.strip()

                    # 解析requirements
                    req_ch_str = row.get('requirements_chinese', '')
                    req_en_str = row.get('requirements_english', '')
                    requirements_ch = req_ch_str.split(';') if req_ch_str else []
                    requirements_en = req_en_str.split(';') if req_en_str else []

                    question = {
                        "id": row.get('id', f"calc_{len(self.question_database['calculation_questions']) + 1}"),
                        "type": "calculation",
                        "context": {
                            "chinese": context_ch,
                            "english": context_en
                        },
                        "parameters": parameters,
                        "requirements": {
                            "chinese": requirements_ch,
                            "english": requirements_en
                        },
                        "knowledge_point": row.get('knowledge_point', ''),
                        "tags": [tag.strip() for tag in row.get('tags', '').split(',') if tag.strip()]
                    }
                    self.question_database["calculation_questions"].append(question)

    def import_from_word(self, filepath):
        """从Word文档导入试题"""
        from word_to_csv_converter import WordToCsvConverter

        converter = WordToCsvConverter()
        questions = converter.parse_questions_from_word_doc(filepath)

        # 检查是否有警告信息
        if hasattr(converter, 'warnings') and converter.warnings:
            warning_msg = "发现以下缺失部分：\n\n" + "\n".join(converter.warnings) + "\n\n如确认则显示空白。"
            result = messagebox.askokcancel("警告", warning_msg)
            if not result:
                return  # 用户取消导入

        for question in questions:
            qtype = question.get('type')
            if qtype == 'single_choice':
                # 确保选项格式正确
                if 'options' not in question:
                    question['options'] = {"A": "", "A_en": "", "B": "", "B_en": "", "C": "", "C_en": "", "D": "", "D_en": ""}
                self.question_database['single_choice_questions'].append(question)
            elif qtype == 'true_false':
                self.question_database['true_false_questions'].append(question)
            elif qtype == 'essay':
                self.question_database['essay_questions'].append(question)
            elif qtype == 'calculation':
                self.question_database['calculation_questions'].append(question)
    
    def import_from_json(self, filepath):
        """从JSON文件导入试题"""
        with open(filepath, 'r', encoding='utf-8') as file:
            data = json.load(file)
            
        # 假设JSON格式与我们定义的格式一致
        for qtype, qlist in data.items():
            if qtype in self.question_database:
                self.question_database[qtype].extend(qlist)
    
    def export_questions(self):
        """导出题库"""
        filepath = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialfile="question_database.json"
        )
        
        if filepath:
            try:
                with open(filepath, 'w', encoding='utf-8') as file:
                    json.dump(self.question_database, file, ensure_ascii=False, indent=2)
                messagebox.showinfo("成功", "题库导出完成！")
            except Exception as e:
                messagebox.showerror("错误", f"导出题库失败: {str(e)}")
    
    def add_question_dialog(self):
        """添加题目的对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("添加题目")
        dialog.geometry("800x600")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 题型选择
        ttk.Label(dialog, text="题型:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        question_type = tk.StringVar(value="single_choice")
        type_combo = ttk.Combobox(dialog, textvariable=question_type, 
                                  values=["single_choice", "true_false", "essay", "calculation"])
        type_combo.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        # 题目内容框架
        content_frame = ttk.LabelFrame(dialog, text="题目内容")
        content_frame.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="nsew")
        
        # 根据题型显示不同的输入字段
        self.create_question_input_fields(content_frame, question_type, dialog)
        
        # 按钮框架
        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="添加", 
                  command=lambda: self.save_new_question(question_type.get(), dialog)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # 配置网格权重
        dialog.columnconfigure(1, weight=1)
        dialog.rowconfigure(1, weight=1)
        content_frame.columnconfigure(1, weight=1)
    
    def create_question_input_fields(self, parent, question_type_var, dialog):
        """创建题目输入字段"""
        # 清除现有内容
        for widget in parent.winfo_children():
            widget.destroy()
        
        def update_fields(*args):
            # 清除现有内容
            for widget in parent.winfo_children():
                widget.destroy()
            
            qtype = question_type_var.get()
            
            if qtype == "single_choice":
                # 单选题字段
                row = 0
                ttk.Label(parent, text="ID:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_id_var = tk.StringVar(value=f"sc_{len(self.question_database['single_choice_questions']) + 1}")
                ttk.Entry(parent, textvariable=self.new_id_var, width=30).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="难度:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_difficulty_var = tk.StringVar(value="medium")
                ttk.Combobox(parent, textvariable=self.new_difficulty_var, 
                            values=["easy", "medium", "hard"]).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文题干:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_chinese_stem_var = tk.StringVar()
                chinese_stem_entry = ttk.Entry(parent, textvariable=self.new_chinese_stem_var, width=60)
                chinese_stem_entry.grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文题干:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_english_stem_var = tk.StringVar()
                english_stem_entry = ttk.Entry(parent, textvariable=self.new_english_stem_var, width=60)
                english_stem_entry.grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项A (中):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_a_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_a_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项A (英):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_a_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_a_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项B (中):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_b_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_b_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项B (英):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_b_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_b_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项C (中):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_c_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_c_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项C (英):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_c_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_c_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项D (中):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_d_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_d_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="选项D (英):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_option_d_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_option_d_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="正确答案:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_correct_answer_var = tk.StringVar()
                ttk.Combobox(parent, textvariable=self.new_correct_answer_var, 
                            values=["A", "B", "C", "D"]).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="知识点:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_knowledge_point_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_knowledge_point_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="标签 (逗号分隔):").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_tags_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_tags_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
            elif qtype == "true_false":
                # 是非题字段
                row = 0
                ttk.Label(parent, text="ID:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_id_var = tk.StringVar(value=f"tf_{len(self.question_database['true_false_questions']) + 1}")
                ttk.Entry(parent, textvariable=self.new_id_var, width=30).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文题干:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_chinese_stem_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_chinese_stem_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文题干:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_english_stem_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_english_stem_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="正确答案:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_correct_answer_var = tk.StringVar()
                ttk.Combobox(parent, textvariable=self.new_correct_answer_var, 
                            values=["T", "F"]).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文解释:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_explanation_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_explanation_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文解释:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_explanation_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_explanation_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="知识点:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_knowledge_point_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_knowledge_point_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="标签 (逗号分隔):").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_tags_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_tags_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
            elif qtype == "essay":
                # 论述题字段
                row = 0
                ttk.Label(parent, text="ID:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_id_var = tk.StringVar(value=f"es_{len(self.question_database['essay_questions']) + 1}")
                ttk.Entry(parent, textvariable=self.new_id_var, width=30).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文题干:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_chinese_stem_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_chinese_stem_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文题干:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_english_stem_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_english_stem_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文评分标准:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_scoring_guide_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_scoring_guide_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文评分标准:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_scoring_guide_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_scoring_guide_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="知识点:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_knowledge_point_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_knowledge_point_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="标签 (逗号分隔):").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_tags_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_tags_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
            elif qtype == "calculation":
                # 计算题字段
                row = 0
                ttk.Label(parent, text="ID:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_id_var = tk.StringVar(value=f"calc_{len(self.question_database['calculation_questions']) + 1}")
                ttk.Entry(parent, textvariable=self.new_id_var, width=30).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文背景:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_context_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_context_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文背景:").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_context_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_context_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文选项 (用#分隔):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_alternatives_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_alternatives_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文选项 (用#分隔):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_alternatives_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_alternatives_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="参数 (格式: key:value;key:value):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_parameters_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_parameters_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="中文要求 (用#分隔):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_requirements_ch_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_requirements_ch_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="英文要求 (用#分隔):").grid(row=row, column=0, sticky=tk.NW, padx=5, pady=2)
                self.new_requirements_en_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_requirements_en_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="知识点:").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_knowledge_point_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_knowledge_point_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
                
                row += 1
                ttk.Label(parent, text="标签 (逗号分隔):").grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
                self.new_tags_var = tk.StringVar()
                ttk.Entry(parent, textvariable=self.new_tags_var, width=60).grid(row=row, column=1, sticky="ew", padx=5, pady=2)
        
        # 绑定事件以更新字段
        question_type_var.trace('w', update_fields)
        # 初始调用以显示默认字段
        update_fields()
    
    def save_new_question(self, qtype, dialog):
        """保存新题目"""
        try:
            if qtype == "single_choice":
                question = {
                    "id": self.new_id_var.get(),
                    "type": "single_choice",
                    "difficulty": self.new_difficulty_var.get(),
                    "chinese_stem": self.new_chinese_stem_var.get(),
                    "english_stem": self.new_english_stem_var.get(),
                    "options": {
                        "A": self.new_option_a_ch_var.get(),
                        "A_en": self.new_option_a_en_var.get(),
                        "B": self.new_option_b_ch_var.get(),
                        "B_en": self.new_option_b_en_var.get(),
                        "C": self.new_option_c_ch_var.get(),
                        "C_en": self.new_option_c_en_var.get(),
                        "D": self.new_option_d_ch_var.get(),
                        "D_en": self.new_option_d_en_var.get()
                    },
                    "correct_answer": self.new_correct_answer_var.get(),
                    "knowledge_point": self.new_knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in self.new_tags_var.get().split(',') if tag.strip()]
                }
                self.question_database["single_choice_questions"].append(question)
                
            elif qtype == "true_false":
                question = {
                    "id": self.new_id_var.get(),
                    "type": "true_false",
                    "chinese_stem": self.new_chinese_stem_var.get(),
                    "english_stem": self.new_english_stem_var.get(),
                    "correct_answer": self.new_correct_answer_var.get(),
                    "explanation": self.new_explanation_ch_var.get(),
                    "explanation_en": self.new_explanation_en_var.get(),
                    "knowledge_point": self.new_knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in self.new_tags_var.get().split(',') if tag.strip()]
                }
                self.question_database["true_false_questions"].append(question)
                
            elif qtype == "essay":
                question = {
                    "id": self.new_id_var.get(),
                    "type": "essay",
                    "chinese_stem": self.new_chinese_stem_var.get(),
                    "english_stem": self.new_english_stem_var.get(),
                    "scoring_guide": self.new_scoring_guide_ch_var.get().split(';'),
                    "scoring_guide_en": self.new_scoring_guide_en_var.get().split(';'),
                    "knowledge_point": self.new_knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in self.new_tags_var.get().split(',') if tag.strip()]
                }
                self.question_database["essay_questions"].append(question)
                
            elif qtype == "calculation":
                # 解析alternatives和requirements
                alternatives_ch = self.new_alternatives_ch_var.get().split('#')
                alternatives_en = self.new_alternatives_en_var.get().split('#')
                requirements_ch = self.new_requirements_ch_var.get().split('#')
                requirements_en = self.new_requirements_en_var.get().split('#')
                
                # 解析parameters
                params_str = self.new_parameters_var.get()
                parameters = {}
                if params_str:
                    for param in params_str.split(';'):
                        if ':' in param:
                            key, value = param.split(':', 1)
                            parameters[key.strip()] = value.strip()
                
                question = {
                    "id": self.new_id_var.get(),
                    "type": "calculation",
                    "context": {
                        "chinese": self.new_context_ch_var.get(),
                        "english": self.new_context_en_var.get()
                    },
                    "alternatives": [{"description": {"chinese": ch, "english": en}} 
                                     for ch, en in zip(alternatives_ch, alternatives_en)],
                    "parameters": parameters,
                    "requirements": {
                        "chinese": requirements_ch,
                        "english": requirements_en
                    },
                    "knowledge_point": self.new_knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in self.new_tags_var.get().split(',') if tag.strip()]
                }
                self.question_database["calculation_questions"].append(question)
            
            messagebox.showinfo("成功", "题目添加成功！")
            dialog.destroy()
            self.refresh_question_list()
            
        except Exception as e:
            messagebox.showerror("错误", f"添加题目失败: {str(e)}")
    
    def edit_question_dialog(self):
        """编辑选中的题目"""
        selection = self.question_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一道题目")
            return
            
        item = self.question_tree.item(selection[0])
        question_id = item['values'][0]  # ID是第一列
        question_type = item['values'][1]  # 题型是第二列
        
        # 找到对应的题目
        question = None
        question_list = None
        
        if question_type == "单选题":
            for q in self.question_database["single_choice_questions"]:
                if q["id"] == question_id:
                    question = q
                    question_list = self.question_database["single_choice_questions"]
                    break
        elif question_type == "是非题":
            for q in self.question_database["true_false_questions"]:
                if q["id"] == question_id:
                    question = q
                    question_list = self.question_database["true_false_questions"]
                    break
        elif question_type == "论述题":
            for q in self.question_database["essay_questions"]:
                if q["id"] == question_id:
                    question = q
                    question_list = self.question_database["essay_questions"]
                    break
        elif question_type == "计算题":
            for q in self.question_database["calculation_questions"]:
                if q["id"] == question_id:
                    question = q
                    question_list = self.question_database["calculation_questions"]
                    break
        
        if not question:
            messagebox.showerror("错误", "未找到对应题目")
            return
        
        # 创建编辑对话框
        dialog = tk.Toplevel(self.root)
        dialog.title(f"编辑题目 - {question_id}")
        dialog.geometry("800x600")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 题目内容框架
        content_frame = ttk.LabelFrame(dialog, text="题目内容")
        content_frame.grid(row=0, column=0, columnspan=2, padx=5, pady=5, sticky="nsew")
        
        # 根据题型填充字段
        if question["type"] == "single_choice":
            # ID (不可编辑)
            ttk.Label(content_frame, text="ID:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
            id_label = ttk.Label(content_frame, text=question["id"])
            id_label.grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
            
            # 难度
            ttk.Label(content_frame, text="难度:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
            difficulty_var = tk.StringVar(value=question.get("difficulty", "medium"))
            ttk.Combobox(content_frame, textvariable=difficulty_var, 
                        values=["easy", "medium", "hard"]).grid(row=1, column=1, sticky="ew", padx=5, pady=2)
            
            # 中文题干
            ttk.Label(content_frame, text="中文题干:").grid(row=2, column=0, sticky=tk.NW, padx=5, pady=2)
            chinese_stem_var = tk.StringVar(value=question.get("chinese_stem", ""))
            ttk.Entry(content_frame, textvariable=chinese_stem_var, width=60).grid(row=2, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文题干
            ttk.Label(content_frame, text="英文题干:").grid(row=3, column=0, sticky=tk.NW, padx=5, pady=2)
            english_stem_var = tk.StringVar(value=question.get("english_stem", ""))
            ttk.Entry(content_frame, textvariable=english_stem_var, width=60).grid(row=3, column=1, sticky="ew", padx=5, pady=2)
            
            # 选项A
            ttk.Label(content_frame, text="选项A (中):").grid(row=4, column=0, sticky=tk.NW, padx=5, pady=2)
            option_a_ch_var = tk.StringVar(value=question["options"].get("A", ""))
            ttk.Entry(content_frame, textvariable=option_a_ch_var, width=60).grid(row=4, column=1, sticky="ew", padx=5, pady=2)
            
            ttk.Label(content_frame, text="选项A (英):").grid(row=5, column=0, sticky=tk.NW, padx=5, pady=2)
            option_a_en_var = tk.StringVar(value=question["options"].get("A_en", ""))
            ttk.Entry(content_frame, textvariable=option_a_en_var, width=60).grid(row=5, column=1, sticky="ew", padx=5, pady=2)
            
            # 选项B
            ttk.Label(content_frame, text="选项B (中):").grid(row=6, column=0, sticky=tk.NW, padx=5, pady=2)
            option_b_ch_var = tk.StringVar(value=question["options"].get("B", ""))
            ttk.Entry(content_frame, textvariable=option_b_ch_var, width=60).grid(row=6, column=1, sticky="ew", padx=5, pady=2)
            
            ttk.Label(content_frame, text="选项B (英):").grid(row=7, column=0, sticky=tk.NW, padx=5, pady=2)
            option_b_en_var = tk.StringVar(value=question["options"].get("B_en", ""))
            ttk.Entry(content_frame, textvariable=option_b_en_var, width=60).grid(row=7, column=1, sticky="ew", padx=5, pady=2)
            
            # 选项C
            ttk.Label(content_frame, text="选项C (中):").grid(row=8, column=0, sticky=tk.NW, padx=5, pady=2)
            option_c_ch_var = tk.StringVar(value=question["options"].get("C", ""))
            ttk.Entry(content_frame, textvariable=option_c_ch_var, width=60).grid(row=8, column=1, sticky="ew", padx=5, pady=2)
            
            ttk.Label(content_frame, text="选项C (英):").grid(row=9, column=0, sticky=tk.NW, padx=5, pady=2)
            option_c_en_var = tk.StringVar(value=question["options"].get("C_en", ""))
            ttk.Entry(content_frame, textvariable=option_c_en_var, width=60).grid(row=9, column=1, sticky="ew", padx=5, pady=2)
            
            # 选项D
            ttk.Label(content_frame, text="选项D (中):").grid(row=10, column=0, sticky=tk.NW, padx=5, pady=2)
            option_d_ch_var = tk.StringVar(value=question["options"].get("D", ""))
            ttk.Entry(content_frame, textvariable=option_d_ch_var, width=60).grid(row=10, column=1, sticky="ew", padx=5, pady=2)
            
            ttk.Label(content_frame, text="选项D (英):").grid(row=11, column=0, sticky=tk.NW, padx=5, pady=2)
            option_d_en_var = tk.StringVar(value=question["options"].get("D_en", ""))
            ttk.Entry(content_frame, textvariable=option_d_en_var, width=60).grid(row=11, column=1, sticky="ew", padx=5, pady=2)
            
            # 正确答案
            ttk.Label(content_frame, text="正确答案:").grid(row=12, column=0, sticky=tk.W, padx=5, pady=2)
            correct_answer_var = tk.StringVar(value=question.get("correct_answer", ""))
            ttk.Combobox(content_frame, textvariable=correct_answer_var, 
                        values=["A", "B", "C", "D"]).grid(row=12, column=1, sticky="ew", padx=5, pady=2)
            
            # 知识点
            ttk.Label(content_frame, text="知识点:").grid(row=13, column=0, sticky=tk.W, padx=5, pady=2)
            knowledge_point_var = tk.StringVar(value=question.get("knowledge_point", ""))
            ttk.Entry(content_frame, textvariable=knowledge_point_var, width=60).grid(row=13, column=1, sticky="ew", padx=5, pady=2)
            
            # 标签
            ttk.Label(content_frame, text="标签 (逗号分隔):").grid(row=14, column=0, sticky=tk.W, padx=5, pady=2)
            tags_var = tk.StringVar(value=",".join(question.get("tags", [])))
            ttk.Entry(content_frame, textvariable=tags_var, width=60).grid(row=14, column=1, sticky="ew", padx=5, pady=2)
            
            # 更新函数
            def update_question():
                idx = question_list.index(question)
                question_list[idx] = {
                    "id": question["id"],  # ID不变
                    "type": "single_choice",
                    "difficulty": difficulty_var.get(),
                    "chinese_stem": chinese_stem_var.get(),
                    "english_stem": english_stem_var.get(),
                    "options": {
                        "A": option_a_ch_var.get(),
                        "A_en": option_a_en_var.get(),
                        "B": option_b_ch_var.get(),
                        "B_en": option_b_en_var.get(),
                        "C": option_c_ch_var.get(),
                        "C_en": option_c_en_var.get(),
                        "D": option_d_ch_var.get(),
                        "D_en": option_d_en_var.get()
                    },
                    "correct_answer": correct_answer_var.get(),
                    "knowledge_point": knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in tags_var.get().split(',') if tag.strip()]
                }
                messagebox.showinfo("成功", "题目更新成功！")
                dialog.destroy()
                self.refresh_question_list()
                
        elif question["type"] == "true_false":
            # ID (不可编辑)
            ttk.Label(content_frame, text="ID:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
            id_label = ttk.Label(content_frame, text=question["id"])
            id_label.grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
            
            # 中文题干
            ttk.Label(content_frame, text="中文题干:").grid(row=1, column=0, sticky=tk.NW, padx=5, pady=2)
            chinese_stem_var = tk.StringVar(value=question.get("chinese_stem", ""))
            ttk.Entry(content_frame, textvariable=chinese_stem_var, width=60).grid(row=1, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文题干
            ttk.Label(content_frame, text="英文题干:").grid(row=2, column=0, sticky=tk.NW, padx=5, pady=2)
            english_stem_var = tk.StringVar(value=question.get("english_stem", ""))
            ttk.Entry(content_frame, textvariable=english_stem_var, width=60).grid(row=2, column=1, sticky="ew", padx=5, pady=2)
            
            # 正确答案
            ttk.Label(content_frame, text="正确答案:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=2)
            correct_answer_var = tk.StringVar(value=question.get("correct_answer", ""))
            ttk.Combobox(content_frame, textvariable=correct_answer_var, 
                        values=["T", "F"]).grid(row=3, column=1, sticky="ew", padx=5, pady=2)
            
            # 中文解释
            ttk.Label(content_frame, text="中文解释:").grid(row=4, column=0, sticky=tk.NW, padx=5, pady=2)
            explanation_ch_var = tk.StringVar(value=question.get("explanation", ""))
            ttk.Entry(content_frame, textvariable=explanation_ch_var, width=60).grid(row=4, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文解释
            ttk.Label(content_frame, text="英文解释:").grid(row=5, column=0, sticky=tk.NW, padx=5, pady=2)
            explanation_en_var = tk.StringVar(value=question.get("explanation_en", ""))
            ttk.Entry(content_frame, textvariable=explanation_en_var, width=60).grid(row=5, column=1, sticky="ew", padx=5, pady=2)
            
            # 知识点
            ttk.Label(content_frame, text="知识点:").grid(row=6, column=0, sticky=tk.W, padx=5, pady=2)
            knowledge_point_var = tk.StringVar(value=question.get("knowledge_point", ""))
            ttk.Entry(content_frame, textvariable=knowledge_point_var, width=60).grid(row=6, column=1, sticky="ew", padx=5, pady=2)
            
            # 标签
            ttk.Label(content_frame, text="标签 (逗号分隔):").grid(row=7, column=0, sticky=tk.W, padx=5, pady=2)
            tags_var = tk.StringVar(value=",".join(question.get("tags", [])))
            ttk.Entry(content_frame, textvariable=tags_var, width=60).grid(row=7, column=1, sticky="ew", padx=5, pady=2)
            
            # 更新函数
            def update_question():
                idx = question_list.index(question)
                question_list[idx] = {
                    "id": question["id"],  # ID不变
                    "type": "true_false",
                    "chinese_stem": chinese_stem_var.get(),
                    "english_stem": english_stem_var.get(),
                    "correct_answer": correct_answer_var.get(),
                    "explanation": explanation_ch_var.get(),
                    "explanation_en": explanation_en_var.get(),
                    "knowledge_point": knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in tags_var.get().split(',') if tag.strip()]
                }
                messagebox.showinfo("成功", "题目更新成功！")
                dialog.destroy()
                self.refresh_question_list()
                
        elif question["type"] == "essay":
            # ID (不可编辑)
            ttk.Label(content_frame, text="ID:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
            id_label = ttk.Label(content_frame, text=question["id"])
            id_label.grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
            
            # 中文题干
            ttk.Label(content_frame, text="中文题干:").grid(row=1, column=0, sticky=tk.NW, padx=5, pady=2)
            chinese_stem_var = tk.StringVar(value=question.get("chinese_stem", ""))
            ttk.Entry(content_frame, textvariable=chinese_stem_var, width=60).grid(row=1, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文题干
            ttk.Label(content_frame, text="英文题干:").grid(row=2, column=0, sticky=tk.NW, padx=5, pady=2)
            english_stem_var = tk.StringVar(value=question.get("english_stem", ""))
            ttk.Entry(content_frame, textvariable=english_stem_var, width=60).grid(row=2, column=1, sticky="ew", padx=5, pady=2)
            
            # 中文评分标准
            ttk.Label(content_frame, text="中文评分标准:").grid(row=3, column=0, sticky=tk.NW, padx=5, pady=2)
            scoring_guide_ch_var = tk.StringVar(value=";".join(question.get("scoring_guide", [])))
            ttk.Entry(content_frame, textvariable=scoring_guide_ch_var, width=60).grid(row=3, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文评分标准
            ttk.Label(content_frame, text="英文评分标准:").grid(row=4, column=0, sticky=tk.NW, padx=5, pady=2)
            scoring_guide_en_var = tk.StringVar(value=";".join(question.get("scoring_guide_en", [])))
            ttk.Entry(content_frame, textvariable=scoring_guide_en_var, width=60).grid(row=4, column=1, sticky="ew", padx=5, pady=2)
            
            # 知识点
            ttk.Label(content_frame, text="知识点:").grid(row=5, column=0, sticky=tk.W, padx=5, pady=2)
            knowledge_point_var = tk.StringVar(value=question.get("knowledge_point", ""))
            ttk.Entry(content_frame, textvariable=knowledge_point_var, width=60).grid(row=5, column=1, sticky="ew", padx=5, pady=2)
            
            # 标签
            ttk.Label(content_frame, text="标签 (逗号分隔):").grid(row=6, column=0, sticky=tk.W, padx=5, pady=2)
            tags_var = tk.StringVar(value=",".join(question.get("tags", [])))
            ttk.Entry(content_frame, textvariable=tags_var, width=60).grid(row=6, column=1, sticky="ew", padx=5, pady=2)
            
            # 更新函数
            def update_question():
                idx = question_list.index(question)
                question_list[idx] = {
                    "id": question["id"],  # ID不变
                    "type": "essay",
                    "chinese_stem": chinese_stem_var.get(),
                    "english_stem": english_stem_var.get(),
                    "scoring_guide": scoring_guide_ch_var.get().split(';'),
                    "scoring_guide_en": scoring_guide_en_var.get().split(';'),
                    "knowledge_point": knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in tags_var.get().split(',') if tag.strip()]
                }
                messagebox.showinfo("成功", "题目更新成功！")
                dialog.destroy()
                self.refresh_question_list()
                
        elif question["type"] == "calculation":
            # ID (不可编辑)
            ttk.Label(content_frame, text="ID:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
            id_label = ttk.Label(content_frame, text=question["id"])
            id_label.grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
            
            # 中文背景
            ttk.Label(content_frame, text="中文背景:").grid(row=1, column=0, sticky=tk.NW, padx=5, pady=2)
            context_ch_var = tk.StringVar(value=question["context"].get("chinese", ""))
            ttk.Entry(content_frame, textvariable=context_ch_var, width=60).grid(row=1, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文背景
            ttk.Label(content_frame, text="英文背景:").grid(row=2, column=0, sticky=tk.NW, padx=5, pady=2)
            context_en_var = tk.StringVar(value=question["context"].get("english", ""))
            ttk.Entry(content_frame, textvariable=context_en_var, width=60).grid(row=2, column=1, sticky="ew", padx=5, pady=2)
            
            # 中文选项
            alternatives_ch = [alt["description"].get("chinese", "") for alt in question.get("alternatives", [])]
            ttk.Label(content_frame, text="中文选项 (用#分隔):").grid(row=3, column=0, sticky=tk.NW, padx=5, pady=2)
            alternatives_ch_var = tk.StringVar(value="#".join(alternatives_ch))
            ttk.Entry(content_frame, textvariable=alternatives_ch_var, width=60).grid(row=3, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文选项
            alternatives_en = [alt["description"].get("english", "") for alt in question.get("alternatives", [])]
            ttk.Label(content_frame, text="英文选项 (用#分隔):").grid(row=4, column=0, sticky=tk.NW, padx=5, pady=2)
            alternatives_en_var = tk.StringVar(value="#".join(alternatives_en))
            ttk.Entry(content_frame, textvariable=alternatives_en_var, width=60).grid(row=4, column=1, sticky="ew", padx=5, pady=2)
            
            # 参数
            params_parts = []
            for k, v in question.get("parameters", {}).items():
                params_parts.append(f"{k}:{v}")
            ttk.Label(content_frame, text="参数 (格式: key:value;key:value):").grid(row=5, column=0, sticky=tk.NW, padx=5, pady=2)
            parameters_var = tk.StringVar(value=";".join(params_parts))
            ttk.Entry(content_frame, textvariable=parameters_var, width=60).grid(row=5, column=1, sticky="ew", padx=5, pady=2)
            
            # 中文要求
            requirements_ch = question["requirements"].get("chinese", [])
            ttk.Label(content_frame, text="中文要求 (用#分隔):").grid(row=6, column=0, sticky=tk.NW, padx=5, pady=2)
            requirements_ch_var = tk.StringVar(value="#".join(requirements_ch))
            ttk.Entry(content_frame, textvariable=requirements_ch_var, width=60).grid(row=6, column=1, sticky="ew", padx=5, pady=2)
            
            # 英文要求
            requirements_en = question["requirements"].get("english", [])
            ttk.Label(content_frame, text="英文要求 (用#分隔):").grid(row=7, column=0, sticky=tk.NW, padx=5, pady=2)
            requirements_en_var = tk.StringVar(value="#".join(requirements_en))
            ttk.Entry(content_frame, textvariable=requirements_en_var, width=60).grid(row=7, column=1, sticky="ew", padx=5, pady=2)
            
            # 知识点
            ttk.Label(content_frame, text="知识点:").grid(row=8, column=0, sticky=tk.W, padx=5, pady=2)
            knowledge_point_var = tk.StringVar(value=question.get("knowledge_point", ""))
            ttk.Entry(content_frame, textvariable=knowledge_point_var, width=60).grid(row=8, column=1, sticky="ew", padx=5, pady=2)
            
            # 标签
            ttk.Label(content_frame, text="标签 (逗号分隔):").grid(row=9, column=0, sticky=tk.W, padx=5, pady=2)
            tags_var = tk.StringVar(value=",".join(question.get("tags", [])))
            ttk.Entry(content_frame, textvariable=tags_var, width=60).grid(row=9, column=1, sticky="ew", padx=5, pady=2)
            
            # 更新函数
            def update_question():
                # 解析alternatives和requirements
                alt_ch_list = alternatives_ch_var.get().split('#')
                alt_en_list = alternatives_en_var.get().split('#')
                req_ch_list = requirements_ch_var.get().split('#')
                req_en_list = requirements_en_var.get().split('#')
                
                # 解析parameters
                params_str = parameters_var.get()
                parameters = {}
                if params_str:
                    for param in params_str.split(';'):
                        if ':' in param:
                            key, value = param.split(':', 1)
                            parameters[key.strip()] = value.strip()
                
                idx = question_list.index(question)
                question_list[idx] = {
                    "id": question["id"],  # ID不变
                    "type": "calculation",
                    "context": {
                        "chinese": context_ch_var.get(),
                        "english": context_en_var.get()
                    },
                    "alternatives": [{"description": {"chinese": ch, "english": en}} 
                                     for ch, en in zip(alt_ch_list, alt_en_list)],
                    "parameters": parameters,
                    "requirements": {
                        "chinese": req_ch_list,
                        "english": req_en_list
                    },
                    "knowledge_point": knowledge_point_var.get(),
                    "tags": [tag.strip() for tag in tags_var.get().split(',') if tag.strip()]
                }
                messagebox.showinfo("成功", "题目更新成功！")
                dialog.destroy()
                self.refresh_question_list()
        
        # 按钮框架
        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=1, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="更新", command=update_question).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        
        # 配置网格权重
        dialog.columnconfigure(1, weight=1)
        dialog.rowconfigure(0, weight=1)
        content_frame.columnconfigure(1, weight=1)
    
    def delete_question(self):
        """删除选中的题目"""
        selection = self.question_tree.selection()
        if not selection:
            messagebox.showwarning("警告", "请先选择一道题目")
            return
            
        item = self.question_tree.item(selection[0])
        question_id = item['values'][0]  # ID是第一列
        question_type = item['values'][1]  # 题型是第二列
        
        result = messagebox.askyesno("确认删除", f"确定要删除题目 {question_id} 吗？")
        if result:
            # 根据题型在对应的列表中删除题目
            if question_type == "单选题":
                self.question_database["single_choice_questions"] = [
                    q for q in self.question_database["single_choice_questions"] 
                    if q["id"] != question_id
                ]
            elif question_type == "是非题":
                self.question_database["true_false_questions"] = [
                    q for q in self.question_database["true_false_questions"] 
                    if q["id"] != question_id
                ]
            elif question_type == "论述题":
                self.question_database["essay_questions"] = [
                    q for q in self.question_database["essay_questions"] 
                    if q["id"] != question_id
                ]
            elif question_type == "计算题":
                self.question_database["calculation_questions"] = [
                    q for q in self.question_database["calculation_questions"] 
                    if q["id"] != question_id
                ]
            
            messagebox.showinfo("成功", "题目删除成功！")
            self.refresh_question_list()
    
    def refresh_question_list(self):
        """刷新题目列表"""
        # 清空现有项目
        for item in self.question_tree.get_children():
            self.question_tree.delete(item)
        
        # 添加所有题目到列表
        for qtype, qlist in self.question_database.items():
            for q in qlist:
                # 获取使用次数和最后使用时间
                usage_info = self.question_usage.get(q["id"], {"count": 0, "last_used": "从未使用"})
                
                # 确定题型显示名称
                type_name = {
                    "single_choice_questions": "单选题",
                    "true_false_questions": "是非题", 
                    "essay_questions": "论述题",
                    "calculation_questions": "计算题"
                }.get(qtype, "未知")
                
                # 添加到树形视图
                self.question_tree.insert("", tk.END, values=(
                    q["id"],
                    type_name,
                    q.get("difficulty", "-"),
                    q.get("knowledge_point", "-"),
                    usage_info["count"],
                    usage_info["last_used"]
                ))
    
    def preview_exam(self):
        """预览试卷"""
        # 构建试卷配置
        exam_config = {
            'school': self.school_var.get(),
            'academic_year': self.academic_year_var.get(),
            'course_name': self.course_name_var.get(),
            'paper_type': self.paper_type_var.get(),
            'course_code': 'B0600432',
            'exam_type': '闭卷、笔试',
            'question_types': {
                'single_choice': {
                    'count': int(self.sc_count_var.get()),
                    'score_per_question': int(self.sc_score_var.get()),
                    'difficulty': 'medium'
                },
                'true_false': {
                    'count': int(self.tf_count_var.get()),
                    'score_per_question': int(self.tf_score_var.get())
                },
                'essay': {
                    'count': int(self.es_count_var.get()),
                    'score_per_question': int(self.es_score_var.get())
                },
                'calculation': {
                    'count': int(self.calc_count_var.get()),
                    'score_per_question': int(self.calc_score_var.get())
                }
            }
        }
        
        # 生成试卷内容
        generator = ForestEconomicsExamGenerator(self.question_database)
        exam_content = generator.generate_exam(exam_config)
        
        # 显示预览
        self.preview_text.delete(1.0, tk.END)
        self.preview_text.insert(tk.END, exam_content['header'] + "\n\n")
        
        # 显示分数表
        score_table = exam_content['score_table']
        self.preview_text.insert(tk.END, "题号\t一\t二\t三\t四\t五\t六\t七\t总分\t总分人\n")
        self.preview_text.insert(tk.END, "分值\t\t\t\t\t\t\t\t\t\n")
        self.preview_text.insert(tk.END, "得分\t\t\t\t\t\t\t\t\t\n")
        self.preview_text.insert(tk.END, "-" * 80 + "\n\n")
        
        # 显示题目
        for q in exam_content['questions']:
            self.preview_text.insert(tk.END, f"{q['content']}\n\n")
    
    def generate_exam(self):
        """生成试卷"""
        # 构建试卷配置
        exam_config = {
            'school': self.school_var.get(),
            'academic_year': self.academic_year_var.get(),
            'course_name': self.course_name_var.get(),
            'paper_type': self.paper_type_var.get(),
            'course_code': 'B0600432',
            'exam_type': '闭卷、笔试',
            'question_types': {
                'single_choice': {
                    'count': int(self.sc_count_var.get()),
                    'score_per_question': int(self.sc_score_var.get()),
                    'difficulty': 'medium'
                },
                'true_false': {
                    'count': int(self.tf_count_var.get()),
                    'score_per_question': int(self.tf_score_var.get())
                },
                'essay': {
                    'count': int(self.es_count_var.get()),
                    'score_per_question': int(self.es_score_var.get())
                },
                'calculation': {
                    'count': int(self.calc_count_var.get()),
                    'score_per_question': int(self.calc_score_var.get())
                }
            }
        }
        
        # 生成试卷
        generator = ForestEconomicsExamGenerator(self.question_database)
        exam_content = generator.generate_exam(exam_config)
        
        # 更新题目使用历史
        current_date = datetime.now().strftime("%Y-%m-%d")
        for q in exam_content['questions']:
            if 'correct_answer' in q:  # 单选题和判断题
                qid = q.get('question_id', '')
                if qid:
                    if qid not in self.question_usage:
                        self.question_usage[qid] = {"count": 0, "last_used": current_date}
                    self.question_usage[qid]["count"] += 1
                    self.question_usage[qid]["last_used"] = current_date
            elif 'scoring_guide' in q:  # 论述题
                qid = q.get('question_id', '')
                if qid:
                    if qid not in self.question_usage:
                        self.question_usage[qid] = {"count": 0, "last_used": current_date}
                    self.question_usage[qid]["count"] += 1
                    self.question_usage[qid]["last_used"] = current_date
            elif 'requirements' in q:  # 计算题
                qid = q.get('question_id', '')
                if qid:
                    if qid not in self.question_usage:
                        self.question_usage[qid] = {"count": 0, "last_used": current_date}
                    self.question_usage[qid]["count"] += 1
                    self.question_usage[qid]["last_used"] = current_date
        
        # 保存为Word文档
        self.save_exam_as_word(exam_content)
        
        # 刷新题目列表以显示更新的使用统计
        self.refresh_question_list()
        
        messagebox.showinfo("成功", "试卷生成完成！")
    
    def save_exam_as_word(self, exam_content):
        """保存试卷为Word格式"""
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            initialfile=f"林业经济学试卷_{self.paper_type_var.get()}.docx"
        )
        
        if filepath:
            doc = Document()
            
            # 设置中文字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rProp.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加标题
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.add_run(exam_content['header']).bold = True
            
            # 添加分值表
            table = doc.add_table(rows=4, cols=10)
            table.style = 'Table Grid'
            
            # 填充表头
            hdr_cells = table.rows[0].cells
            for i, text in enumerate(["题号", "一", "二", "三", "四", "五", "六", "七", "总分", "总分人"]):
                hdr_cells[i].text = text
            
            # 添加题目内容
            for q in exam_content['questions']:
                doc.add_paragraph(f"\n{q['content']}")
            
            # 保存文档
            doc.save(filepath)
            messagebox.showinfo("成功", f"试卷已保存到: {filepath}")

class ForestEconomicsExamGenerator:
    def __init__(self, question_database):
        self.question_database = question_database
    
    def generate_exam(self, exam_config):
        """
        生成试卷
        :param exam_config: 试卷配置参数
        :return: 生成的试卷内容
        """
        exam_content = {
            'header': self.generate_header(exam_config),
            'score_table': self.generate_score_table(exam_config),
            'questions': []
        }
        
        # 生成各类型题目
        if 'single_choice' in exam_config['question_types']:
            single_choice = self.generate_single_choice_questions(
                exam_config['question_types']['single_choice']
            )
            exam_content['questions'].extend(single_choice)
        
        if 'true_false' in exam_config['question_types']:
            true_false = self.generate_true_false_questions(
                exam_config['question_types']['true_false']
            )
            exam_content['questions'].extend(true_false)
        
        if 'essay' in exam_config['question_types']:
            essay = self.generate_essay_questions(
                exam_config['question_types']['essay']
            )
            exam_content['questions'].extend(essay)
        
        if 'calculation' in exam_config['question_types']:
            calculation = self.generate_calculation_questions(
                exam_config['question_types']['calculation']
            )
            exam_content['questions'].extend(calculation)
        
        return exam_content
    
    def generate_header(self, config):
        """生成试卷头部信息"""
        return f"""{config['school']}{config['academic_year']}期末考试试卷
课程名称：《{config['course_name']}（双语）》 （{config['paper_type']}）卷
课程代号：{config['course_code']}
考试形式：{config['exam_type']}
使用对象："""
    
    def generate_score_table(self, config):
        """生成题型分值表"""
        question_types = config['question_types']
        total_score = sum([q['count'] * q['score_per_question'] for q in question_types.values()])
        
        header = "题号"
        scores = "分值"
        points = "得分"
        
        for i, (q_type, q_config) in enumerate(question_types.items()):
            header += f"\t{chr(97+i)}\t"  # a, b, c...
            scores += f"\t{q_config['count'] * q_config['score_per_question']}\t"
            points += f"\t\t"
        
        header += "总分\t总分人"
        scores += f"\t{total_score}\t"
        points += "\t\t"
        
        return {
            'header': header,
            'scores': scores,
            'points': points
        }
    
    def generate_single_choice_questions(self, config):
        """生成单选题"""
        # 从题库中随机选择指定数量的题目
        available_questions = [
            q for q in self.question_database['single_choice_questions']
            if q['difficulty'] == config.get('difficulty', 'medium')
        ]
        
        # 如果数量不够，使用所有可用题目
        if len(available_questions) < config['count']:
            selected_questions = available_questions
        else:
            selected_questions = random.sample(available_questions, config['count'])
        
        questions = []
        for i, q in enumerate(selected_questions, 1):
            question_text = f"""
{i}. {q['chinese_stem']}

A. {q['options']['A']}
B. {q['options']['B']}  
C. {q['options']['C']}
D. {q['options']['D']}

{i}. {q['english_stem']}

A. {q['options']['A_en']}
B. {q['options']['B_en']}
C. {q['options']['C_en']}
D. {q['options']['D_en']}
            """.strip()
            
            questions.append({
                'type': 'single_choice',
                'number': i,
                'content': question_text,
                'correct_answer': q['correct_answer'],
                'question_id': q['id']
            })
        
        return questions
    
    def generate_true_false_questions(self, config):
        """生成是非题"""
        available_questions = self.question_database['true_false_questions']
        
        # 如果数量不够，使用所有可用题目
        if len(available_questions) < config['count']:
            selected_questions = available_questions
        else:
            selected_questions = random.sample(available_questions, config['count'])
        
        questions = []
        for i, q in enumerate(selected_questions, 1):
            question_text = f"""
{q['chinese_stem']}

{q['english_stem']}
            """.strip()
            
            questions.append({
                'type': 'true_false',
                'number': i,
                'content': question_text,
                'correct_answer': q['correct_answer'],
                'question_id': q['id']
            })
        
        return questions
    
    def generate_essay_questions(self, config):
        """生成论述题"""
        available_questions = self.question_database['essay_questions']
        
        # 如果数量不够，使用所有可用题目
        if len(available_questions) < config['count']:
            selected_questions = available_questions
        else:
            selected_questions = random.sample(available_questions, config['count'])
        
        questions = []
        for i, q in enumerate(selected_questions, 1):
            question_text = f"""
{i}. {q['chinese_stem']}

{i}. {q['english_stem']}
            """.strip()
            
            questions.append({
                'type': 'essay',
                'number': i,
                'content': question_text,
                'scoring_guide': q['scoring_guide'],
                'question_id': q['id']
            })
        
        return questions
    
    def generate_calculation_questions(self, config):
        """生成计算题"""
        available_questions = self.question_database['calculation_questions']
        
        # 如果数量不够，使用所有可用题目
        if len(available_questions) < config['count']:
            selected_questions = available_questions
        else:
            selected_questions = random.sample(available_questions, config['count'])
        
        questions = []
        for i, q in enumerate(selected_questions, 1):
            # 构建参数字符串
            params_str = ""
            for key, value in q['parameters'].items():
                params_str += f"{key}:{value}; "
            
            # 构建选项字符串
            alternatives_str = ""
            for j, alt in enumerate(q['alternatives']):
                alternatives_str += f"- {alt['description']['chinese']}\n"
                alternatives_str += f"- {alt['description']['english']}\n"
            
            # 构建要求字符串
            requirements_str = ""
            for req_ch, req_en in zip(q['requirements']['chinese'], q['requirements']['english']):
                requirements_str += f"要求: {req_ch}\n     {req_en}\n"
            
            question_text = f"""
{q['context']['chinese']}

{q['context']['english']}

{alternatives_str}
参数: {params_str}

{requirements_str}
            """.strip()
            
            questions.append({
                'type': 'calculation',
                'number': i,
                'content': question_text,
                'question_id': q['id']
            })
        
        return questions

def main():
    root = tk.Tk()
    app = ExamGeneratorGUI(root)
    root.mainloop()

if __name__ == '__main__':
    main()