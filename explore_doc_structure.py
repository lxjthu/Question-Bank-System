"""
探索Word文档的深层结构
处理复杂布局和特殊内容
"""

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.section import Section
import os

def clean_text(text):
    """
    清理文本中的特殊字符
    """
    if text is None:
        return ""
    # 移除不可见字符和特殊字符
    cleaned = text.replace('\x00', '').replace('\ufffd', '').strip()
    return cleaned

def explore_document_structure(file_path):
    """
    探索文档深层结构
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return
    
    try:
        doc = Document(file_path)
        print(f"探索文档结构: {file_path}")
        print("="*60)
        
        # 检查文档的基本信息
        print("文档基本信息:")
        print("-" * 30)
        print(f"总段落数: {len(doc.paragraphs)}")
        print(f"总表格数: {len(doc.tables)}")
        print(f"总节数: {len(doc.sections)}")
        
        # 检查每个节
        print(f"\n节数: {len(doc.sections)}")
        for i, section in enumerate(doc.sections):
            print(f"  第 {i+1} 节:")
            print(f"    页边距: 上={section.top_margin}, 下={section.bottom_margin}, 左={section.left_margin}, 右={section.right_margin}")
            print(f"    页面大小: {section.page_width} x {section.page_height}")
            print(f"    页眉: {section.header is not None}")
            print(f"    页脚: {section.footer is not None}")
        
        # 检查页眉和页脚
        print("\n页眉页脚内容:")
        print("-" * 30)
        for i, section in enumerate(doc.sections):
            print(f"第 {i+1} 节:")
            
            # 检查页眉
            if section.header:
                header_content = []
                for paragraph in section.header.paragraphs:
                    text = clean_text(paragraph.text)
                    if text:
                        header_content.append(text)
                if header_content:
                    print(f"  页眉: {header_content}")
                else:
                    print("  页眉: (无内容)")
            else:
                print("  页眉: (无)")
            
            # 检查页脚
            if section.footer:
                footer_content = []
                for paragraph in section.footer.paragraphs:
                    text = clean_text(paragraph.text)
                    if text:
                        footer_content.append(text)
                if footer_content:
                    print(f"  页脚: {footer_content}")
                else:
                    print("  页脚: (无内容)")
            else:
                print("  页脚: (无)")
        
        # 尝试更深入地读取所有段落
        print("\n完整文档内容探索:")
        print("-" * 30)
        
        # 读取所有段落（包括特殊位置的段落）
        all_content = []
        page_content = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = clean_text(paragraph.text)
            if text:  # 即使是空内容也记录位置
                alignment = paragraph.paragraph_format.alignment
                alignment_text = {0: "左对齐", 1: "居中", 2: "右对齐", 3: "两端对齐"}.get(alignment, "未知")
                
                para_info = {
                    'index': i,
                    'text': text,
                    'alignment': alignment_text
                }
                
                # 检查格式化信息
                formatting_info = []
                for j, run in enumerate(paragraph.runs):
                    run_text = clean_text(run.text)
                    if run_text:
                        run_format = []
                        if run.bold: run_format.append("粗体")
                        if run.italic: run_format.append("斜体")
                        if run.underline: run_format.append("下划线")
                        if run.font.name: run_format.append(f"字体:{run.font.name}")
                        if run.font.size: 
                            size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                            run_format.append(f"大小:{size}")
                        
                        if run_format:
                            formatting_info.append(f"Run {j+1}({' '.join(run_format)}): '{run_text}'")
                
                para_info['formatting'] = formatting_info
                all_content.append(para_info)
                
                # 如果是页码，则保存当前页内容
                if "第" in text and "页(共" in text:
                    page_content.append({
                        'page_text': text,
                        'content': all_content[-10:]  # 保存页码前的一些内容
                    })
        
        # 显示所有找到的内容
        for info in all_content:
            print(f"段落 {info['index']+1} ({info['alignment']}): {info['text']}")
            for fmt in info['formatting']:
                print(f"  - {fmt}")
        
        # 检查是否文档内容存储在其他地方（例如文本框）
        print("\n正在检查XML原始内容...")
        print("-" * 30)
        
        # 访问底层XML结构
        try:
            paragraphs_xml = []
            for i, p in enumerate(doc.paragraphs):
                if hasattr(p, '_element'):
                    xml_text = p._element.xml
                    if xml_text and len(xml_text) < 500:  # 只显示较短的XML
                        paragraphs_xml.append((i, xml_text))
            
            if paragraphs_xml:
                print("前几个段落的XML结构:")
                for idx, xml_content in paragraphs_xml[:5]:
                    print(f"  段落 {idx+1}: {xml_content[:200]}...")
            else:
                print("无法访问XML结构")
        except:
            print("无法访问XML结构")
        
    except Exception as e:
        print(f"探索文档结构时发生错误: {str(e)}")
        import traceback
        traceback.print_exc()

def extract_all_text_elements(file_path):
    """
    尝试从文档的所有部分提取文本元素
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return ""
    
    try:
        doc = Document(file_path)
        all_text = []
        
        # 添加主体段落
        print("正在从主体内容提取文本...")
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if text:
                all_text.append(text)
        
        # 添加表格中的文本
        print("正在从表格提取文本...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = clean_text(paragraph.text)
                        if text:
                            all_text.append(text)
        
        # 添加所有节的页眉页脚
        print("正在从页眉页脚提取文本...")
        for section in doc.sections:
            # 页眉
            if section.header:
                for paragraph in section.header.paragraphs:
                    text = clean_text(paragraph.text)
                    if text:
                        all_text.append(text)
            
            # 页脚
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    text = clean_text(paragraph.text)
                    if text:
                        all_text.append(text)
        
        return "\n".join(all_text)
    
    except Exception as e:
        print(f"提取文本时发生错误: {str(e)}")
        return ""

if __name__ == "__main__":
    file_path = r"D:\code\试卷\林业经济学（双语）试卷A-打印.docx"
    
    # 探索文档结构
    explore_document_structure(file_path)
    
    print("\n" + "="*60)
    print("从所有部分提取的完整文本:")
    print("-" * 30)
    
    # 提取所有文本元素
    all_text = extract_all_text_elements(file_path)
    print(all_text)