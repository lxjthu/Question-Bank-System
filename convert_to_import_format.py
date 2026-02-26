"""
Convert 林业经济学题库.docx to import template format.

Source format (per question type):
  单选题: ID, 难度, 知识点, 标签, 中文题干, 英文题干, A./A_en./B./B_en./C./C_en./D./D_en., 正确答案
  是非题: ID, 知识点, 标签, 中文题干, 英文题干, 答案, 解释, English Explanation
  论述题: ID, 知识点, 标签, 中文题干, 英文题干, 评分标准, English Scoring Guide
  计算题: ID, 知识点, 标签, 中文背景, 英文背景, [投资方案...], 参数, 要求, English Requirements

Target format:
  [题型][答案]
  中文题干
  [A]选项  [A_en] English option  ...
  <参考答案>...</参考答案>
  <解析>
  知识点: ...
  标签: ...
  英文题目: ...
  难度: ...
  </解析>
"""
import re
from docx import Document


def parse_source_doc(filepath):
    """Parse the source question bank document into structured question dicts."""
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]

    questions = []
    current = None
    current_type = None  # 'single', 'truefalse', 'essay', 'calc'

    for para_text in paragraphs:
        line = para_text.strip()
        if not line:
            # Empty line = question boundary
            if current:
                questions.append(current)
                current = None
            continue

        # Detect question type from ID line
        if re.match(r'^单选题ID:', line):
            current_type = 'single'
            current = {'type': 'single', 'lines': []}
        elif re.match(r'^是非题ID:', line):
            current_type = 'truefalse'
            current = {'type': 'truefalse', 'lines': []}
        elif re.match(r'^论述题ID:', line):
            current_type = 'essay'
            current = {'type': 'essay', 'lines': []}
        elif re.match(r'^计算题ID:', line):
            current_type = 'calc'
            current = {'type': 'calc', 'lines': []}
        elif line.startswith('林业经济学') and '题库' in line:
            # Section header, skip
            continue

        if current is not None:
            current['lines'].append(line)

    # Don't forget last question
    if current:
        questions.append(current)

    return questions


def extract_single_choice(lines):
    """Extract fields from single-choice question lines."""
    q = {
        'difficulty': '', 'knowledge_point': '', 'tags': '',
        'content_zh': '', 'content_en': '', 'answer': '',
        'options_zh': [], 'options_en': [],
    }
    for line in lines:
        if line.startswith('单选题ID:'):
            continue
        elif line.startswith('难度:') or line.startswith('难度：'):
            q['difficulty'] = re.sub(r'^难度[:：]\s*', '', line)
        elif line.startswith('知识点:') or line.startswith('知识点：'):
            q['knowledge_point'] = re.sub(r'^知识点[:：]\s*', '', line)
        elif line.startswith('标签:') or line.startswith('标签：'):
            q['tags'] = re.sub(r'^标签[:：]\s*', '', line)
        elif line.startswith('中文题干:') or line.startswith('中文题干：'):
            q['content_zh'] = re.sub(r'^中文题干[:：]\s*', '', line)
        elif line.startswith('英文题干:') or line.startswith('英文题干：'):
            q['content_en'] = re.sub(r'^英文题干[:：]\s*', '', line)
        elif line.startswith('正确答案:') or line.startswith('正确答案：'):
            q['answer'] = re.sub(r'^正确答案[:：]\s*', '', line).strip()
        elif re.match(r'^[A-D]\.\s', line):
            q['options_zh'].append(re.sub(r'^[A-D]\.\s*', '', line))
        elif re.match(r'^[A-D]_en\.\s', line):
            q['options_en'].append(re.sub(r'^[A-D]_en\.\s*', '', line))
    return q


def extract_truefalse(lines):
    """Extract fields from true/false question lines."""
    q = {
        'knowledge_point': '', 'tags': '',
        'content_zh': '', 'content_en': '', 'answer': '',
        'explanation_zh': '', 'explanation_en': '',
    }
    for line in lines:
        if line.startswith('是非题ID:'):
            continue
        elif line.startswith('知识点:') or line.startswith('知识点：'):
            q['knowledge_point'] = re.sub(r'^知识点[:：]\s*', '', line)
        elif line.startswith('标签:') or line.startswith('标签：'):
            q['tags'] = re.sub(r'^标签[:：]\s*', '', line)
        elif line.startswith('中文题干:') or line.startswith('中文题干：'):
            q['content_zh'] = re.sub(r'^中文题干[:：]\s*', '', line)
        elif line.startswith('英文题干:') or line.startswith('英文题干：'):
            q['content_en'] = re.sub(r'^英文题干[:：]\s*', '', line)
        elif line.startswith('答案:') or line.startswith('答案：'):
            q['answer'] = re.sub(r'^答案[:：]\s*', '', line).strip()
        elif line.startswith('解释:') or line.startswith('解释：'):
            q['explanation_zh'] = re.sub(r'^解释[:：]\s*', '', line)
        elif line.startswith('English Explanation:'):
            q['explanation_en'] = re.sub(r'^English Explanation[:：]\s*', '', line)
    return q


def extract_essay(lines):
    """Extract fields from essay question lines."""
    q = {
        'knowledge_point': '', 'tags': '',
        'content_zh': '', 'content_en': '',
        'scoring_zh': '', 'scoring_en': '',
    }
    for line in lines:
        if line.startswith('论述题ID:'):
            continue
        elif line.startswith('知识点:') or line.startswith('知识点：'):
            q['knowledge_point'] = re.sub(r'^知识点[:：]\s*', '', line)
        elif line.startswith('标签:') or line.startswith('标签：'):
            q['tags'] = re.sub(r'^标签[:：]\s*', '', line)
        elif line.startswith('中文题干:') or line.startswith('中文题干：'):
            q['content_zh'] = re.sub(r'^中文题干[:：]\s*', '', line)
        elif line.startswith('英文题干:') or line.startswith('英文题干：'):
            q['content_en'] = re.sub(r'^英文题干[:：]\s*', '', line)
        elif line.startswith('评分标准:') or line.startswith('评分标准：'):
            q['scoring_zh'] = re.sub(r'^评分标准[:：]\s*', '', line)
        elif line.startswith('English Scoring Guide:'):
            q['scoring_en'] = re.sub(r'^English Scoring Guide[:：]\s*', '', line)
    return q


def extract_calc(lines):
    """Extract fields from calculation question lines."""
    q = {
        'knowledge_point': '', 'tags': '',
        'bg_zh': '', 'bg_en': '',
        'extra_zh': [],  # investment options, params, etc (Chinese)
        'extra_en': [],  # investment options, params, etc (English)
        'params': '',
        'req_zh': '', 'req_en': '',
    }
    for line in lines:
        if line.startswith('计算题ID:'):
            continue
        elif line.startswith('知识点:') or line.startswith('知识点：'):
            q['knowledge_point'] = re.sub(r'^知识点[:：]\s*', '', line)
        elif line.startswith('标签:') or line.startswith('标签：'):
            q['tags'] = re.sub(r'^标签[:：]\s*', '', line)
        elif line.startswith('中文背景:') or line.startswith('中文背景：'):
            q['bg_zh'] = re.sub(r'^中文背景[:：]\s*', '', line)
        elif line.startswith('英文背景:') or line.startswith('英文背景：'):
            q['bg_en'] = re.sub(r'^英文背景[:：]\s*', '', line)
        elif line.startswith('参数:') or line.startswith('参数：'):
            q['params'] = re.sub(r'^参数[:：]\s*', '', line)
        elif line.startswith('要求:') or line.startswith('要求：'):
            q['req_zh'] = re.sub(r'^要求[:：]\s*', '', line)
        elif line.startswith('English Requirements:'):
            q['req_en'] = re.sub(r'^English Requirements[:：]\s*', '', line)
        elif re.match(r'^投资方案\d+[:：]', line):
            q['extra_zh'].append(line)
        elif re.match(r'^Investment Option \d+[:：]', line) or re.match(r'^Option \d+[:：]', line):
            q['extra_en'].append(line)
    return q


def write_import_doc(questions, output_path):
    """Write questions in import template format to a new docx file."""
    doc = Document()

    for i, raw_q in enumerate(questions):
        qtype = raw_q['type']

        if qtype == 'single':
            q = extract_single_choice(raw_q['lines'])
            # [单选][答案]
            header = f"[单选][{q['answer']}]"
            doc.add_paragraph(header)
            # Chinese content
            doc.add_paragraph(q['content_zh'])
            # Interleaved options
            for j in range(len(q['options_zh'])):
                letter = chr(65 + j)
                doc.add_paragraph(f"[{letter}]{q['options_zh'][j]}")
                if j < len(q['options_en']):
                    doc.add_paragraph(f"[{letter}_en] {q['options_en'][j]}")
            # Explanation block with metadata
            doc.add_paragraph('<解析>')
            if q['knowledge_point']:
                doc.add_paragraph(f"知识点: {q['knowledge_point']}")
            if q['tags']:
                doc.add_paragraph(f"标签: {q['tags']}")
            if q['content_en']:
                doc.add_paragraph(f"英文题目: {q['content_en']}")
            if q['difficulty']:
                doc.add_paragraph(f"难度: {q['difficulty']}")
            doc.add_paragraph('</解析>')

        elif qtype == 'truefalse':
            q = extract_truefalse(raw_q['lines'])
            header = f"[是非][{q['answer']}]"
            doc.add_paragraph(header)
            doc.add_paragraph(q['content_zh'])
            # Explanation
            doc.add_paragraph('<解析>')
            if q['explanation_zh']:
                doc.add_paragraph(q['explanation_zh'])
            if q['knowledge_point']:
                doc.add_paragraph(f"知识点: {q['knowledge_point']}")
            if q['tags']:
                doc.add_paragraph(f"标签: {q['tags']}")
            if q['content_en']:
                doc.add_paragraph(f"英文题目: {q['content_en']}")
            doc.add_paragraph('</解析>')

        elif qtype == 'essay':
            q = extract_essay(raw_q['lines'])
            doc.add_paragraph('[简答>论述]')
            doc.add_paragraph(q['content_zh'])
            # Reference answer = scoring criteria
            if q['scoring_zh']:
                doc.add_paragraph('<参考答案>')
                doc.add_paragraph(q['scoring_zh'])
                if q['scoring_en']:
                    doc.add_paragraph(q['scoring_en'])
                doc.add_paragraph('</参考答案>')
            # Explanation with metadata
            doc.add_paragraph('<解析>')
            if q['knowledge_point']:
                doc.add_paragraph(f"知识点: {q['knowledge_point']}")
            if q['tags']:
                doc.add_paragraph(f"标签: {q['tags']}")
            if q['content_en']:
                doc.add_paragraph(f"英文题目: {q['content_en']}")
            doc.add_paragraph('</解析>')

        elif qtype == 'calc':
            q = extract_calc(raw_q['lines'])
            doc.add_paragraph('[简答>计算]')
            # Combine Chinese content: background + extras + params + requirements
            content_parts = [q['bg_zh']]
            content_parts.extend(q['extra_zh'])
            if q['params']:
                content_parts.append(f"参数: {q['params']}")
            if q['req_zh']:
                content_parts.append(f"要求: {q['req_zh']}")
            doc.add_paragraph('\n'.join(p for p in content_parts if p))
            # Reference answer = English background + requirements
            ref_parts = []
            if q['bg_en']:
                ref_parts.append(q['bg_en'])
            ref_parts.extend(q['extra_en'])
            if q['req_en']:
                ref_parts.append(q['req_en'])
            if ref_parts:
                doc.add_paragraph('<参考答案>')
                doc.add_paragraph('\n'.join(ref_parts))
                doc.add_paragraph('</参考答案>')
            # Explanation with metadata
            doc.add_paragraph('<解析>')
            if q['knowledge_point']:
                doc.add_paragraph(f"知识点: {q['knowledge_point']}")
            if q['tags']:
                doc.add_paragraph(f"标签: {q['tags']}")
            if q['bg_en']:
                # Use English background as the English question stem
                en_content = q['bg_en']
                if q['req_en']:
                    en_content += ' ' + q['req_en']
                doc.add_paragraph(f"英文题目: {en_content}")
            doc.add_paragraph('</解析>')

        # Add empty paragraph as separator between questions
        doc.add_paragraph('')

    doc.save(output_path)
    print(f"Saved to: {output_path}")


def main():
    src = '林业经济学题库.docx'
    dst = '林业经济学题库_导入格式.docx'

    print(f"Reading: {src}")
    questions = parse_source_doc(src)

    type_counts = {}
    for q in questions:
        type_counts[q['type']] = type_counts.get(q['type'], 0) + 1
    print(f"Parsed {len(questions)} questions: {type_counts}")

    write_import_doc(questions, dst)

    # Verify by running the converter
    print("\nVerifying with converter...")
    from word_to_csv_converter import convert_word_to_csv, parse_csv_to_questions
    csv_path = convert_word_to_csv(dst)
    parsed = parse_csv_to_questions(csv_path)
    print(f"Converter parsed: {len(parsed)} questions")

    has_en = sum(1 for q in parsed if q.get('content_en'))
    has_opts = sum(1 for q in parsed if q.get('options_en'))
    has_kp = sum(1 for q in parsed if q.get('knowledge_point'))
    has_diff = sum(1 for q in parsed if q.get('difficulty'))
    print(f"  with content_en: {has_en}")
    print(f"  with options_en: {has_opts}")
    print(f"  with knowledge_point: {has_kp}")
    print(f"  with difficulty: {has_diff}")

    # Show first question of each type
    types_seen = set()
    for q in parsed:
        t = q['type']
        if t not in types_seen:
            types_seen.add(t)
            print(f"\n  [{t}] content: {q['content'][:60]}...")
            if q.get('content_en'):
                print(f"    content_en: {q['content_en'][:60]}...")
            if q.get('options'):
                print(f"    options: {q['options'][:2]}...")
            if q.get('options_en'):
                print(f"    options_en: {q['options_en'][:2]}...")
            print(f"    answer: {q.get('answer')}")
            print(f"    kp: {q.get('knowledge_point')}, diff: {q.get('difficulty')}")

    import os
    if os.path.exists(csv_path):
        os.remove(csv_path)


if __name__ == '__main__':
    main()
