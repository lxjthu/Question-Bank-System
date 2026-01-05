"""
API路由定义
"""
from flask import Blueprint, request, jsonify, send_file
import json
import os
import tempfile
from werkzeug.utils import secure_filename
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from app.models import question_bank, question_usage, saved_exams
from app.utils import allowed_file, WordToCsvConverter

# 创建蓝图
api_bp = Blueprint('api', __name__)

@api_bp.route('/questions', methods=['GET'])
def get_questions():
    """获取所有题目"""
    return jsonify(question_bank)

@api_bp.route('/questions', methods=['POST'])
def add_question():
    """添加新题目"""
    question_data = request.json

    # 根据题型添加到对应的数组
    if question_data['type'] == 'single_choice':
        question_bank['single_choice_questions'].append(question_data)
    elif question_data['type'] == 'true_false':
        question_bank['true_false_questions'].append(question_data)
    elif question_data['type'] == 'essay':
        question_bank['essay_questions'].append(question_data)
    elif question_data['type'] == 'calculation':
        question_bank['calculation_questions'].append(question_data)

    return jsonify({"success": True, "message": "题目添加成功"})

@api_bp.route('/questions/<question_type>/<int:index>', methods=['PUT'])
def update_question(question_type, index):
    """更新题目"""
    question_data = request.json

    if question_type == 'single_choice':
        question_bank['single_choice_questions'][index] = question_data
    elif question_type == 'true_false':
        question_bank['true_false_questions'][index] = question_data
    elif question_type == 'essay':
        question_bank['essay_questions'][index] = question_data
    elif question_type == 'calculation':
        question_bank['calculation_questions'][index] = question_data

    return jsonify({"success": True, "message": "题目更新成功"})

@api_bp.route('/questions/<question_type>/<int:index>', methods=['DELETE'])
def delete_question(question_type, index):
    """删除题目"""
    if question_type == 'single_choice':
        if 0 <= index < len(question_bank['single_choice_questions']):
            question_bank['single_choice_questions'].pop(index)
    elif question_type == 'true_false':
        if 0 <= index < len(question_bank['true_false_questions']):
            question_bank['true_false_questions'].pop(index)
    elif question_type == 'essay':
        if 0 <= index < len(question_bank['essay_questions']):
            question_bank['essay_questions'].pop(index)
    elif question_type == 'calculation':
        if 0 <= index < len(question_bank['calculation_questions']):
            question_bank['calculation_questions'].pop(index)

    return jsonify({"success": True, "message": "题目删除成功"})

@api_bp.route('/convert-word', methods=['POST'])
def convert_word():
    """转换Word文档为题库"""
    if 'file' not in request.files:
        return jsonify({"success": False, "message": "没有上传文件"})

    file = request.files['file']
    if file.filename == '':
        return jsonify({"success": False, "message": "未选择文件"})

    if file and allowed_file(file.filename):
        # 保存上传的文件
        filename = secure_filename(file.filename)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            file.save(tmp_file.name)
            temp_path = tmp_file.name

        try:
            # 创建转换器实例
            converter = WordToCsvConverter()

            # 从Word文档直接解析题目（不通过CSV中间步骤）
            questions = converter.parse_questions_from_word_doc(temp_path)

            if questions:
                # 将解析的题目添加到题库
                for question in questions:
                    qtype = question.get('type')
                    if qtype == 'single_choice':
                        # 确保选项格式正确
                        if 'options' not in question:
                            question['options'] = {"A": "", "A_en": "", "B": "", "B_en": "", "C": "", "C_en": "", "D": "", "D_en": ""}
                        question_bank['single_choice_questions'].append(question)
                    elif qtype == 'true_false':
                        question_bank['true_false_questions'].append(question)
                    elif qtype == 'essay':
                        question_bank['essay_questions'].append(question)
                    elif qtype == 'calculation':
                        question_bank['calculation_questions'].append(question)

                return jsonify({"success": True, "message": f"成功导入 {len(questions)} 道题目", "warnings": getattr(converter, 'warnings', [])})
            else:
                return jsonify({"success": False, "message": "未解析到任何题目"})

        except Exception as e:
            return jsonify({"success": False, "message": f"转换过程中发生错误: {str(e)}"})

        finally:
            # 清理临时文件
            try:
                os.unlink(temp_path)
                if os.path.exists(temp_path + '.csv'):
                    os.unlink(temp_path + '.csv')
            except:
                pass
    else:
        return jsonify({"success": False, "message": "仅支持.docx格式的文件"})

@api_bp.route('/export/json', methods=['GET'])
def export_json():
    """导出题库为JSON"""
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.json', mode='w', encoding='utf-8')
    json.dump(question_bank, temp_file, ensure_ascii=False, indent=2)
    temp_file.close()

    return send_file(temp_file.name, as_attachment=True, download_name='question_bank.json',
                     mimetype='application/json')

@api_bp.route('/export/csv', methods=['GET'])
def export_csv():
    """导出题库为CSV"""
    import csv
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.csv', mode='w', encoding='utf-8-sig')

    writer = csv.writer(temp_file)

    # 写入表头
    writer.writerow(['ID', 'Type', 'Difficulty', 'Chinese Stem', 'English Stem',
                     'Option A (CH)', 'Option A (EN)', 'Option B (CH)', 'Option B (EN)',
                     'Option C (CH)', 'Option C (EN)', 'Option D (CH)', 'Option D (EN)',
                     'Correct Answer', 'Knowledge Point', 'Tags'])

    # 写入单选题
    for q in question_bank['single_choice_questions']:
        writer.writerow([
            q['id'], 'single_choice', q['difficulty'], q['chinese_stem'], q['english_stem'],
            q['options']['A'], q['options']['A_en'], q['options']['B'], q['options']['B_en'],
            q['options']['C'], q['options']['C_en'], q['options']['D'], q['options']['D_en'],
            q['correct_answer'], q['knowledge_point'], ','.join(q['tags'])
        ])

    # 写入其他题型
    for q in question_bank['true_false_questions']:
        writer.writerow([
            q['id'], 'true_false', '', q['chinese_stem'], q['english_stem'],
            q['correct_answer'], '', q['explanation'], q['explanation_en'],
            '', '', '', '',
            q['correct_answer'], q['knowledge_point'], ','.join(q['tags'])
        ])

    for q in question_bank['essay_questions']:
        writer.writerow([
            q['id'], 'essay', '', q['chinese_stem'], q['english_stem'],
            '', '', '', '',
            '', '', '', '',
            '', q['knowledge_point'], ','.join(q['tags'])
        ])

    for q in question_bank['calculation_questions']:
        writer.writerow([
            q['id'], 'calculation', '', q['context']['chinese'], q['context']['english'],
            '', '', '', '',
            '', '', '', '',
            '', q['knowledge_point'], ','.join(q['tags'])
        ])

    temp_file.close()

    return send_file(temp_file.name, as_attachment=True, download_name='question_bank.csv',
                     mimetype='text/csv')

@api_bp.route('/generate-exam', methods=['POST'])
def generate_exam():
    """生成试卷"""
    config = request.json

    # 检查是否有手动选择的题目
    selected_questions = config.get('selected_questions', [])

    exam_content = f"""{config['school']}{config['academic_year']}期末考试试卷
课程名称：《{config['course_name']}（双语）》 （{config['paper_type']}）卷
课程代号：B0600432
考试形式：闭卷、笔试
使用对象：

题号	一	二	三	四	五	六	七	总分	总分人
分值
得分

 -------------------------------------------------------------------------------

得分    评阅人


"""

    # 按题型生成试卷
    question_types_map = {
        'single_choice': '单选题',
        'true_false': '是非题',
        'essay': '论述题',
        'calculation': '计算题'
    }

    # 生成各题型内容
    for qtype, qtype_name in question_types_map.items():
        q_config = config['question_types'][qtype]
        count = q_config['count']
        score_per_question = q_config['score_per_question']

        if count > 0:
            # 从手动选择的题目中筛选当前题型
            selected_of_type = [q for q in selected_questions if q['type'] == qtype]

            # 如果手动选择的题目数量不足，从题库中补充
            if len(selected_of_type) < count:
                remaining_count = count - len(selected_of_type)
                # 从题库中随机选择补充题目
                all_questions_of_type = question_bank[f"{qtype}_questions"]
                import random
                # 避免选择已手动选择的题目
                available_questions = [q for q in all_questions_of_type if q['id'] not in [sq['id'] for sq in selected_of_type]]
                additional_questions = random.sample(available_questions, min(remaining_count, len(available_questions)))
                all_questions_for_type = selected_of_type + additional_questions
            else:
                all_questions_for_type = selected_of_type[:count]

            exam_content += f"{qtype_name}：（共{len(all_questions_for_type)}题，每题{score_per_question}分，共{len(all_questions_for_type) * score_per_question}分）\n"

            if qtype == 'single_choice':
                exam_content += '请将答案统一填写在答题栏里\n'
                exam_content += '题号\t' + '\t'.join([str(i+1) for i in range(len(all_questions_for_type))]) + '\n'
                exam_content += '答案\t' + '\t'.join([' ' for _ in range(len(all_questions_for_type))]) + '\n\n'

            for i, q in enumerate(all_questions_for_type):
                q_data = q.get('data', q) if isinstance(q, dict) else q
                exam_content += f"{i+1}. {q_data.get('chinese_stem', '题目内容缺失')}\n"

                if qtype == 'single_choice':
                    exam_content += f"A. {q_data.get('options', {}).get('A', '')}\n"
                    exam_content += f"B. {q_data.get('options', {}).get('B', '')}\n"
                    exam_content += f"C. {q_data.get('options', {}).get('C', '')}\n"
                    exam_content += f"D. {q_data.get('options', {}).get('D', '')}\n\n"
                elif qtype == 'true_false':
                    exam_content += f"正确答案: {q_data.get('correct_answer', '')}\n\n"
                elif qtype == 'essay':
                    scoring_guide = q_data.get('scoring_guide', [])
                    exam_content += f"评分标准: {'; '.join(scoring_guide) if isinstance(scoring_guide, list) else str(scoring_guide)}\n\n"
                elif qtype == 'calculation':
                    context = q_data.get('context', {})
                    requirements = q_data.get('requirements', {})
                    exam_content += f"背景: {context.get('chinese', '')}\n"
                    req_chinese = requirements.get('chinese', [])
                    exam_content += f"要求: {'; '.join(req_chinese) if isinstance(req_chinese, list) else str(req_chinese)}\n\n"

    return jsonify({"exam_content": exam_content})

@api_bp.route('/search-questions', methods=['POST'])
def search_questions():
    """搜索题目"""
    try:
        data = request.json
        if not data:
            return jsonify({"success": False, "message": "请求数据为空"}), 400

        keyword = data.get('keyword', '').lower()
        question_type = data.get('type', '')
        knowledge_point = data.get('knowledge_point', '')

        results = {
            "single_choice_questions": [],
            "true_false_questions": [],
            "essay_questions": [],
            "calculation_questions": []
        }

        # 搜索单选题
        for q in question_bank['single_choice_questions']:
            match = False
            if keyword and (keyword in q.get('chinese_stem', '').lower() or
                           keyword in q.get('english_stem', '').lower() or
                           keyword in q.get('knowledge_point', '').lower()):
                match = True
            elif not keyword:
                match = True

            if question_type and q.get('type') != question_type:
                match = False
            if knowledge_point and knowledge_point != q.get('knowledge_point', ''):
                match = False

            if match:
                results['single_choice_questions'].append(q)

        # 搜索是非题
        for q in question_bank['true_false_questions']:
            match = False
            if keyword and (keyword in q.get('chinese_stem', '').lower() or
                           keyword in q.get('english_stem', '').lower() or
                           keyword in q.get('knowledge_point', '').lower()):
                match = True
            elif not keyword:
                match = True

            if question_type and q.get('type') != question_type:
                match = False
            if knowledge_point and knowledge_point != q.get('knowledge_point', ''):
                match = False

            if match:
                results['true_false_questions'].append(q)

        # 搜索论述题
        for q in question_bank['essay_questions']:
            match = False
            if keyword and (keyword in q.get('chinese_stem', '').lower() or
                           keyword in q.get('english_stem', '').lower() or
                           keyword in q.get('knowledge_point', '').lower()):
                match = True
            elif not keyword:
                match = True

            if question_type and q.get('type') != question_type:
                match = False
            if knowledge_point and knowledge_point != q.get('knowledge_point', ''):
                match = False

            if match:
                results['essay_questions'].append(q)

        # 搜索计算题
        for q in question_bank['calculation_questions']:
            match = False
            if keyword and (keyword in q.get('context', {}).get('chinese', '').lower() or
                           keyword in q.get('context', {}).get('english', '').lower() or
                           keyword in q.get('knowledge_point', '').lower()):
                match = True
            elif not keyword:
                match = True

            if question_type and q.get('type') != question_type:
                match = False
            if knowledge_point and knowledge_point != q.get('knowledge_point', ''):
                match = False

            if match:
                results['calculation_questions'].append(q)

        return jsonify(results)
    except Exception as e:
        return jsonify({"success": False, "message": f"搜索题目时发生错误: {str(e)}"}), 500

@api_bp.route('/save-exam', methods=['POST'])
def save_exam():
    """保存试卷组卷"""
    try:
        data = request.json
        if not data:
            return jsonify({"success": False, "message": "请求数据为空"}), 400

        exam_data = data.get('exam_data', {})
        exam_id = data.get('exam_id', f"exam_{int(datetime.now().timestamp())}")

        saved_exams[exam_id] = {
            'id': exam_id,
            'name': exam_data.get('name', f'试卷_{exam_id}'),
            'config': exam_data.get('config', {}),
            'questions': exam_data.get('questions', {}),
            'created_at': datetime.now().isoformat()
        }

        return jsonify({"success": True, "exam_id": exam_id, "message": "试卷保存成功"})
    except Exception as e:
        return jsonify({"success": False, "message": f"保存试卷时发生错误: {str(e)}"}), 500

@api_bp.route('/load-exam/<exam_id>', methods=['GET'])
def load_exam(exam_id):
    """加载保存的试卷"""
    try:
        if exam_id in saved_exams:
            return jsonify(saved_exams[exam_id])
        else:
            return jsonify({"success": False, "message": "试卷不存在"}), 404
    except Exception as e:
        return jsonify({"success": False, "message": f"加载试卷时发生错误: {str(e)}"}), 500

@api_bp.route('/list-saved-exams', methods=['GET'])
def list_saved_exams():
    """列出所有保存的试卷"""
    try:
        return jsonify(list(saved_exams.values()))
    except Exception as e:
        return jsonify({"success": False, "message": f"获取试卷列表时发生错误: {str(e)}"}), 500

@api_bp.route('/delete-exam/<exam_id>', methods=['DELETE'])
def delete_exam(exam_id):
    """删除保存的试卷"""
    try:
        if exam_id in saved_exams:
            del saved_exams[exam_id]
            return jsonify({"success": True, "message": "试卷删除成功"})
        else:
            return jsonify({"success": False, "message": "试卷不存在"}), 404
    except Exception as e:
        return jsonify({"success": False, "message": f"删除试卷时发生错误: {str(e)}"}), 500

@api_bp.route('/download-template/<template_type>', methods=['GET'])
def download_template(template_type):
    """下载Word模板"""
    if template_type not in ['single_choice', 'true_false', 'essay', 'calculation']:
        return jsonify({"success": False, "message": "无效的模板类型"}), 400

    # 创建转换器实例
    converter = WordToCsvConverter()

    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        converter.create_word_template(template_type, tmp_file.name)
        temp_path = tmp_file.name

    return send_file(temp_path, as_attachment=True,
                     download_name=f'{template_type}_template.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@api_bp.route('/export-word', methods=['POST'])
def export_word():
    """导出试卷为Word文档（增强版）"""
    try:
        exam_data = request.json

        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            # 创建Word文档
            doc = Document()

            # 设置页面边距
            section = doc.sections[0]
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

            # 添加标题
            if 'header' in exam_data:
                header_para = doc.add_paragraph()
                header_run = header_para.add_run(exam_data['header'])
                header_run.font.name = 'SimSun'
                header_run.font.size = Pt(16)
                header_run.font.bold = True
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 添加空行
                doc.add_paragraph()

            # 添加题型和题目
            if 'questions' in exam_data:
                questions = exam_data['questions']

                # 单选题
                if 'single_choice' in questions and questions['single_choice']:
                    sc_questions = questions['single_choice']
                    if sc_questions:
                        sc_config = exam_data.get('question_types', {}).get('single_choice', {})
                        score_per_question = sc_config.get('score_per_question', 1)

                        section_title = f"一、单选题（每题{score_per_question}分，共{len(sc_questions) * score_per_question}分）"
                        section_para = doc.add_paragraph()
                        section_run = section_para.add_run(section_title)
                        section_run.font.name = 'SimSun'
                        section_run.font.size = Pt(12)
                        section_run.font.bold = True

                        for i, q in enumerate(sc_questions, 1):
                            # 题目
                            q_para = doc.add_paragraph()
                            q_run = q_para.add_run(f"{i}. {q.get('chinese_stem', '')}")
                            q_run.font.name = 'SimSun'
                            q_run.font.size = Pt(11)

                            # 选项
                            options = q.get('options', {})
                            for opt_key in ['A', 'B', 'C', 'D']:
                                if opt_key in options:
                                    opt_para = doc.add_paragraph()
                                    opt_run = opt_para.add_run(f"   {opt_key}. {options[opt_key]}")
                                    opt_run.font.name = 'SimSun'
                                    opt_run.font.size = Pt(10)

                            # 答案
                            correct_answer = q.get('correct_answer', '')
                            if correct_answer:
                                ans_para = doc.add_paragraph()
                                ans_run = ans_para.add_run(f"   正确答案：{correct_answer}")
                                ans_run.font.name = 'SimSun'
                                ans_run.font.size = Pt(10)
                                ans_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色答案

                # 是非题
                if 'true_false' in questions and questions['true_false']:
                    tf_questions = questions['true_false']
                    if tf_questions:
                        tf_config = exam_data.get('question_types', {}).get('true_false', {})
                        score_per_question = tf_config.get('score_per_question', 1)

                        section_title = f"二、是非题（每题{score_per_question}分，共{len(tf_questions) * score_per_question}分）"
                        section_para = doc.add_paragraph()
                        section_run = section_para.add_run(section_title)
                        section_run.font.name = 'SimSun'
                        section_run.font.size = Pt(12)
                        section_run.font.bold = True

                        for i, q in enumerate(tf_questions, 1):
                            q_para = doc.add_paragraph()
                            q_run = q_para.add_run(f"{i}. {q.get('chinese_stem', '')}")
                            q_run.font.name = 'SimSun'
                            q_run.font.size = Pt(11)

                            # 答案
                            correct_answer = q.get('correct_answer', '')
                            if correct_answer:
                                ans_para = doc.add_paragraph()
                                ans_run = ans_para.add_run(f"   答案：{correct_answer}")
                                ans_run.font.name = 'SimSun'
                                ans_run.font.size = Pt(10)
                                ans_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色答案

                # 论述题
                if 'essay' in questions and questions['essay']:
                    es_questions = questions['essay']
                    if es_questions:
                        es_config = exam_data.get('question_types', {}).get('essay', {})
                        score_per_question = es_config.get('score_per_question', 1)

                        section_title = f"三、论述题（每题{score_per_question}分，共{len(es_questions) * score_per_question}分）"
                        section_para = doc.add_paragraph()
                        section_run = section_para.add_run(section_title)
                        section_run.font.name = 'SimSun'
                        section_run.font.size = Pt(12)
                        section_run.font.bold = True

                        for i, q in enumerate(es_questions, 1):
                            q_para = doc.add_paragraph()
                            q_run = q_para.add_run(f"{i}. {q.get('chinese_stem', '')}")
                            q_run.font.name = 'SimSun'
                            q_run.font.size = Pt(11)

                # 计算题
                if 'calculation' in questions and questions['calculation']:
                    calc_questions = questions['calculation']
                    if calc_questions:
                        calc_config = exam_data.get('question_types', {}).get('calculation', {})
                        score_per_question = calc_config.get('score_per_question', 1)

                        section_title = f"四、计算题（每题{score_per_question}分，共{len(calc_questions) * score_per_question}分）"
                        section_para = doc.add_paragraph()
                        section_run = section_para.add_run(section_title)
                        section_run.font.name = 'SimSun'
                        section_run.font.size = Pt(12)
                        section_run.font.bold = True

                        for i, q in enumerate(calc_questions, 1):
                            context = q.get('context', {})
                            q_para = doc.add_paragraph()
                            q_run = q_para.add_run(f"{i}. {context.get('chinese', '')}")
                            q_run.font.name = 'SimSun'
                            q_run.font.size = Pt(11)

            # 保存文档
            doc.save(tmp_file.name)

        return send_file(tmp_file.name, as_attachment=True,
                         download_name='exam.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({"success": False, "message": f"导出Word文档时发生错误: {str(e)}"}), 500