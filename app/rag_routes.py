"""RAG API Blueprint — AI 智能出题后端。

端点：
  GET    /api/rag/docs                  列出已摄入文档
  GET    /api/rag/docs/<doc_id>/meta    获取文档章节/节名
  DELETE /api/rag/docs/<doc_id>         删除文档
  POST   /api/rag/ingest                上传并摄入文档
  GET    /api/rag/tasks/<task_id>       轮询 OCR 后台任务状态
  POST   /api/rag/generate              RAG 检索 + DeepSeek 出题
"""
from __future__ import annotations

import os
import sys
import uuid
import sqlite3
import threading
from pathlib import Path

from flask import Blueprint, request, jsonify

rag_bp = Blueprint('rag', __name__)

# ── 单例组件（懒加载）─────────────────────────────────────────────────────────
_lock = threading.Lock()
_components = None  # (Embedder, VectorStore, BM25Index) — 全局共享，避免双 QdrantClient 锁冲突
_ingestor = None    # rag_pipeline.ingest.Ingestor
_retriever = None   # rag_pipeline.retriever.HybridRetriever

# OCR 后台任务状态表  {task_id: {status, message, doc_id}}
_ocr_tasks: dict = {}

# DS 直出模式 OCR 后台任务状态表  {task_id: {...}}
# 注：_ds_tasks 在下方 DS 模式区域定义，此处仅声明占位供模块级函数引用
_ds_tasks: dict = {}

# RAG 文档上传目录
_UPLOAD_DIR = Path(__file__).parent.parent / "rag_uploads"


# ── 路径工具 ──────────────────────────────────────────────────────────────────

def _project_root() -> Path:
    return Path(__file__).parent.parent


def _rag_pipeline_dir() -> Path:
    return _project_root() / "rag_pipeline"


def _rag_db_path() -> Path:
    return _rag_pipeline_dir() / "kg.db"


def _ensure_upload_dir():
    _UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def _add_to_path():
    """将项目根目录加入 sys.path，使 rag_pipeline / pptx_ocr 可直接导入。"""
    root = str(_project_root())
    if root not in sys.path:
        sys.path.insert(0, root)


# ── 单例懒加载 ────────────────────────────────────────────────────────────────

def _get_components():
    """
    初始化并返回共享的 (Embedder, VectorStore, BM25Index)。

    Qdrant 本地文件模式同一时间只允许一个 QdrantClient 持有锁，
    因此 Ingestor 和 HybridRetriever 必须共用同一个 VectorStore 实例。
    本函数保证全局只创建一次，只在已持有 _lock 的情况下调用。
    """
    global _components
    if _components is not None:
        return _components
    _add_to_path()
    from rag_pipeline.config import EMBEDDING_MODEL, QDRANT_PATH, BM25_PATH
    from rag_pipeline.embedder import Embedder
    from rag_pipeline.vector_store import VectorStore
    from rag_pipeline.bm25_index import BM25Index

    emb = Embedder(model_name=EMBEDDING_MODEL)
    vs = VectorStore(QDRANT_PATH)
    vs.ensure_collection(emb.dim)
    bm25 = BM25Index(BM25_PATH)
    _components = (emb, vs, bm25)
    return _components


def _get_ingestor():
    global _ingestor
    if _ingestor is not None:
        return _ingestor
    with _lock:
        if _ingestor is not None:
            return _ingestor
        emb, vs, bm25 = _get_components()
        from rag_pipeline.ingest import Ingestor
        _ingestor = Ingestor(emb=emb, vs=vs, bm25=bm25)
    return _ingestor


def _get_retriever():
    global _retriever
    if _retriever is not None:
        return _retriever
    with _lock:
        if _retriever is not None:
            return _retriever
        emb, vs, bm25 = _get_components()
        from rag_pipeline.retriever import HybridRetriever
        _retriever = HybridRetriever(vector_store=vs, bm25_index=bm25, embedder=emb)
    return _retriever


# ── SQLite 帮助函数 ───────────────────────────────────────────────────────────

def _db_conn():
    conn = sqlite3.connect(str(_rag_db_path()))
    conn.row_factory = sqlite3.Row
    return conn


# ── .env 读写工具 ─────────────────────────────────────────────────────────────

def _env_path() -> Path:
    return _project_root() / ".env"


def _read_env_vars(*keys) -> dict:
    """从 .env 文件读取指定 key，不影响进程环境变量。"""
    from dotenv import dotenv_values
    vals = dotenv_values(str(_env_path())) if _env_path().exists() else {}
    return {k: vals.get(k, '') for k in keys}


def _write_env_vars(updates: dict) -> None:
    """原地更新 .env 文件中的 key=value 行；不存在的 key 追加到末尾。"""
    env_file = _env_path()
    if env_file.exists():
        lines = env_file.read_text(encoding='utf-8').splitlines(keepends=True)
    else:
        lines = []

    updated_keys: set = set()
    new_lines = []
    for line in lines:
        stripped = line.rstrip('\r\n')
        # 跳过注释和空行，直接保留
        if not stripped or stripped.startswith('#'):
            new_lines.append(line if line.endswith('\n') else line + '\n')
            continue
        if '=' in stripped:
            key = stripped.split('=', 1)[0].strip()
            if key in updates:
                new_lines.append(f"{key}={updates[key]}\n")
                updated_keys.add(key)
                continue
        new_lines.append(line if line.endswith('\n') else line + '\n')

    # 追加尚未出现的 key
    for key, val in updates.items():
        if key not in updated_keys:
            new_lines.append(f"{key}={val}\n")

    env_file.write_text(''.join(new_lines), encoding='utf-8')


# ═══════════════════════════════════════════════════════════════════════════════
# 端点
# ═══════════════════════════════════════════════════════════════════════════════

@rag_bp.route('/api/rag/docs', methods=['GET'])
def list_rag_docs():
    """列出所有已摄入文档及其 chunk 数量。"""
    try:
        if not _rag_db_path().exists():
            return jsonify([])
        with _db_conn() as conn:
            rows = conn.execute(
                "SELECT doc_id, source_type, COUNT(*) as chunk_count "
                "FROM chunks GROUP BY doc_id ORDER BY doc_id"
            ).fetchall()
        return jsonify([
            {
                'doc_id': r['doc_id'],
                'source_type': r['source_type'],
                'chunk_count': r['chunk_count'],
            }
            for r in rows
        ])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@rag_bp.route('/api/rag/docs/<path:doc_id>/meta', methods=['GET'])
def rag_doc_meta(doc_id):
    """获取文档的章节名和节名（用于前端筛选）。

    sections 响应字段：
      num          — section_num
      name         — section_name（原始内部名，如 "Slide 1"，用于生成 API 查询）
      chapter_num  — 所属章节编号
      chapter_name — 所属章节名称（用于前端按章节联动过滤知识点）
      display_name — 展示用名称（slides 类型取首行内容，textbook 同 name）
    """
    try:
        if not _rag_db_path().exists():
            return jsonify({'chapters': [], 'sections': []})
        with _db_conn() as conn:
            chapters = conn.execute(
                "SELECT DISTINCT chapter_num, chapter_name FROM chunks "
                "WHERE doc_id=? AND chapter_name != '' ORDER BY chapter_num",
                (doc_id,),
            ).fetchall()

            # 判断文档类型，用于决定 display_name 格式
            src_row = conn.execute(
                "SELECT source_type FROM chunks WHERE doc_id=? LIMIT 1", (doc_id,)
            ).fetchone()
            source_type = src_row['source_type'] if src_row else ''

            # 获取每节的去重信息（含所属章节）
            sections_raw = conn.execute(
                "SELECT DISTINCT section_num, section_name, chapter_num, chapter_name "
                "FROM chunks WHERE doc_id=? AND section_name != '' ORDER BY section_num",
                (doc_id,),
            ).fetchall()

            sections = []
            for r in sections_raw:
                display_name = r['section_name']
                if source_type == 'slides':
                    # 取该页第一个 chunk 的文本首行作为知识点描述
                    text_row = conn.execute(
                        "SELECT text FROM chunks WHERE doc_id=? AND section_num=? "
                        "ORDER BY rowid LIMIT 1",
                        (doc_id, r['section_num']),
                    ).fetchone()
                    if text_row and text_row['text']:
                        first_line = text_row['text'].strip().split('\n')[0].strip()[:60]
                        if first_line:
                            display_name = f"第{r['section_num']}页: {first_line}"
                sections.append({
                    'num': r['section_num'],
                    'name': r['section_name'],
                    'chapter_num': r['chapter_num'],
                    'chapter_name': r['chapter_name'],
                    'display_name': display_name,
                })

        return jsonify({
            'chapters': [
                {'num': r['chapter_num'], 'name': r['chapter_name']} for r in chapters
            ],
            'sections': sections,
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@rag_bp.route('/api/rag/docs/<path:doc_id>', methods=['DELETE'])
def rag_delete_doc(doc_id):
    """从向量库、SQLite chunks 和知识图谱表中删除文档。"""
    try:
        _add_to_path()
        ingestor = _get_ingestor()
        ingestor.delete_doc(doc_id)
        # 同步删除该文档的幻灯片 KG（若有）
        from rag_pipeline import db as rag_db
        rag_db.delete_kg_by_doc(doc_id)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@rag_bp.route('/api/rag/ingest', methods=['POST'])
def rag_ingest():
    """上传文件并摄入知识库。"""
    _ensure_upload_dir()

    if 'file' not in request.files:
        return jsonify({'error': '未提供文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    doc_type = request.form.get('doc_type', 'auto')  # auto / textbook / slides
    filename = file.filename
    suffix = Path(filename).suffix.lower()
    stem = Path(filename).stem
    doc_id = stem[:40]  # 截断过长名称

    save_path = _UPLOAD_DIR / filename
    file.save(str(save_path))

    # ── 文本格式：同步摄入 ─────────────────────────────────────────────────
    if suffix in ('.md', '.txt'):
        try:
            ingestor = _get_ingestor()
            if doc_type == 'slides':
                n = ingestor.ingest_slides(doc_id, save_path, doc_title=stem)
            else:
                n = ingestor.ingest_textbook(doc_id, save_path)
            return jsonify({'success': True, 'doc_id': doc_id, 'chunk_count': n})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # ── DOCX：提取段落 → 临时 MD → 摄入 ──────────────────────────────────
    if suffix == '.docx':
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(save_path))
            md_lines = [f'# {stem}', '']
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    md_lines.append(text)
                    md_lines.append('')
            tmp_md = _UPLOAD_DIR / (stem + '_converted.md')
            tmp_md.write_text('\n'.join(md_lines), encoding='utf-8')
            ingestor = _get_ingestor()
            n = ingestor.ingest_textbook(doc_id, tmp_md)
            return jsonify({'success': True, 'doc_id': doc_id, 'chunk_count': n})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # ── PPTX / PDF：异步 OCR ─────────────────────────────────────────────
    if suffix in ('.pptx', '.ppt', '.pdf'):
        task_id = str(uuid.uuid4())[:8]
        subject_hint = request.form.get('subject_hint', stem)  # 学科提示，默认取文件名
        _ocr_tasks[task_id] = {
            'status': 'running',
            'message': 'OCR 处理中，请稍候...',
            'doc_id': doc_id,
        }

        def _ocr_worker(task_id, suffix, save_path, doc_id, doc_type, stem, subject_hint):
            try:
                _add_to_path()
                # 每次 OCR 任务启动时从 .env 读取最新配置，确保界面保存后立即生效
                ocr_cfg = _read_env_vars('PADDLEOCR_API_URL', 'PADDLEOCR_TOKEN', 'PADDLEOCR_TIMEOUT')
                ocr_url     = ocr_cfg.get('PADDLEOCR_API_URL', '')
                ocr_token   = ocr_cfg.get('PADDLEOCR_TOKEN', '')
                ocr_timeout = int(ocr_cfg.get('PADDLEOCR_TIMEOUT') or 120)

                ocr_output = _UPLOAD_DIR / (stem + '_ocr')
                if suffix == '.pdf':
                    from pptx_ocr.pipeline import process_pdf
                    result_md = process_pdf(
                        save_path, output_dir=ocr_output,
                        api_url=ocr_url, api_token=ocr_token, api_timeout=ocr_timeout,
                    )
                else:
                    from pptx_ocr.pipeline import process_pptx
                    result_md = process_pptx(
                        save_path, output_dir=ocr_output,
                        api_url=ocr_url, api_token=ocr_token, api_timeout=ocr_timeout,
                    )

                ingestor = _get_ingestor()
                use_slides = (suffix in ('.pptx', '.ppt') or doc_type == 'slides')
                if use_slides:
                    n = ingestor.ingest_slides(doc_id, result_md, doc_title=stem)
                    # ── 幻灯片专属：调用 DeepSeek 构建知识图谱 ──────────────
                    _ocr_tasks[task_id]['message'] = (
                        f'文本块摄入完成（{n} 块），正在调用 DeepSeek 提取知识图谱...'
                    )
                    try:
                        from rag_pipeline.slides_kg import build_kg_for_slides

                        def _kg_progress(cur, total, name):
                            _ocr_tasks[task_id]['message'] = (
                                f'知识图谱提取中... [{cur}/{total}] {name}'
                            )

                        n_topics = build_kg_for_slides(
                            doc_id, result_md,
                            subject_hint=subject_hint,
                            progress_callback=_kg_progress,
                        )
                        _ocr_tasks[task_id] = {
                            'status': 'done',
                            'message': f'完成！{n} 个文本块，{n_topics} 个主题组的知识图谱已提取。',
                            'doc_id': doc_id,
                        }
                    except Exception as kg_err:
                        # KG 提取失败不影响 RAG 检索，降级处理
                        _ocr_tasks[task_id] = {
                            'status': 'done',
                            'message': (
                                f'文本块摄入成功（{n} 块）；'
                                f'知识图谱提取失败（{kg_err}），可手动重试。'
                            ),
                            'doc_id': doc_id,
                        }
                else:
                    n = ingestor.ingest_textbook(doc_id, result_md)
                    _ocr_tasks[task_id] = {
                        'status': 'done',
                        'message': f'完成！共摄入 {n} 个文本块。',
                        'doc_id': doc_id,
                    }
            except Exception as e:
                _ocr_tasks[task_id] = {
                    'status': 'error',
                    'message': str(e),
                    'doc_id': doc_id,
                }

        t = threading.Thread(
            target=_ocr_worker,
            args=(task_id, suffix, save_path, doc_id, doc_type, stem, subject_hint),
            daemon=True,
        )
        t.start()
        return jsonify({'success': True, 'task_id': task_id, 'doc_id': doc_id, 'async': True})

    return jsonify({'error': f'不支持的文件类型：{suffix}'}), 400


@rag_bp.route('/api/rag/tasks/<task_id>', methods=['GET'])
def rag_task_status(task_id):
    """轮询 OCR 后台任务状态。"""
    task = _ocr_tasks.get(task_id)
    if task is None:
        return jsonify({'error': '任务不存在'}), 404
    return jsonify(task)


@rag_bp.route('/api/rag/generate', methods=['POST'])
def rag_generate():
    """RAG 检索 → 拼提示词 → DeepSeek API → 返回题目文本。"""
    data = request.json or {}

    doc_ids = data.get('doc_ids', [])           # 限定文档范围
    chapters = data.get('chapters', [])          # 章节名（用于检索 query）
    knowledge_points = data.get('knowledge_points', [])  # 知识点/节名
    prompt_template = data.get('prompt', '')     # 含 {context} / {question_list}
    question_list = data.get('question_list', '')

    if not prompt_template:
        return jsonify({'error': '未提供提示词模板'}), 400

    # ── 构建检索 queries ──────────────────────────────────────────────────
    queries = []
    for ch in chapters[:6]:
        queries.append(ch)
    for kp in knowledge_points[:6]:
        queries.append(kp)
    if not queries:
        queries = ['课程核心知识点', '重要概念原理', '典型案例分析']

    # ── RAG 检索 context ──────────────────────────────────────────────────
    context = ''
    chunks_used = 0
    try:
        retriever = _get_retriever()
        seen_ids: set = set()
        chunks = []

        for query in queries[:12]:
            try:
                results = retriever.search(
                    query,
                    top_n=6,
                    doc_ids=doc_ids if doc_ids else None,
                )
                for r in results:
                    if r['chunk_id'] not in seen_ids:
                        chunks.append(r)
                        seen_ids.add(r['chunk_id'])
            except Exception:
                continue

        context_parts = []
        total_chars = 0
        for chunk in chunks:
            text = chunk.get('text', '')
            if total_chars + len(text) > 6000:
                break
            context_parts.append(text)
            total_chars += len(text)

        context = '\n\n---\n\n'.join(context_parts) if context_parts else ''
        chunks_used = len(context_parts)
    except Exception:
        # RAG 组件未就绪时退化为无上下文模式
        context = ''
        chunks_used = 0

    if not context:
        context = '（知识库暂无相关内容，请根据题目配置和科目知识直接出题）'

    # ── 替换提示词占位符 ──────────────────────────────────────────────────
    final_prompt = (
        prompt_template
        .replace('{context}', context)
        .replace('{question_list}', question_list)
    )

    # ── 调用 DeepSeek API ─────────────────────────────────────────────────
    try:
        _add_to_path()
        # 用 dotenv_values() 直接从文件读取，不受系统/进程环境变量干扰
        from dotenv import dotenv_values
        env_file = _project_root() / '.env'
        env_vals = dotenv_values(env_file) if env_file.exists() else {}
        api_key = (
            env_vals.get('DEEPSEEK_API_KEY')
            or env_vals.get('DEEPSEEK_TOKEN')
            or os.environ.get('DEEPSEEK_API_KEY', '')
            or os.environ.get('DEEPSEEK_TOKEN', '')
        )
        if not api_key:
            return jsonify({
                'error': '未设置 DEEPSEEK_API_KEY，请在项目根目录 .env 文件中添加：\nDEEPSEEK_API_KEY=your_key_here'
            }), 400

        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url='https://api.deepseek.com')

        response = client.chat.completions.create(
            model='deepseek-chat',
            messages=[
                {
                    'role': 'system',
                    'content': '你是一位专业的教学内容设计专家，擅长根据课程材料设计题库。',
                },
                {'role': 'user', 'content': final_prompt},
            ],
            max_tokens=8000,
            temperature=0.7,
        )

        content = response.choices[0].message.content
        return jsonify({
            'success': True,
            'content': content,
            'stats': {
                'chunks_used': chunks_used,
                'context_chars': len(context),
            },
        })
    except Exception as e:
        return jsonify({'error': f'DeepSeek API 调用失败：{str(e)}'}), 500


@rag_bp.route('/api/rag/config', methods=['GET'])
def rag_get_config():
    """读取 .env 中的 API 配置（仅返回 RAG 相关字段）。"""
    vals = _read_env_vars(
        'PADDLEOCR_API_URL', 'PADDLEOCR_TOKEN',
        'DEEPSEEK_API_KEY', 'DEEPSEEK_TOKEN',
    )
    return jsonify({
        'paddleocr_url':   vals.get('PADDLEOCR_API_URL', ''),
        'paddleocr_token': vals.get('PADDLEOCR_TOKEN', ''),
        'deepseek_api_key': (
            vals.get('DEEPSEEK_API_KEY')
            or vals.get('DEEPSEEK_TOKEN')
            or ''
        ),
    })


@rag_bp.route('/api/rag/config', methods=['PUT'])
def rag_set_config():
    """将前端提交的 API 配置写入 .env 文件。"""
    data = request.json or {}
    updates: dict = {}
    if 'paddleocr_url' in data:
        updates['PADDLEOCR_API_URL'] = data['paddleocr_url'].strip()
    if 'paddleocr_token' in data:
        updates['PADDLEOCR_TOKEN'] = data['paddleocr_token'].strip()
    if 'deepseek_api_key' in data:
        updates['DEEPSEEK_API_KEY'] = data['deepseek_api_key'].strip()
        # 同时清除旧版 key 名（避免两个同时生效时歧义）
        updates.setdefault('DEEPSEEK_TOKEN', '')
    if not updates:
        return jsonify({'error': '未提供任何配置'}), 400
    try:
        _write_env_vars(updates)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════════
# DeepSeek 直出模式 (DS Mode)
# 不依赖向量库和嵌入模型，纯 DeepSeek 知识提取 + 出题
# ═══════════════════════════════════════════════════════════════════════════════


def _do_ocr_then_parse(task_id: str, suffix: str, save_path, doc_id: str,
                       filename: str, subject: str, stem: str,
                       resume: bool = True) -> None:
    """模块级 OCR 后台工作函数（支持断点续传）。

    Args:
        task_id:  任务 ID，用于更新 _ds_tasks 状态。
        suffix:   文件后缀（.pdf / .pptx / .ppt）。
        save_path: 已保存的源文件路径。
        doc_id:   文档 ID。
        filename: 原始文件名。
        subject:  学科名称。
        stem:     文件名主干（用于构造 ocr_output 目录）。
        resume:   是否从检查点续传（默认 True）。
    """
    import datetime
    import json as _json
    from pathlib import Path as _Path

    save_path = _Path(save_path)
    ocr_output = _UPLOAD_DIR / (stem + '_ocr')

    def _progress(done, total, msg):
        _ds_tasks[task_id].update({
            'message': msg,
            'done_chunks': done,
            'total_chunks': total,
        })

    try:
        _add_to_path()
        ocr_cfg = _read_env_vars('PADDLEOCR_API_URL', 'PADDLEOCR_TOKEN', 'PADDLEOCR_TIMEOUT')
        ocr_url     = ocr_cfg.get('PADDLEOCR_API_URL', '')
        ocr_token   = ocr_cfg.get('PADDLEOCR_TOKEN', '')
        ocr_timeout = int(ocr_cfg.get('PADDLEOCR_TIMEOUT') or 120)

        if suffix == '.pdf':
            from pptx_ocr.pipeline import process_pdf
            result_md = process_pdf(
                save_path, output_dir=ocr_output,
                api_url=ocr_url, api_token=ocr_token, api_timeout=ocr_timeout,
                resume=resume, progress_callback=_progress,
            )
        else:
            from pptx_ocr.pipeline import process_pptx
            result_md = process_pptx(
                save_path, output_dir=ocr_output,
                api_url=ocr_url, api_token=ocr_token, api_timeout=ocr_timeout,
            )

        _ds_tasks[task_id]['message'] = 'OCR 完成，正在分析文档结构...'

        # ── 1. 读取 OCR 结果并清理页码标记 ────────────────────────────────────
        result_text = result_md.read_text(encoding='utf-8', errors='replace')
        cleaned_text, had_page_markers = _clean_ocr_md(result_text)

        # ── 2. 若发现页码标记，调用 DeepSeek 判断章节标题级别 ─────────────────
        chapter_level = 1
        if had_page_markers:
            _ds_tasks[task_id]['message'] = (
                '检测到页码标记，已自动清理；正在识别章节层级...'
            )
            ds_cfg = _read_env_vars('DEEPSEEK_API_KEY', 'DEEPSEEK_TOKEN')
            ds_key = (
                ds_cfg.get('DEEPSEEK_API_KEY')
                or ds_cfg.get('DEEPSEEK_TOKEN')
                or os.environ.get('DEEPSEEK_API_KEY', '')
                or os.environ.get('DEEPSEEK_TOKEN', '')
            )
            if ds_key:
                try:
                    from openai import OpenAI as _OAI
                    _ds_client = _OAI(api_key=ds_key, base_url='https://api.deepseek.com')
                    chapter_level = _detect_chapter_level(cleaned_text, subject, _ds_client)
                    _ds_tasks[task_id]['message'] = (
                        f'已识别章节层级为 H{chapter_level}（{"#" * chapter_level}），正在分段...'
                    )
                except Exception:
                    chapter_level = 1  # 检测失败时安全回退

        # ── 3. 按检测到的级别切分章节 ─────────────────────────────────────────
        chapters = _parse_md_to_chapters(cleaned_text, chapter_level)
        if not chapters:
            _ds_tasks[task_id] = {
                'status': 'error',
                'message': '未能从 OCR 结果中解析出章节内容，请确认文档含有标题结构',
                'doc_id': doc_id, 'type': 'ocr',
            }
            return

        with _ds_db_conn() as conn:
            conn.execute(
                "INSERT OR REPLACE INTO ds_docs "
                "(doc_id, filename, subject, status, error_msg, created_at) "
                "VALUES (?,?,?,'uploaded','',?)",
                (doc_id, filename, subject, datetime.datetime.now().isoformat()),
            )
            conn.execute("DELETE FROM ds_chapters WHERE doc_id=?", (doc_id,))
            conn.execute("DELETE FROM ds_kps WHERE doc_id=?", (doc_id,))
            for ch in chapters:
                conn.execute(
                    "INSERT INTO ds_chapters "
                    "(doc_id, chapter_num, chapter_name, raw_text) VALUES (?,?,?,?)",
                    (doc_id, ch['num'], ch['name'], ch['text']),
                )
            conn.commit()

        _ds_tasks[task_id] = {
            'status': 'done',
            'message': f'完成！已解析 {len(chapters)} 个章节，可点击「提取」按钮提取知识点',
            'doc_id': doc_id,
            'chapter_count': len(chapters),
            'type': 'ocr',
        }

    except Exception as e:
        err_str = str(e)
        # 判断是否为网络超时 / 连接错误（可续传）
        _lower = err_str.lower()
        is_timeout = (
            'timed out' in _lower
            or 'timeout' in _lower
            or 'connectionerror' in _lower
            or 'connection' in _lower and 'error' in _lower
            or 'read timeout' in _lower
        )

        if is_timeout:
            # 从检查点文件读取已完成块数
            ck_file = ocr_output / 'checkpoint.json'
            done_count = _ds_tasks[task_id].get('done_chunks', 0)
            total_count = _ds_tasks[task_id].get('total_chunks', 0)
            if ck_file.exists():
                try:
                    ck = _json.loads(ck_file.read_text(encoding='utf-8'))
                    done_count = len(ck.get('done', []))
                    total_count = ck.get('total', total_count)
                except Exception:
                    pass

            _ds_tasks[task_id] = {
                'status': 'timeout',
                'message': (
                    f'OCR 超时（已完成 {done_count}/{total_count} 块）。'
                    f'已保存进度，点击「重试」按钮继续'
                ),
                'doc_id': doc_id,
                'type': 'ocr',
                'done_chunks': done_count,
                'total_chunks': total_count,
                # 重试所需参数
                '_retry': {
                    'suffix': suffix,
                    'save_path': str(save_path),
                    'filename': filename,
                    'subject': subject,
                    'stem': stem,
                },
            }
        else:
            _ds_tasks[task_id] = {
                'status': 'error',
                'message': err_str,
                'doc_id': doc_id, 'type': 'ocr',
            }


def _ds_db_path() -> Path:
    """DS 模式独立 SQLite 数据库，不影响 RAG 的 kg.db。"""
    return _project_root() / "ds_knowledge.db"


def _ds_db_conn():
    conn = sqlite3.connect(str(_ds_db_path()))
    conn.row_factory = sqlite3.Row
    return conn


def _init_ds_db():
    with _ds_db_conn() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS ds_docs (
                doc_id     TEXT PRIMARY KEY,
                filename   TEXT,
                subject    TEXT,
                status     TEXT DEFAULT 'uploaded',
                error_msg  TEXT,
                created_at TEXT
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS ds_chapters (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id       TEXT,
                chapter_num  INTEGER,
                chapter_name TEXT,
                raw_text     TEXT
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS ds_kps (
                id             INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id         TEXT,
                chapter_name   TEXT,
                chapter_num    INTEGER,
                kp_name        TEXT,
                kp_content     TEXT,
                relations_json TEXT DEFAULT '[]'
            )
        """)
        conn.commit()


# _ds_tasks 已在模块顶部声明，此处不重复定义


def _clean_ocr_md(text: str) -> tuple:
    """清理 OCR pipeline 生成的 result.md 中的人工产物。

    处理步骤：
    1. 去掉 pipeline 生成的文件头（# 文件名 / > Source / > Generated / > Chunks / ---）
    2. 去掉页码标记行（## Page N、## 第N页 等）
    3. 去掉多余的 --- 分隔线（OCR 分页符）
    4. 合并连续空行（>2 行 → 1 空行）

    返回 (cleaned_text: str, had_page_markers: bool)
    """
    import re

    # 步骤 1：识别并去掉 pipeline 文件头
    # 文件头特征：开头有 "# 文件名" 行，随后有 "> Source:" / "> Generated:" 等元信息行，
    # 以第一个单独的 "---" 行结束。
    header_match = re.search(r'^---[ \t]*\n', text, re.MULTILINE)
    if header_match:
        before_sep = text[:header_match.start()]
        # 判断 --- 之前确实含有 pipeline 元信息（> Source / > Generated / > Chunks）
        if re.search(r'^>\s*(?:Source|Generated|Chunks)', before_sep, re.MULTILINE):
            text = text[header_match.end():]

    # 步骤 2：去掉页码标记行
    # 匹配：## Page 1、## Page 123、## 第1页、## 第 12 面 等，整行独立
    page_pattern = re.compile(
        r'^##\s+(?:Page\s+\d+|第\s*\d+\s*[页面])\s*$',
        re.MULTILINE | re.IGNORECASE,
    )
    cleaned, n_markers = page_pattern.subn('', text)

    # 步骤 3：去掉多余的 --- 分隔线（OCR 各 chunk 之间的分页符）
    # 只在页码标记被清理后执行，避免误删文档原有 --- 分隔
    if n_markers > 0:
        cleaned = re.sub(r'\n[ \t]*---[ \t]*\n', '\n', cleaned)

    # 步骤 4：合并 3 行及以上连续空行为 1 行
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)

    return cleaned.strip(), n_markers > 0


def _detect_chapter_level(text_sample: str, subject: str, ds_client) -> int:
    """调用 DeepSeek 判断 Markdown 文档中哪个标题级别代表"章"。

    参数：
        text_sample: 文档前 2000 字左右的样本（已清理页码标记）
        subject:     学科名称（辅助判断）
        ds_client:   已初始化的 OpenAI-compatible 客户端

    返回：1、2 或 3（对应 #、## 或 ###），失败时默认返回 1。
    """
    import json as _json

    prompt = (
        '请分析以下教材文档（OCR 识别结果）的部分内容，判断哪个 Markdown 标题级别代表"章"。\n\n'
        f'【学科/科目】：{subject or "（未指定）"}\n\n'
        '【说明】\n'
        '- 章是文档最高层的内容组织单元，通常含"第X章""Chapter X"或同等表述\n'
        '- 标题级别：# = 1级，## = 2级，### = 3级\n'
        '- 若文档无明显章级别，选择最接近章粒度的标题级别\n\n'
        f'【文档样本】：\n{text_sample[:2000]}\n\n'
        '请直接输出 JSON（不加代码块标记或其他文字）：\n'
        '{"chapter_level": 1, "reason": "简要说明（20字以内）"}'
    )
    try:
        resp = ds_client.chat.completions.create(
            model='deepseek-chat',
            messages=[
                {'role': 'system', 'content': '你是一位文档结构分析专家，善于识别教材的章节层级。'},
                {'role': 'user', 'content': prompt},
            ],
            max_tokens=120,
            temperature=0,
        )
        raw = resp.choices[0].message.content.strip()
        if raw.startswith('```'):
            raw = '\n'.join(raw.split('\n')[1:])
        if raw.endswith('```'):
            raw = '\n'.join(raw.split('\n')[:-1])
        parsed = _json.loads(raw.strip())
        level = int(parsed.get('chapter_level', 1))
        if level in (1, 2, 3):
            return level
    except Exception:
        pass
    return 1  # 默认值


def _parse_md_to_chapters(text: str, heading_level: int = 1) -> list:
    """将 Markdown/TXT 文本按指定标题级别切分为章节列表。

    参数：
        text:          文档文本（应已用 _clean_ocr_md 清理）
        heading_level: 章标题级别，1=#，2=##，3=###（默认 1）

    返回 [{'num': int, 'name': str, 'text': str}, ...]，过滤空章节。
    """
    import re
    # 精确匹配 heading_level 个 #，后面不能再跟 #（避免 ## 被 # 规则误匹配）
    pattern = re.compile(
        r'^#{' + str(heading_level) + r'}(?!#)\s+(.+)',
        re.MULTILINE,
    )
    chapters: list = []
    current_title = '（前言/概述）'
    current_num = 0
    current_lines: list = []
    for line in text.splitlines():
        m = pattern.match(line)
        if m:
            if current_lines:
                chapters.append({
                    'num': current_num,
                    'name': current_title,
                    'text': '\n'.join(current_lines).strip(),
                })
            current_num += 1
            current_title = m.group(1).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_lines:
        chapters.append({
            'num': current_num,
            'name': current_title,
            'text': '\n'.join(current_lines).strip(),
        })
    return [c for c in chapters if c['text'].strip()]


def _parse_file_to_chapters(filepath: Path, suffix: str,
                            heading_level: int = 1) -> list:
    """从文件解析章节列表 [{num, name, text}]。

    对 .md/.txt 文件自动应用 _clean_ocr_md 清理（对非 OCR 文件无副作用）。
    若需智能章节级别检测，请在调用前先用 _detect_chapter_level 获取 heading_level。
    """
    if suffix in ('.md', '.txt'):
        text = filepath.read_text(encoding='utf-8', errors='replace')
        cleaned, _ = _clean_ocr_md(text)
        return _parse_md_to_chapters(cleaned, heading_level)
    if suffix == '.docx':
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(filepath))
            md_lines: list = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if not t:
                    continue
                style = para.style.name if para.style else ''
                if 'Heading 1' in style or 'heading 1' in style.lower():
                    md_lines.append(f'# {t}')
                elif 'Heading 2' in style or 'heading 2' in style.lower():
                    md_lines.append(f'## {t}')
                else:
                    md_lines.append(t)
            combined = '\n'.join(md_lines)
            cleaned, _ = _clean_ocr_md(combined)
            return _parse_md_to_chapters(cleaned, heading_level)
        except Exception:
            return []
    return []


_DS_EXTRACT_SYSTEM = '你是一位专业的教材分析专家，擅长从教材中提取结构化知识图谱。'

_DS_EXTRACT_PROMPT = """请分析以下教材章节内容，提取该章节的核心知识点及其关系，构建知识图谱。

【学科/科目】：{subject}
【章节名称】：{chapter_name}

【章节原文】：
{chapter_text}

---
**任务要求：**
1. 提取 3-8 个核心知识点
2. 每个知识点的 content 字段应尽量引用原文关键表述，包含：定义/概念、特征/特点、分类/类型、示例/案例等
3. 仅标注同一章节内知识点之间的关系
4. 直接输出纯 JSON，不要添加代码块标记（```）或任何其他文字

**输出格式（严格遵守，直接输出 JSON）：**
{{
  "chapter": "章节名称",
  "knowledge_points": [
    {{
      "name": "知识点名称（5-20 字）",
      "content": "详细说明（引用原文，200-500 字，包含定义、特征、分类、示例）",
      "relations": [
        {{"type": "从属", "target": "上位知识点名称"}},
        {{"type": "比较", "target": "对比知识点名称"}},
        {{"type": "并列", "target": "同级知识点名称"}}
      ]
    }}
  ]
}}

关系类型（可组合使用，无关系则填 []）：从属、比较、并列、交叉、前提"""


# ── DS 端点 ───────────────────────────────────────────────────────────────────

@rag_bp.route('/api/rag/ds-upload', methods=['POST'])
def ds_upload():
    """上传文件并解析章节结构（DS 直出模式，不摄入向量库）。"""
    _ensure_upload_dir()
    _init_ds_db()

    if 'file' not in request.files:
        return jsonify({'error': '未提供文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    subject = request.form.get('subject', '')
    filename = file.filename
    suffix = Path(filename).suffix.lower()
    stem = Path(filename).stem
    doc_id = stem[:40]

    if suffix not in ('.md', '.txt', '.docx', '.pptx', '.ppt', '.pdf'):
        return jsonify({'error': f'不支持的文件格式：{suffix}'}), 400

    save_path = _UPLOAD_DIR / filename
    file.save(str(save_path))

    # ── PDF / PPTX：先走 PaddleOCR 异步流程，OCR 完成后再解析章节 ──────────────
    if suffix in ('.pptx', '.ppt', '.pdf'):
        task_id = str(uuid.uuid4())[:8]
        _ds_tasks[task_id] = {
            'status': 'running',
            'message': 'OCR 处理中，请稍候...',
            'doc_id': doc_id,
            'type': 'ocr',
        }

        t = threading.Thread(
            target=_do_ocr_then_parse,
            args=(task_id, suffix, save_path, doc_id, filename, subject, stem),
            daemon=True,
        )
        t.start()
        return jsonify({'success': True, 'task_id': task_id, 'doc_id': doc_id, 'async': True})

    # ── 文本格式：同步解析章节 ───────────────────────────────────────────────────
    try:
        # 读取原始文本并清理 OCR 页码标记（对非 OCR 文件无副作用）
        if suffix in ('.md', '.txt'):
            raw_text = save_path.read_text(encoding='utf-8', errors='replace')
            cleaned_text, had_markers = _clean_ocr_md(raw_text)
        else:
            cleaned_text, had_markers = None, False

        # 若发现页码标记（上传了 OCR 结果文件），自动检测章节级别
        chapter_level = 1
        if had_markers:
            ds_cfg = _read_env_vars('DEEPSEEK_API_KEY', 'DEEPSEEK_TOKEN')
            ds_key = (
                ds_cfg.get('DEEPSEEK_API_KEY')
                or ds_cfg.get('DEEPSEEK_TOKEN')
                or os.environ.get('DEEPSEEK_API_KEY', '')
                or os.environ.get('DEEPSEEK_TOKEN', '')
            )
            if ds_key:
                try:
                    from openai import OpenAI as _OAI
                    _ds_client = _OAI(api_key=ds_key, base_url='https://api.deepseek.com')
                    chapter_level = _detect_chapter_level(cleaned_text, subject, _ds_client)
                except Exception:
                    chapter_level = 1

        # 按检测级别解析章节
        if cleaned_text is not None:
            chapters = _parse_md_to_chapters(cleaned_text, chapter_level)
        else:
            chapters = _parse_file_to_chapters(save_path, suffix, chapter_level)

        if not chapters:
            return jsonify({'error': '未能从文件中解析出章节内容，请确认文件含有标题结构'}), 400

        import datetime
        with _ds_db_conn() as conn:
            conn.execute(
                "INSERT OR REPLACE INTO ds_docs "
                "(doc_id, filename, subject, status, error_msg, created_at) "
                "VALUES (?,?,?,'uploaded','',?)",
                (doc_id, filename, subject, datetime.datetime.now().isoformat()),
            )
            conn.execute("DELETE FROM ds_chapters WHERE doc_id=?", (doc_id,))
            conn.execute("DELETE FROM ds_kps WHERE doc_id=?", (doc_id,))
            for ch in chapters:
                conn.execute(
                    "INSERT INTO ds_chapters (doc_id, chapter_num, chapter_name, raw_text) VALUES (?,?,?,?)",
                    (doc_id, ch['num'], ch['name'], ch['text']),
                )
            conn.commit()

        return jsonify({
            'success': True,
            'doc_id': doc_id,
            'chapter_count': len(chapters),
            'chapters': [{'num': c['num'], 'name': c['name']} for c in chapters],
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@rag_bp.route('/api/rag/ds-extract/<path:doc_id>', methods=['POST'])
def ds_extract(doc_id):
    """异步：逐章调用 DeepSeek 提取知识点，构建知识图谱。

    resume=True 时跳过已提取的章节，实现断点续传。
    """
    _init_ds_db()
    data = request.json or {}
    subject = data.get('subject', '')
    resume = bool(data.get('resume', False))

    with _ds_db_conn() as conn:
        conn.execute("UPDATE ds_docs SET subject=? WHERE doc_id=?", (subject, doc_id))
        conn.execute(
            "UPDATE ds_docs SET status='extracting', error_msg='' WHERE doc_id=?",
            (doc_id,),
        )
        if not resume:
            # 全新提取：清空已有知识点
            conn.execute("DELETE FROM ds_kps WHERE doc_id=?", (doc_id,))
        conn.commit()
        all_chapters = conn.execute(
            "SELECT chapter_num, chapter_name, raw_text FROM ds_chapters "
            "WHERE doc_id=? ORDER BY chapter_num",
            (doc_id,),
        ).fetchall()
        if resume:
            # 续传：跳过已有 KP 的章节
            extracted_nums = {
                r[0] for r in conn.execute(
                    "SELECT DISTINCT chapter_num FROM ds_kps WHERE doc_id=?",
                    (doc_id,),
                ).fetchall()
            }
            chapters = [ch for ch in all_chapters if ch['chapter_num'] not in extracted_nums]
        else:
            chapters = list(all_chapters)

    if not all_chapters:
        return jsonify({'error': '文档未上传或章节解析失败，请先上传文档'}), 404

    if not chapters:
        # 续传但所有章节都已提取完毕，直接标记 done
        with _ds_db_conn() as conn:
            kp_count = conn.execute(
                "SELECT COUNT(*) FROM ds_kps WHERE doc_id=?", (doc_id,)
            ).fetchone()[0]
            conn.execute("UPDATE ds_docs SET status='done' WHERE doc_id=?", (doc_id,))
            conn.commit()
        return jsonify({
            'success': True, 'task_id': None,
            'total': 0, 'already_done': True, 'kp_count': kp_count,
        })

    total_chapters = len(all_chapters)
    task_id = str(uuid.uuid4())[:8]
    _ds_tasks[task_id] = {
        'status': 'running',
        'message': '准备开始...',
        'doc_id': doc_id,
        'progress': 0,
        'total': len(chapters),
        'total_chapters': total_chapters,
        'pause_requested': False,
    }

    def _worker(task_id, doc_id, subject, chapters):
        import json as _json
        try:
            _add_to_path()
            from dotenv import dotenv_values
            env_file = _project_root() / '.env'
            env_vals = dotenv_values(env_file) if env_file.exists() else {}
            api_key = (
                env_vals.get('DEEPSEEK_API_KEY')
                or env_vals.get('DEEPSEEK_TOKEN')
                or os.environ.get('DEEPSEEK_API_KEY', '')
                or os.environ.get('DEEPSEEK_TOKEN', '')
            )
            if not api_key:
                with _ds_db_conn() as conn:
                    conn.execute(
                        "UPDATE ds_docs SET status='error', error_msg=? WHERE doc_id=?",
                        ('未设置 DEEPSEEK_API_KEY', doc_id),
                    )
                    conn.commit()
                _ds_tasks[task_id] = {
                    'status': 'error',
                    'message': '未设置 DEEPSEEK_API_KEY，请在 API 配置中填写',
                    'doc_id': doc_id,
                }
                return

            from openai import OpenAI
            client = OpenAI(api_key=api_key, base_url='https://api.deepseek.com')
            total = len(chapters)

            for i, ch in enumerate(chapters):
                # 检查暂停请求
                if _ds_tasks.get(task_id, {}).get('pause_requested', False):
                    with _ds_db_conn() as conn:
                        kp_count = conn.execute(
                            "SELECT COUNT(*) FROM ds_kps WHERE doc_id=?", (doc_id,)
                        ).fetchone()[0]
                        ch_done = conn.execute(
                            "SELECT COUNT(DISTINCT chapter_num) FROM ds_kps WHERE doc_id=?",
                            (doc_id,),
                        ).fetchone()[0]
                        conn.execute(
                            "UPDATE ds_docs SET status='paused' WHERE doc_id=?", (doc_id,)
                        )
                        conn.commit()
                    _ds_tasks[task_id] = {
                        'status': 'paused',
                        'message': f'已暂停。已完成 {ch_done}/{_ds_tasks[task_id].get("total_chapters", total)} 章，共 {kp_count} 个知识点',
                        'doc_id': doc_id,
                        'progress': i,
                        'total': total,
                        'total_chapters': _ds_tasks[task_id].get('total_chapters', total),
                        'kp_count': kp_count,
                    }
                    return

                _ds_tasks[task_id].update({
                    'message': f'正在提取第 {i + 1}/{total} 章：{ch["chapter_name"]}',
                    'progress': i,
                    'total': total,
                })
                ch_text = ch['raw_text']
                if len(ch_text) > 8000:
                    ch_text = ch_text[:8000] + '\n\n[（内容过长，已截断）]'

                prompt = _DS_EXTRACT_PROMPT.format(
                    subject=subject or '（未指定）',
                    chapter_name=ch['chapter_name'],
                    chapter_text=ch_text,
                )
                try:
                    resp = client.chat.completions.create(
                        model='deepseek-chat',
                        messages=[
                            {'role': 'system', 'content': _DS_EXTRACT_SYSTEM},
                            {'role': 'user', 'content': prompt},
                        ],
                        max_tokens=3000,
                        temperature=0.3,
                    )
                    raw = resp.choices[0].message.content.strip()
                    # 清除可能的 Markdown 代码块标记
                    if raw.startswith('```'):
                        raw = '\n'.join(raw.split('\n')[1:])
                    if raw.endswith('```'):
                        raw = '\n'.join(raw.split('\n')[:-1])
                    raw = raw.strip()

                    parsed = _json.loads(raw)
                    kps = parsed.get('knowledge_points', [])
                    with _ds_db_conn() as conn:
                        for kp in kps:
                            conn.execute(
                                "INSERT INTO ds_kps "
                                "(doc_id, chapter_name, chapter_num, kp_name, kp_content, relations_json) "
                                "VALUES (?,?,?,?,?,?)",
                                (
                                    doc_id,
                                    ch['chapter_name'],
                                    ch['chapter_num'],
                                    kp.get('name', ''),
                                    kp.get('content', ''),
                                    _json.dumps(kp.get('relations', []), ensure_ascii=False),
                                ),
                            )
                        conn.commit()
                except Exception as ch_err:
                    # 单章失败不影响整体，记录并继续
                    _ds_tasks[task_id]['message'] = (
                        f'第 {i + 1} 章提取出错（{ch_err}），跳过，继续下一章...'
                    )

            # 全部完成
            with _ds_db_conn() as conn:
                kp_count = conn.execute(
                    "SELECT COUNT(*) FROM ds_kps WHERE doc_id=?", (doc_id,)
                ).fetchone()[0]
                conn.execute(
                    "UPDATE ds_docs SET status='done' WHERE doc_id=?", (doc_id,)
                )
                conn.commit()
            _ds_tasks[task_id] = {
                'status': 'done',
                'message': f'提取完成！共提取 {kp_count} 个知识点（{total} 章）',
                'doc_id': doc_id,
                'kp_count': kp_count,
                'total': total,
            }
        except Exception as e:
            with _ds_db_conn() as conn:
                conn.execute(
                    "UPDATE ds_docs SET status='error', error_msg=? WHERE doc_id=?",
                    (str(e), doc_id),
                )
                conn.commit()
            _ds_tasks[task_id] = {
                'status': 'error',
                'message': str(e),
                'doc_id': doc_id,
            }

    t = threading.Thread(
        target=_worker, args=(task_id, doc_id, subject, list(chapters)), daemon=True
    )
    t.start()
    return jsonify({'success': True, 'task_id': task_id, 'total': len(chapters)})


@rag_bp.route('/api/rag/ds-tasks/<task_id>', methods=['GET'])
def ds_task_status(task_id):
    """轮询 DS 知识点提取任务状态。"""
    task = _ds_tasks.get(task_id)
    if task is None:
        return jsonify({'error': '任务不存在'}), 404
    # 不把内部重试参数暴露给前端
    result = {k: v for k, v in task.items() if k != '_retry'}
    return jsonify(result)


@rag_bp.route('/api/rag/ds-tasks/<task_id>/pause', methods=['POST'])
def ds_pause_task(task_id):
    """请求暂停正在进行的知识点提取任务（将在当前章节完成后生效）。"""
    task = _ds_tasks.get(task_id)
    if task is None:
        return jsonify({'error': '任务不存在'}), 404
    if task.get('status') != 'running':
        return jsonify({'error': f'任务当前状态为 {task.get("status")}，无法暂停'}), 400
    _ds_tasks[task_id]['pause_requested'] = True
    return jsonify({'success': True, 'message': '暂停请求已发送，将在当前章节完成后暂停'})


@rag_bp.route('/api/rag/ds-retry-ocr/<task_id>', methods=['POST'])
def ds_retry_ocr(task_id):
    """断点续传：基于已有检查点重新启动 OCR 任务。"""
    task = _ds_tasks.get(task_id)
    if not task:
        return jsonify({'error': '原任务不存在'}), 404
    if task.get('status') != 'timeout':
        return jsonify({'error': '该任务不是超时状态，无法续传'}), 400

    retry_params = task.get('_retry')
    if not retry_params:
        return jsonify({'error': '缺少重试参数，请重新上传文件'}), 400

    new_task_id = str(uuid.uuid4())[:8]
    doc_id = task['doc_id']
    _ds_tasks[new_task_id] = {
        'status': 'running',
        'message': 'OCR 断点续传中，正在跳过已完成分块...',
        'doc_id': doc_id,
        'type': 'ocr',
    }

    t = threading.Thread(
        target=_do_ocr_then_parse,
        args=(
            new_task_id,
            retry_params['suffix'],
            retry_params['save_path'],
            doc_id,
            retry_params['filename'],
            retry_params['subject'],
            retry_params['stem'],
        ),
        kwargs={'resume': True},
        daemon=True,
    )
    t.start()
    return jsonify({'success': True, 'task_id': new_task_id, 'doc_id': doc_id})


@rag_bp.route('/api/rag/ds-graph', methods=['GET'])
def ds_graph():
    """返回 DS 知识图谱数据（节点 + 边），供前端 D3 力导向图可视化。

    Query params:
        doc_id  (可重复): 限定文档，不传则返回所有已完成文档
        chapter (可重复): 章节名称过滤
    """
    import json as _json
    _init_ds_db()

    doc_ids = request.args.getlist('doc_id')
    chapters_filter = request.args.getlist('chapter')

    with _ds_db_conn() as conn:
        # 若未指定 doc_id，取所有 status='done' 的文档
        if not doc_ids:
            rows_d = conn.execute(
                "SELECT doc_id FROM ds_docs WHERE status='done'"
            ).fetchall()
            doc_ids = [r['doc_id'] for r in rows_d]

        if not doc_ids:
            return jsonify({'nodes': [], 'links': [], 'chapters': []})

        ph = ','.join('?' for _ in doc_ids)
        query = (
            "SELECT id, doc_id, chapter_name, chapter_num, "
            "kp_name, kp_content, relations_json "
            f"FROM ds_kps WHERE doc_id IN ({ph})"
        )
        params: list = list(doc_ids)
        if chapters_filter:
            cp = ','.join('?' for _ in chapters_filter)
            query += f" AND chapter_name IN ({cp})"
            params.extend(chapters_filter)
        query += " ORDER BY doc_id, chapter_num, id"
        rows = conn.execute(query, params).fetchall()

    if not rows:
        return jsonify({'nodes': [], 'links': [], 'chapters': []})

    nodes: list = []
    # name_to_id: (doc_id, kp_name) -> node_id，同时保留 kp_name -> node_id 作为跨文档回退
    name_to_id: dict = {}
    chapters_map: dict = {}  # chapter_name -> chapter_num

    for row in rows:
        nid = f"kp_{row['id']}"
        name_to_id[(row['doc_id'], row['kp_name'])] = nid
        name_to_id.setdefault(row['kp_name'], nid)  # 跨文档回退（先写先得）
        nodes.append({
            'id': nid,
            'name': row['kp_name'],
            'chapter': row['chapter_name'],
            'chapter_num': row['chapter_num'],
            'doc_id': row['doc_id'],
            'content': row['kp_content'],
        })
        chapters_map.setdefault(row['chapter_name'], row['chapter_num'])

    links: list = []
    for row in rows:
        try:
            rels = _json.loads(row['relations_json'] or '[]')
            src = f"kp_{row['id']}"
            for rel in rels:
                tgt_name = (rel.get('target') or '').strip()
                rel_type = rel.get('type', '')
                tgt = (
                    name_to_id.get((row['doc_id'], tgt_name))
                    or name_to_id.get(tgt_name)
                )
                if tgt and tgt != src:
                    links.append({'source': src, 'target': tgt, 'type': rel_type})
        except Exception:
            pass

    chapters = sorted(
        [{'name': k, 'num': v} for k, v in chapters_map.items()],
        key=lambda c: c['num'],
    )
    return jsonify({
        'nodes': nodes,
        'links': links,
        'chapters': chapters,
        'total_nodes': len(nodes),
        'total_links': len(links),
    })


@rag_bp.route('/api/rag/ds-docs', methods=['GET'])
def list_ds_docs():
    """列出所有 DS 模式文档及其状态。"""
    _init_ds_db()
    with _ds_db_conn() as conn:
        rows = conn.execute(
            "SELECT doc_id, filename, subject, status, error_msg FROM ds_docs "
            "ORDER BY created_at DESC"
        ).fetchall()
    result = []
    for r in rows:
        with _ds_db_conn() as conn:
            kp_count = conn.execute(
                "SELECT COUNT(*) FROM ds_kps WHERE doc_id=?", (r['doc_id'],)
            ).fetchone()[0]
            ch_count = conn.execute(
                "SELECT COUNT(*) FROM ds_chapters WHERE doc_id=?", (r['doc_id'],)
            ).fetchone()[0]
            ch_with_kps = conn.execute(
                "SELECT COUNT(DISTINCT chapter_num) FROM ds_kps WHERE doc_id=?",
                (r['doc_id'],),
            ).fetchone()[0]
        result.append({
            'doc_id': r['doc_id'],
            'filename': r['filename'] or '',
            'subject': r['subject'] or '',
            'status': r['status'],
            'error_msg': r['error_msg'] or '',
            'kp_count': kp_count,
            'chapter_count': ch_count,
            'ch_with_kps': ch_with_kps,
        })
    return jsonify(result)


@rag_bp.route('/api/rag/ds-docs/<path:doc_id>/kps', methods=['GET'])
def ds_doc_kps(doc_id):
    """获取文档的章节列表和知识点列表（用于前端三级联动筛选）。"""
    _init_ds_db()
    with _ds_db_conn() as conn:
        chapters = conn.execute(
            "SELECT DISTINCT chapter_num, chapter_name FROM ds_chapters "
            "WHERE doc_id=? ORDER BY chapter_num",
            (doc_id,),
        ).fetchall()
        kps = conn.execute(
            "SELECT id, chapter_name, chapter_num, kp_name, kp_content, relations_json "
            "FROM ds_kps WHERE doc_id=? ORDER BY chapter_num, id",
            (doc_id,),
        ).fetchall()
    return jsonify({
        'chapters': [{'num': c['chapter_num'], 'name': c['chapter_name']} for c in chapters],
        'kps': [
            {
                'id': k['id'],
                'chapter_name': k['chapter_name'],
                'chapter_num': k['chapter_num'],
                'name': k['kp_name'],
                'content': k['kp_content'],
                'relations': k['relations_json'],
            }
            for k in kps
        ],
    })


@rag_bp.route('/api/rag/ds-docs/<path:doc_id>', methods=['DELETE'])
def ds_delete_doc(doc_id):
    """删除 DS 文档及其所有章节和知识点数据。"""
    _init_ds_db()
    with _ds_db_conn() as conn:
        conn.execute("DELETE FROM ds_docs WHERE doc_id=?", (doc_id,))
        conn.execute("DELETE FROM ds_chapters WHERE doc_id=?", (doc_id,))
        conn.execute("DELETE FROM ds_kps WHERE doc_id=?", (doc_id,))
        conn.commit()
    return jsonify({'success': True})


@rag_bp.route('/api/rag/ds-generate', methods=['POST'])
def ds_generate():
    """基于 DS 提取的知识图谱生成题目（DeepSeek 直出，无向量检索）。"""
    import json as _json
    _init_ds_db()
    data = request.json or {}
    doc_ids = data.get('doc_ids', [])
    chapters = data.get('chapters', [])
    kp_names = data.get('kp_names', [])
    prompt_template = data.get('prompt', '')
    question_list = data.get('question_list', '')

    if not prompt_template:
        return jsonify({'error': '未提供提示词模板'}), 400

    # ── 从 DS 数据库读取相关知识点 ────────────────────────────────────────────
    kps_data = []
    try:
        with _ds_db_conn() as conn:
            for doc_id in (doc_ids or []):
                query = (
                    "SELECT chapter_name, chapter_num, kp_name, kp_content, relations_json "
                    "FROM ds_kps WHERE doc_id=?"
                )
                params: list = [doc_id]
                if chapters:
                    placeholders = ','.join(['?' for _ in chapters])
                    query += f" AND chapter_name IN ({placeholders})"
                    params.extend(chapters)
                if kp_names:
                    placeholders = ','.join(['?' for _ in kp_names])
                    query += f" AND kp_name IN ({placeholders})"
                    params.extend(kp_names)
                query += " ORDER BY chapter_num, id"
                rows = conn.execute(query, params).fetchall()
                kps_data.extend(rows)
    except Exception:
        kps_data = []

    # ── 构建知识图谱 context ──────────────────────────────────────────────────
    if kps_data:
        context_parts: list = []
        current_chapter = None
        total_chars = 0
        for kp in kps_data:
            if total_chars > 7000:
                context_parts.append('\n（已达上下文长度上限，后续知识点省略）')
                break
            if kp['chapter_name'] != current_chapter:
                current_chapter = kp['chapter_name']
                context_parts.append(f'\n## {current_chapter}\n')
            context_parts.append(f'### 知识点：{kp["kp_name"]}')
            context_parts.append(kp['kp_content'])
            try:
                rels = _json.loads(kp['relations_json'] or '[]')
                if rels:
                    rel_str = '；'.join(
                        f"{r.get('type', '')} → {r.get('target', '')}"
                        for r in rels
                        if r.get('target')
                    )
                    if rel_str:
                        context_parts.append(f'关联关系：{rel_str}')
            except Exception:
                pass
            context_parts.append('')
            total_chars += len(kp['kp_content'])
        context = '\n'.join(context_parts)
    else:
        context = '（未选择知识点或知识点提取尚未完成，请先上传文档并完成知识点提取）'

    # ── 替换占位符 ────────────────────────────────────────────────────────────
    final_prompt = (
        prompt_template
        .replace('{context}', context)
        .replace('{question_list}', question_list)
    )

    # ── 调用 DeepSeek API ─────────────────────────────────────────────────────
    try:
        _add_to_path()
        from dotenv import dotenv_values
        env_file = _project_root() / '.env'
        env_vals = dotenv_values(env_file) if env_file.exists() else {}
        api_key = (
            env_vals.get('DEEPSEEK_API_KEY')
            or env_vals.get('DEEPSEEK_TOKEN')
            or os.environ.get('DEEPSEEK_API_KEY', '')
            or os.environ.get('DEEPSEEK_TOKEN', '')
        )
        if not api_key:
            return jsonify({
                'error': '未设置 DEEPSEEK_API_KEY，请在 API 配置中填写'
            }), 400

        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url='https://api.deepseek.com')
        response = client.chat.completions.create(
            model='deepseek-chat',
            messages=[
                {
                    'role': 'system',
                    'content': '你是一位专业的教学内容设计专家，擅长根据课程知识图谱设计题库。',
                },
                {'role': 'user', 'content': final_prompt},
            ],
            max_tokens=8000,
            temperature=0.7,
        )
        content = response.choices[0].message.content
        return jsonify({
            'success': True,
            'content': content,
            'stats': {
                'kps_used': len(kps_data),
                'context_chars': len(context),
            },
        })
    except Exception as e:
        return jsonify({'error': f'DeepSeek API 调用失败：{str(e)}'}), 500
