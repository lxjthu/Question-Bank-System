"""
Word文档导出器
用于将试卷导出为真正的Word文档格式，并提供预览和样式设置功能
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser, font
import json
import os


class WordExporter:
    def __init__(self):
        self.document = None
        self.exam_data = None
        
        # 默认样式设置
        self.styles = {
            'header_font': 'SimSun',
            'header_size': 16,
            'header_bold': True,
            'question_font': 'SimSun',
            'question_size': 12,
            'question_bold': False,
            'option_font': 'SimSun',
            'option_size': 11,
            'option_bold': False,
            'answer_font': 'SimSun',
            'answer_size': 10,
            'answer_bold': False,
            'header_color': '#000000',
            'question_color': '#000000',
            'option_color': '#000000',
            'answer_color': '#FF0000',  # 红色答案
            'page_margins': {
                'top': 1.0,
                'bottom': 1.0,
                'left': 1.0,
                'right': 1.0
            }
        }
    
    def set_exam_data(self, exam_data):
        """设置试卷数据"""
        self.exam_data = exam_data
    
    def create_document(self):
        """创建Word文档"""
        self.document = Document()
        
        # 设置页面边距
        section = self.document.sections[0]
        section.top_margin = Inches(self.styles['page_margins']['top'])
        section.bottom_margin = Inches(self.styles['page_margins']['bottom'])
        section.left_margin = Inches(self.styles['page_margins']['left'])
        section.right_margin = Inches(self.styles['page_margins']['right'])
        
        # 添加标题
        if self.exam_data and 'header' in self.exam_data:
            header_para = self.document.add_paragraph()
            header_run = header_para.add_run(self.exam_data['header'])
            header_run.font.name = self.styles['header_font']
            header_run.font.size = Pt(self.styles['header_size'])
            header_run.font.bold = self.styles['header_bold']
            header_run.font.color.rgb = self.hex_to_rgb(self.styles['header_color'])
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加空行
            self.document.add_paragraph()
        
        # 添加题型和题目
        if self.exam_data and 'questions' in self.exam_data:
            questions = self.exam_data['questions']
            
            # 单选题
            if 'single_choice' in questions and questions['single_choice']:
                sc_questions = questions['single_choice']
                if sc_questions:
                    sc_config = self.exam_data.get('question_types', {}).get('single_choice', {})
                    score_per_question = sc_config.get('score_per_question', 1)
                    
                    section_title = f"一、单选题（每题{score_per_question}分，共{len(sc_questions) * score_per_question}分）"
                    section_para = self.document.add_paragraph()
                    section_run = section_para.add_run(section_title)
                    section_run.font.name = self.styles['question_font']
                    section_run.font.size = Pt(self.styles['question_size'])
                    section_run.font.bold = True
                    
                    for i, q in enumerate(sc_questions, 1):
                        self._add_single_choice_question(i, q)
            
            # 是非题
            if 'true_false' in questions and questions['true_false']:
                tf_questions = questions['true_false']
                if tf_questions:
                    tf_config = self.exam_data.get('question_types', {}).get('true_false', {})
                    score_per_question = tf_config.get('score_per_question', 1)
                    
                    section_title = f"二、是非题（每题{score_per_question}分，共{len(tf_questions) * score_per_question}分）"
                    section_para = self.document.add_paragraph()
                    section_run = section_para.add_run(section_title)
                    section_run.font.name = self.styles['question_font']
                    section_run.font.size = Pt(self.styles['question_size'])
                    section_run.font.bold = True
                    
                    for i, q in enumerate(tf_questions, 1):
                        self._add_true_false_question(i, q)
            
            # 论述题
            if 'essay' in questions and questions['essay']:
                es_questions = questions['essay']
                if es_questions:
                    es_config = self.exam_data.get('question_types', {}).get('essay', {})
                    score_per_question = es_config.get('score_per_question', 1)
                    
                    section_title = f"三、论述题（每题{score_per_question}分，共{len(es_questions) * score_per_question}分）"
                    section_para = self.document.add_paragraph()
                    section_run = section_para.add_run(section_title)
                    section_run.font.name = self.styles['question_font']
                    section_run.font.size = Pt(self.styles['question_size'])
                    section_run.font.bold = True
                    
                    for i, q in enumerate(es_questions, 1):
                        self._add_essay_question(i, q)
            
            # 计算题
            if 'calculation' in questions and questions['calculation']:
                calc_questions = questions['calculation']
                if calc_questions:
                    calc_config = self.exam_data.get('question_types', {}).get('calculation', {})
                    score_per_question = calc_config.get('score_per_question', 1)
                    
                    section_title = f"四、计算题（每题{score_per_question}分，共{len(calc_questions) * score_per_question}分）"
                    section_para = self.document.add_paragraph()
                    section_run = section_para.add_run(section_title)
                    section_run.font.name = self.styles['question_font']
                    section_run.font.size = Pt(self.styles['question_size'])
                    section_run.font.bold = True
                    
                    for i, q in enumerate(calc_questions, 1):
                        self._add_calculation_question(i, q)
    
    def _add_single_choice_question(self, index, question):
        """添加单选题"""
        # 题目
        q_para = self.document.add_paragraph()
        q_run = q_para.add_run(f"{index}. {question.get('chinese_stem', '')}")
        q_run.font.name = self.styles['question_font']
        q_run.font.size = Pt(self.styles['question_size'])
        q_run.font.bold = self.styles['question_bold']
        q_run.font.color.rgb = self.hex_to_rgb(self.styles['question_color'])
        
        # 选项
        options = question.get('options', {})
        for opt_key in ['A', 'B', 'C', 'D']:
            if opt_key in options:
                opt_para = self.document.add_paragraph()
                opt_run = opt_para.add_run(f"{opt_key}. {options[opt_key]}")
                opt_run.font.name = self.styles['option_font']
                opt_run.font.size = Pt(self.styles['option_size'])
                opt_run.font.bold = self.styles['option_bold']
                opt_run.font.color.rgb = self.hex_to_rgb(self.styles['option_color'])
        
        # 答案
        correct_answer = question.get('correct_answer', '')
        if correct_answer:
            ans_para = self.document.add_paragraph()
            ans_run = ans_para.add_run(f"正确答案：{correct_answer}")
            ans_run.font.name = self.styles['answer_font']
            ans_run.font.size = Pt(self.styles['answer_size'])
            ans_run.font.bold = self.styles['answer_bold']
            ans_run.font.color.rgb = self.hex_to_rgb(self.styles['answer_color'])
    
    def _add_true_false_question(self, index, question):
        """添加是非题"""
        q_para = self.document.add_paragraph()
        q_run = q_para.add_run(f"{index}. {question.get('chinese_stem', '')}")
        q_run.font.name = self.styles['question_font']
        q_run.font.size = Pt(self.styles['question_size'])
        q_run.font.bold = self.styles['question_bold']
        q_run.font.color.rgb = self.hex_to_rgb(self.styles['question_color'])
        
        # 答案
        correct_answer = question.get('correct_answer', '')
        if correct_answer:
            ans_para = self.document.add_paragraph()
            ans_run = ans_para.add_run(f"答案：{correct_answer}")
            ans_run.font.name = self.styles['answer_font']
            ans_run.font.size = Pt(self.styles['answer_size'])
            ans_run.font.bold = self.styles['answer_bold']
            ans_run.font.color.rgb = self.hex_to_rgb(self.styles['answer_color'])
    
    def _add_essay_question(self, index, question):
        """添加论述题"""
        q_para = self.document.add_paragraph()
        q_run = q_para.add_run(f"{index}. {question.get('chinese_stem', '')}")
        q_run.font.name = self.styles['question_font']
        q_run.font.size = Pt(self.styles['question_size'])
        q_run.font.bold = self.styles['question_bold']
        q_run.font.color.rgb = self.hex_to_rgb(self.styles['question_color'])
    
    def _add_calculation_question(self, index, question):
        """添加计算题"""
        context = question.get('context', {})
        q_para = self.document.add_paragraph()
        q_run = q_para.add_run(f"{index}. {context.get('chinese', '')}")
        q_run.font.name = self.styles['question_font']
        q_run.font.size = Pt(self.styles['question_size'])
        q_run.font.bold = self.styles['question_bold']
        q_run.font.color.rgb = self.hex_to_rgb(self.styles['question_color'])
    
    def hex_to_rgb(self, hex_color):
        """将十六进制颜色转换为RGB"""
        from docx.shared import RGBColor
        hex_color = hex_color.lstrip('#')
        rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        return RGBColor(rgb[0], rgb[1], rgb[2])
    
    def save_document(self, file_path):
        """保存文档"""
        if self.document:
            self.document.save(file_path)
            return True
        return False
    
    def update_style(self, style_updates):
        """更新样式设置"""
        for key, value in style_updates.items():
            if key in self.styles:
                self.styles[key] = value


class WordExportPreview:
    def __init__(self, exam_data):
        self.exam_data = exam_data
        self.exporter = WordExporter()
        self.exporter.set_exam_data(exam_data)
        
        # 创建主窗口
        self.root = tk.Tk()
        self.root.title("Word导出预览和样式设置")
        self.root.geometry("1000x700")
        
        self.setup_ui()
    
    def setup_ui(self):
        """设置用户界面"""
        # 创建主框架
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 左侧：样式设置面板
        style_frame = ttk.LabelFrame(main_frame, text="样式设置", padding=10)
        style_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        
        # 标题样式设置
        ttk.Label(style_frame, text="标题样式:").grid(row=0, column=0, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="字体:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.header_font_var = tk.StringVar(value=self.exporter.styles['header_font'])
        header_font_combo = ttk.Combobox(style_frame, textvariable=self.header_font_var, 
                                        values=['SimSun', 'SimHei', 'KaiTi', 'Microsoft YaHei', 'Arial', 'Times New Roman'])
        header_font_combo.grid(row=1, column=1, sticky=tk.EW, pady=2)
        
        ttk.Label(style_frame, text="字号:").grid(row=2, column=0, sticky=tk.W, pady=2)
        self.header_size_var = tk.IntVar(value=self.exporter.styles['header_size'])
        header_size_spin = ttk.Spinbox(style_frame, from_=8, to=72, textvariable=self.header_size_var)
        header_size_spin.grid(row=2, column=1, sticky=tk.EW, pady=2)
        
        self.header_bold_var = tk.BooleanVar(value=self.exporter.styles['header_bold'])
        header_bold_check = ttk.Checkbutton(style_frame, text="加粗", variable=self.header_bold_var)
        header_bold_check.grid(row=3, column=0, columnspan=2, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="颜色:").grid(row=4, column=0, sticky=tk.W, pady=2)
        self.header_color_btn = tk.Button(style_frame, bg=self.exporter.styles['header_color'], 
                                         width=3, command=lambda: self.choose_color('header'))
        self.header_color_btn.grid(row=4, column=1, sticky=tk.W, pady=2)
        
        # 题目样式设置
        ttk.Separator(style_frame, orient='horizontal').grid(row=5, column=0, columnspan=2, sticky='ew', pady=10)
        
        ttk.Label(style_frame, text="题目样式:").grid(row=6, column=0, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="字体:").grid(row=7, column=0, sticky=tk.W, pady=2)
        self.question_font_var = tk.StringVar(value=self.exporter.styles['question_font'])
        question_font_combo = ttk.Combobox(style_frame, textvariable=self.question_font_var, 
                                         values=['SimSun', 'SimHei', 'KaiTi', 'Microsoft YaHei', 'Arial', 'Times New Roman'])
        question_font_combo.grid(row=7, column=1, sticky=tk.EW, pady=2)
        
        ttk.Label(style_frame, text="字号:").grid(row=8, column=0, sticky=tk.W, pady=2)
        self.question_size_var = tk.IntVar(value=self.exporter.styles['question_size'])
        question_size_spin = ttk.Spinbox(style_frame, from_=8, to=72, textvariable=self.question_size_var)
        question_size_spin.grid(row=8, column=1, sticky=tk.EW, pady=2)
        
        self.question_bold_var = tk.BooleanVar(value=self.exporter.styles['question_bold'])
        question_bold_check = ttk.Checkbutton(style_frame, text="加粗", variable=self.question_bold_var)
        question_bold_check.grid(row=9, column=0, columnspan=2, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="颜色:").grid(row=10, column=0, sticky=tk.W, pady=2)
        self.question_color_btn = tk.Button(style_frame, bg=self.exporter.styles['question_color'], 
                                          width=3, command=lambda: self.choose_color('question'))
        self.question_color_btn.grid(row=10, column=1, sticky=tk.W, pady=2)
        
        # 选项样式设置
        ttk.Separator(style_frame, orient='horizontal').grid(row=11, column=0, columnspan=2, sticky='ew', pady=10)
        
        ttk.Label(style_frame, text="选项样式:").grid(row=12, column=0, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="字体:").grid(row=13, column=0, sticky=tk.W, pady=2)
        self.option_font_var = tk.StringVar(value=self.exporter.styles['option_font'])
        option_font_combo = ttk.Combobox(style_frame, textvariable=self.option_font_var, 
                                       values=['SimSun', 'SimHei', 'KaiTi', 'Microsoft YaHei', 'Arial', 'Times New Roman'])
        option_font_combo.grid(row=13, column=1, sticky=tk.EW, pady=2)
        
        ttk.Label(style_frame, text="字号:").grid(row=14, column=0, sticky=tk.W, pady=2)
        self.option_size_var = tk.IntVar(value=self.exporter.styles['option_size'])
        option_size_spin = ttk.Spinbox(style_frame, from_=8, to=72, textvariable=self.option_size_var)
        option_size_spin.grid(row=14, column=1, sticky=tk.EW, pady=2)
        
        self.option_bold_var = tk.BooleanVar(value=self.exporter.styles['option_bold'])
        option_bold_check = ttk.Checkbutton(style_frame, text="加粗", variable=self.option_bold_var)
        option_bold_check.grid(row=15, column=0, columnspan=2, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="颜色:").grid(row=16, column=0, sticky=tk.W, pady=2)
        self.option_color_btn = tk.Button(style_frame, bg=self.exporter.styles['option_color'], 
                                        width=3, command=lambda: self.choose_color('option'))
        self.option_color_btn.grid(row=16, column=1, sticky=tk.W, pady=2)
        
        # 答案样式设置
        ttk.Separator(style_frame, orient='horizontal').grid(row=17, column=0, columnspan=2, sticky='ew', pady=10)
        
        ttk.Label(style_frame, text="答案样式:").grid(row=18, column=0, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="字体:").grid(row=19, column=0, sticky=tk.W, pady=2)
        self.answer_font_var = tk.StringVar(value=self.exporter.styles['answer_font'])
        answer_font_combo = ttk.Combobox(style_frame, textvariable=self.answer_font_var, 
                                       values=['SimSun', 'SimHei', 'KaiTi', 'Microsoft YaHei', 'Arial', 'Times New Roman'])
        answer_font_combo.grid(row=19, column=1, sticky=tk.EW, pady=2)
        
        ttk.Label(style_frame, text="字号:").grid(row=20, column=0, sticky=tk.W, pady=2)
        self.answer_size_var = tk.IntVar(value=self.exporter.styles['answer_size'])
        answer_size_spin = ttk.Spinbox(style_frame, from_=8, to=72, textvariable=self.answer_size_var)
        answer_size_spin.grid(row=20, column=1, sticky=tk.EW, pady=2)
        
        self.answer_bold_var = tk.BooleanVar(value=self.exporter.styles['answer_bold'])
        answer_bold_check = ttk.Checkbutton(style_frame, text="加粗", variable=self.answer_bold_var)
        answer_bold_check.grid(row=21, column=0, columnspan=2, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="颜色:").grid(row=22, column=0, sticky=tk.W, pady=2)
        self.answer_color_btn = tk.Button(style_frame, bg=self.exporter.styles['answer_color'], 
                                        width=3, command=lambda: self.choose_color('answer'))
        self.answer_color_btn.grid(row=22, column=1, sticky=tk.W, pady=2)
        
        # 页面边距设置
        ttk.Separator(style_frame, orient='horizontal').grid(row=23, column=0, columnspan=2, sticky='ew', pady=10)
        
        ttk.Label(style_frame, text="页面边距(英寸):").grid(row=24, column=0, sticky=tk.W, pady=2)
        
        ttk.Label(style_frame, text="上:").grid(row=25, column=0, sticky=tk.W, pady=2)
        self.margin_top_var = tk.DoubleVar(value=self.exporter.styles['page_margins']['top'])
        margin_top_spin = ttk.Spinbox(style_frame, from_=0.1, to=5.0, increment=0.1, textvariable=self.margin_top_var)
        margin_top_spin.grid(row=25, column=1, sticky=tk.EW, pady=2)
        
        ttk.Label(style_frame, text="下:").grid(row=26, column=0, sticky=tk.W, pady=2)
        self.margin_bottom_var = tk.DoubleVar(value=self.exporter.styles['page_margins']['bottom'])
        margin_bottom_spin = ttk.Spinbox(style_frame, from_=0.1, to=5.0, increment=0.1, textvariable=self.margin_bottom_var)
        margin_bottom_spin.grid(row=26, column=1, sticky=tk.EW, pady=2)
        
        ttk.Label(style_frame, text="左:").grid(row=27, column=0, sticky=tk.W, pady=2)
        self.margin_left_var = tk.DoubleVar(value=self.exporter.styles['page_margins']['left'])
        margin_left_spin = ttk.Spinbox(style_frame, from_=0.1, to=5.0, increment=0.1, textvariable=self.margin_left_var)
        margin_left_spin.grid(row=27, column=1, sticky=tk.EW, pady=2)
        
        ttk.Label(style_frame, text="右:").grid(row=28, column=0, sticky=tk.W, pady=2)
        self.margin_right_var = tk.DoubleVar(value=self.exporter.styles['page_margins']['right'])
        margin_right_spin = ttk.Spinbox(style_frame, from_=0.1, to=5.0, increment=0.1, textvariable=self.margin_right_var)
        margin_right_spin.grid(row=28, column=1, sticky=tk.EW, pady=2)
        
        # 按钮框架
        button_frame = ttk.Frame(style_frame)
        button_frame.grid(row=29, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="应用样式", command=self.apply_styles).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="导出Word", command=self.export_word).pack(side=tk.LEFT)
        
        # 右侧：预览框架
        preview_frame = ttk.LabelFrame(main_frame, text="预览", padding=10)
        preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # 预览文本框
        self.preview_text = tk.Text(preview_frame, wrap=tk.WORD)
        scrollbar = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=self.preview_text.yview)
        self.preview_text.configure(yscrollcommand=scrollbar.set)
        
        self.preview_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 显示初始预览
        self.update_preview()
    
    def choose_color(self, element_type):
        """选择颜色"""
        color = colorchooser.askcolor(title=f"选择{element_type}颜色")[1]
        if color:
            if element_type == 'header':
                self.header_color_btn.config(bg=color)
            elif element_type == 'question':
                self.question_color_btn.config(bg=color)
            elif element_type == 'option':
                self.option_color_btn.config(bg=color)
            elif element_type == 'answer':
                self.answer_color_btn.config(bg=color)
    
    def apply_styles(self):
        """应用样式设置"""
        new_styles = {
            'header_font': self.header_font_var.get(),
            'header_size': self.header_size_var.get(),
            'header_bold': self.header_bold_var.get(),
            'question_font': self.question_font_var.get(),
            'question_size': self.question_size_var.get(),
            'question_bold': self.question_bold_var.get(),
            'option_font': self.option_font_var.get(),
            'option_size': self.option_size_var.get(),
            'option_bold': self.option_bold_var.get(),
            'answer_font': self.answer_font_var.get(),
            'answer_size': self.answer_size_var.get(),
            'answer_bold': self.answer_bold_var.get(),
            'header_color': self.header_color_btn.cget('bg'),
            'question_color': self.question_color_btn.cget('bg'),
            'option_color': self.option_color_btn.cget('bg'),
            'answer_color': self.answer_color_btn.cget('bg'),
            'page_margins': {
                'top': self.margin_top_var.get(),
                'bottom': self.margin_bottom_var.get(),
                'left': self.margin_left_var.get(),
                'right': self.margin_right_var.get()
            }
        }
        
        self.exporter.update_style(new_styles)
        self.update_preview()
    
    def update_preview(self):
        """更新预览"""
        # 清空预览
        self.preview_text.delete(1.0, tk.END)
        
        # 添加预览内容
        if self.exam_data:
            header = self.exam_data.get('header', '试卷标题')
            self.preview_text.insert(tk.END, header + "\n\n")
            
            questions = self.exam_data.get('questions', {})
            
            # 单选题预览
            if 'single_choice' in questions and questions['single_choice']:
                sc_questions = questions['single_choice']
                if sc_questions:
                    self.preview_text.insert(tk.END, "一、单选题\n")
                    for i, q in enumerate(sc_questions[:3], 1):  # 只显示前3题作为预览
                        self.preview_text.insert(tk.END, f"{i}. {q.get('chinese_stem', '')}\n")
                        options = q.get('options', {})
                        for opt_key in ['A', 'B', 'C', 'D']:
                            if opt_key in options:
                                self.preview_text.insert(tk.END, f"   {opt_key}. {options[opt_key]}\n")
                        self.preview_text.insert(tk.END, f"   正确答案：{q.get('correct_answer', '')}\n\n")
            
            # 是非题预览
            if 'true_false' in questions and questions['true_false'] and len(questions['true_false']) > 0:
                tf_questions = questions['true_false'][:2]  # 只显示前2题作为预览
                self.preview_text.insert(tk.END, "二、是非题\n")
                for i, q in enumerate(tf_questions, 1):
                    self.preview_text.insert(tk.END, f"{i}. {q.get('chinese_stem', '')}\n")
                    self.preview_text.insert(tk.END, f"   答案：{q.get('correct_answer', '')}\n\n")
    
    def export_word(self):
        """导出Word文档"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
            title="保存Word文档"
        )
        
        if file_path:
            try:
                self.exporter.create_document()
                if self.exporter.save_document(file_path):
                    messagebox.showinfo("成功", f"Word文档已保存到：{file_path}")
                else:
                    messagebox.showerror("错误", "保存文档失败")
            except Exception as e:
                messagebox.showerror("错误", f"导出过程中出现错误：{str(e)}")
    
    def run(self):
        """运行预览窗口"""
        self.root.mainloop()


def main():
    # 示例试卷数据
    sample_exam_data = {
        "header": "中南财经政法大学\n2025–2026学年第1学期\n林业经济学 试卷（A卷）\n总分：100分",
        "question_types": {
            "single_choice": {
                "count": 5,
                "score_per_question": 2
            },
            "true_false": {
                "count": 5,
                "score_per_question": 2
            },
            "essay": {
                "count": 2,
                "score_per_question": 15
            },
            "calculation": {
                "count": 1,
                "score_per_question": 20
            }
        },
        "questions": {
            "single_choice": [
                {
                    "id": "sc_001",
                    "chinese_stem": "Merit goods在林业经济学中指的是什么？（）",
                    "english_stem": "What does 'Merit goods' refer to in forest economics? ()",
                    "options": {
                        "A": "社会价值低于私人消费者价值的商品",
                        "A_en": "Merchandise with social value lower than private consumer value",
                        "B": "社会价值超过私人消费者价值的商品",
                        "B_en": "Merchandise with social value exceeding private consumer value",
                        "C": "仅具有私人价值的商品",
                        "C_en": "Merchandise with private value only",
                        "D": "仅具有社会价值的商品",
                        "D_en": "Merchandise with social value only"
                    },
                    "correct_answer": "B",
                    "difficulty": "medium",
                    "knowledge_point": "Merit goods definition",
                    "tags": ["basic", "definition", "public goods"]
                }
            ],
            "true_false": [
                {
                    "id": "tf_001",
                    "chinese_stem": "Contingent valuation method是一种了解消费者偏好的显示型（Revealed preference)方法。",
                    "english_stem": "The Contingent Valuation Method (CVM) is a revealed preference approach for eliciting consumer preferences.",
                    "correct_answer": "F",
                    "explanation": "CVM是陈述型（Stated preference）方法，而非显示型（Revealed preference）方法。",
                    "explanation_en": "CVM is a stated preference approach, not a revealed preference approach.",
                    "knowledge_point": "CVM methodology",
                    "tags": ["valuation", "method", "preference"]
                }
            ],
            "essay": [
                {
                    "id": "es_001",
                    "chinese_stem": "有哪些不同类型的外部性？请列举出至少三种，并举例分析每种外部性分别是如何产生的？有何应对措施？",
                    "english_stem": "What are the different types of externalities? Please identify at least three distinct categories and provide illustrative examples demonstrating how each type of externality arises. What policy interventions or remedial measures can be implemented to address them?",
                    "scoring_guide": ["识别三种以上外部性类型 (5分)", "每种类型提供具体例子 (5分)", "分析产生机制 (3分)", "提出应对措施 (2分)"],
                    "scoring_guide_en": ["Identify three or more types of externalities (5 points)", "Provide specific examples for each type (5 points)", "Analyze generation mechanisms (3 points)", "Propose countermeasures (2 points)"],
                    "knowledge_point": "Types of externalities",
                    "tags": ["externality", "classification", "policy"]
                }
            ],
            "calculation": [
                {
                    "id": "calc_001",
                    "context": {
                        "chinese": "某村庄有一片荒地，如果自然更新，每30年可以产出价值10.9万元的木材，采伐成本为1万元。村委会现在要从以下投资决策中做出选择：",
                        "english": "A village possesses a tract of wasteland. Under natural regeneration, it would yield timber valued at ¥109,000 every 30 years, with harvesting costs of ¥10,000. The village committee must now choose among the following investment alternatives:"
                    },
                    "alternatives": [],
                    "parameters": {"折现率": "0.05"},
                    "requirements": {
                        "chinese": ["请计算不同投资决策的净现值，并根据净现值考虑选择哪种投资决策？", "假如你是村委会的决策人，在投资决策时，除了净现值以外，你还将考虑哪些因素？分析在考虑这些因素后，你会如何改变决策？"],
                        "english": ["Calculate the net present value (NPV) of each investment alternative and determine which option should be selected based on NPV analysis.", "If you were the decision-maker for the village committee, what additional factors beyond NPV would you consider in the investment decision? Analyze how consideration of these factors might alter your decision."]
                    },
                    "knowledge_point": "NPV calculation",
                    "tags": ["NPV", "investment", "decision"]
                }
            ]
        }
    }
    
    # 创建并运行预览窗口
    preview = WordExportPreview(sample_exam_data)
    preview.run()


if __name__ == "__main__":
    main()