"""
改进版Python脚本用于读取Word文档内容和基本格式信息
解决编码问题
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
import os
import sys

def clean_text(text):
    """
    清理文本中的特殊字符
    """
    if text is None:
        return ""
    # 移除不可见字符和特殊字符
    cleaned = text.replace('\x00', '').replace('\ufffd', '')
    return cleaned

def read_word_document(file_path):
    """
    读取Word文档内容和基本格式信息
    """
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 - {file_path}")
        return
    
    try:
        doc = Document(file_path)
        print(f"正在读取文档: {file_path}")
        print("="*50)
        
        # 读取段落内容和格式
        print("段落内容:")
        print("-" * 30)
        for i, paragraph in enumerate(doc.paragraphs):
            text = clean_text(paragraph.text)
            if text.strip():  # 只显示非空段落
                print(f"段落 {i+1}:")
                print(f"  文本: {text}")
                
                # 检查段落格式
                para_format = paragraph.paragraph_format
                print(f"  对齐方式: {para_format.alignment if para_format.alignment is not None else '默认'}")
                
                # 读取段落中的格式化文本
                if paragraph.runs:
                    print("  格式化文本:")
                    for j, run in enumerate(paragraph.runs):
                        run_text = clean_text(run.text)
                        if run_text.strip():
                            print(f"    Run {j+1}: '{run_text}'")
                            print(f"      粗体: {run.bold}")
                            print(f"      斜体: {run.italic}")
                            print(f"      下划线: {run.underline}")
                            if run.font.size:
                                print(f"      字体大小: {run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size}")
                            if run.font.name:
                                print(f"      字体名称: {run.font.name}")
                    print()
        
        # 读取表格内容
        print("\n表格内容:")
        print("-" * 30)
        for i, table in enumerate(doc.tables):
            print(f"表格 {i+1}:")
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    row_data.append(cell_text)
                print(f"  行 {row_idx+1}: {row_data}")
        
        # 读取文档属性
        print("\n文档属性:")
        print("-" * 30)
        core_props = doc.core_properties
        print(f"标题: {clean_text(core_props.title) if core_props.title else '无'}")
        print(f"作者: {clean_text(core_props.author) if core_props.author else '无'}")
        print(f"主题: {clean_text(core_props.subject) if core_props.subject else '无'}")
        print(f"创建时间: {core_props.created if core_props.created else '无'}")
        print(f"修改时间: {core_props.modified if core_props.modified else '无'}")
        
    except Exception as e:
        print(f"读取文档时发生错误: {str(e)}")
        import traceback
        traceback.print_exc()

def extract_text_only(file_path):
    """
    仅提取文档文本内容
    """
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if text.strip():
                full_text.append(text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    if cell_text.strip():
                        full_text.append(cell_text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"提取文本时发生错误: {str(e)}")
        return ""

def extract_text_with_formatting(file_path):
    """
    提取带格式的文本内容
    """
    try:
        doc = Document(file_path)
        formatted_text = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = clean_text(paragraph.text)
            if text.strip():
                # 获取段落对齐方式
                alignment = paragraph.paragraph_format.alignment
                alignment_text = ""
                if alignment == 0:
                    alignment_text = "(左对齐)"
                elif alignment == 1:
                    alignment_text = "(居中)"
                elif alignment == 2:
                    alignment_text = "(右对齐)"
                elif alignment == 3:
                    alignment_text = "(两端对齐)"
                
                formatted_text.append(f"段落 {i+1}{alignment_text}: {text}")
                
                # 检查格式化文本
                if paragraph.runs:
                    for j, run in enumerate(paragraph.runs):
                        run_text = clean_text(run.text)
                        if run_text.strip():
                            formatting_info = []
                            if run.bold:
                                formatting_info.append("粗体")
                            if run.italic:
                                formatting_info.append("斜体")
                            if run.underline:
                                formatting_info.append("下划线")
                            if run.font.size:
                                size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                                formatting_info.append(f"字体大小: {size}")
                            if run.font.name:
                                formatting_info.append(f"字体: {run.font.name}")
                            
                            if formatting_info:
                                formatted_text.append(f"  Run {j+1} ({', '.join(formatting_info)}): '{run_text}'")
        
        # 提取表格内容
        for i, table in enumerate(doc.tables):
            formatted_text.append(f"\n表格 {i+1}:")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = clean_text(cell.text)
                    if cell_text.strip():
                        formatted_text.append(f"  行{row_idx+1}列{cell_idx+1}: {cell_text}")
        
        return '\n'.join(formatted_text)
    except Exception as e:
        print(f"提取格式化文本时发生错误: {str(e)}")
        return ""

if __name__ == "__main__":
    # 请将此处路径替换为您的Word文档路径
    file_path = r"D:\code\试卷\林业经济学（双语）试卷A-打印.docx"
    
    # 读取完整文档信息
    print("正在读取文档信息...")
    read_word_document(file_path)
    
    print("\n" + "="*70)
    print("仅文本内容:")
    print("-" * 30)
    text_content = extract_text_only(file_path)
    print(text_content)
    
    print("\n" + "="*70)
    print("格式化文本内容:")
    print("-" * 30)
    formatted_content = extract_text_with_formatting(file_path)
    print(formatted_content)