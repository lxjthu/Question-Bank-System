"""
Word文档到CSV转换器
用于将Word格式的试题模板转换为CSV格式
"""

import json
import csv
import os
from docx import Document
import re
from datetime import datetime


class WordToCsvConverter:
    def __init__(self):
        self.question_types = {
            'single_choice': self.parse_single_choice,
            'true_false': self.parse_true_false,
            'essay': self.parse_essay,
            'calculation': self.parse_calculation
        }
        self.warnings = []  # 存储警告信息
    
    def convert_word_to_csv(self, word_file_path, output_csv_path):
        """
        将Word文档转换为CSV格式
        
        Args:
            word_file_path: Word文档路径
            output_csv_path: 输出CSV路径
        """
        try:
            # 读取Word文档
            doc = Document(word_file_path)
            
            # 提取所有文本
            all_text = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    all_text.append(text)
            
            # 解析题目
            questions = self.parse_questions_from_text(all_text)
            
            # 保存为CSV
            self.save_questions_to_csv(questions, output_csv_path)
            
            return True, f"成功转换 {len(questions)} 道题目"
        
        except Exception as e:
            return False, f"转换失败: {str(e)}"
    
    def parse_questions_from_text(self, text_lines):
        """
        从文本行中解析题目
        
        Args:
            text_lines: 文本行列表
            
        Returns:
            题目列表
        """
        questions = []
        current_question = None
        current_type = None
        
        i = 0
        while i < len(text_lines):
            line = text_lines[i].strip()
            
            # 检查题目类型
            if line.startswith("单选题ID:"):
                current_type = "single_choice"
                current_question = {"type": current_type}
                current_question["id"] = line.replace("单选题ID:", "").strip()
                i += 1
                continue
            elif line.startswith("是非题ID:"):
                current_type = "true_false"
                current_question = {"type": current_type}
                current_question["id"] = line.replace("是非题ID:", "").strip()
                i += 1
                continue
            elif line.startswith("论述题ID:"):
                current_type = "essay"
                current_question = {"type": current_type}
                current_question["id"] = line.replace("论述题ID:", "").strip()
                i += 1
                continue
            elif line.startswith("计算题ID:"):
                current_type = "calculation"
                current_question = {"type": current_type}
                current_question["id"] = line.replace("计算题ID:", "").strip()
                i += 1
                continue
            
            if current_question and current_type in self.question_types:
                # 解析当前题目的内容
                question_data = self.question_types[current_type](text_lines, i)
                
                if question_data:
                    # 合并解析结果
                    current_question.update(question_data)
                    questions.append(current_question)
                
                # 跳过已解析的行
                i = question_data.get("line_index", i + 1)  # 假设解析函数返回行索引
                current_question = None
                continue
            
            i += 1
        
        # 简化版解析 - 直接按行解析
        return self.simple_parse_questions(text_lines)
    
    def simple_parse_questions(self, text_lines):
        """
        简化版题目解析
        """
        questions = []
        current_question = {}
        current_type = None

        i = 0
        while i < len(text_lines):
            line = text_lines[i].strip()

            if line.startswith("单选题ID:"):
                current_type = "single_choice"
                current_question = {
                    "type": current_type,
                    "id": line.replace("单选题ID:", "").strip(),
                    "chinese_stem": "",
                    "english_stem": "",
                    "options": {"A": "", "A_en": "", "B": "", "B_en": "", "C": "", "C_en": "", "D": "", "D_en": ""},
                    "correct_answer": "",
                    "difficulty": "medium",  # 默认难度
                    "knowledge_point": "待设置",  # 默认知识点
                    "tags": ["待分类"]  # 默认标签
                }

                # 查找题干和难度、知识点、标签
                i += 1
                while i < len(text_lines):
                    sub_line = text_lines[i].strip()
                    if sub_line.startswith("难度:"):
                        current_question["difficulty"] = sub_line.replace("难度:", "").strip()
                    elif sub_line.startswith("知识点:"):
                        current_question["knowledge_point"] = sub_line.replace("知识点:", "").strip()
                    elif sub_line.startswith("标签:"):
                        tags_str = sub_line.replace("标签:", "").strip()
                        current_question["tags"] = [tag.strip() for tag in tags_str.split(',') if tag.strip()]
                    elif sub_line.startswith("A.") or sub_line.startswith("B.") or sub_line.startswith("C.") or sub_line.startswith("D."):
                        # 遇到选项，跳出当前循环
                        break
                    elif sub_line and not any(sub_line.startswith(prefix) for prefix in ["正确答案:", "单选题ID:", "是非题ID:", "论述题ID:", "计算题ID:"]):
                        if not current_question["chinese_stem"]:
                            current_question["chinese_stem"] = sub_line
                        elif not current_question["english_stem"]:
                            current_question["english_stem"] = sub_line
                    elif sub_line.startswith("单选题ID:") or sub_line.startswith("是非题ID:") or sub_line.startswith("论述题ID:") or sub_line.startswith("计算题ID:"):
                        # 遇到下一个题目，结束当前题目解析
                        break
                    i += 1

                # 查找选项（包括英文选项）
                option_pattern = re.compile(r'^(A|B|C|D)\.(.*)')
                while i < len(text_lines):
                    sub_line = text_lines[i].strip()
                    match = option_pattern.match(sub_line)
                    if match:
                        option = match.group(1)
                        text = match.group(2).strip()
                        current_question["options"][option] = text
                        i += 1

                        # 检查下一行是否是英文选项
                        if i < len(text_lines):
                            next_line = text_lines[i].strip()
                            if next_line.startswith(f"{option}_en."):
                                en_text = next_line.replace(f"{option}_en.", "").strip()
                                current_question["options"][f"{option}_en"] = en_text
                                i += 1
                    elif sub_line.startswith("正确答案:"):
                        current_question["correct_answer"] = sub_line.replace("正确答案:", "").strip()
                        i += 1
                        break
                    elif sub_line.startswith("单选题ID:") or sub_line.startswith("是非题ID:") or sub_line.startswith("论述题ID:") or sub_line.startswith("计算题ID:"):
                        # 遇到下一个题目，结束当前题目解析
                        break
                    else:
                        i += 1

                # 检查必填字段是否缺失
                missing_parts = []
                if not current_question["chinese_stem"]:
                    missing_parts.append("中文题干")
                if not current_question["english_stem"]:
                    missing_parts.append("英文题干")
                if not current_question["correct_answer"]:
                    missing_parts.append("正确答案")

                # 检查选项是否完整
                for opt in ['A', 'B', 'C', 'D']:
                    if not current_question["options"][opt]:
                        missing_parts.append(f"选项{opt}")

                if missing_parts:
                    warning_msg = f"警告：单选题 {current_question['id']} 的 {', '.join(missing_parts)} 部分缺失，将显示空白"
                    self.warnings.append(warning_msg)

                questions.append(current_question)
                current_question = {}

            elif line.startswith("是非题ID:"):
                current_type = "true_false"
                current_question = {
                    "type": current_type,
                    "id": line.replace("是非题ID:", "").strip(),
                    "chinese_stem": "",
                    "english_stem": "",
                    "correct_answer": "",
                    "explanation": "",
                    "explanation_en": "",
                    "knowledge_point": "待设置",
                    "tags": ["待分类"]
                }

                # 查找题干和答案
                i += 1
                while i < len(text_lines):
                    sub_line = text_lines[i].strip()
                    if sub_line.startswith("知识点:"):
                        current_question["knowledge_point"] = sub_line.replace("知识点:", "").strip()
                    elif sub_line.startswith("标签:"):
                        tags_str = sub_line.replace("标签:", "").strip()
                        current_question["tags"] = [tag.strip() for tag in tags_str.split(',') if tag.strip()]
                    elif sub_line.startswith("答案:"):
                        current_question["correct_answer"] = sub_line.replace("答案:", "").strip()
                    elif sub_line.startswith("解释:"):
                        current_question["explanation"] = sub_line.replace("解释:", "").strip()
                    elif sub_line.startswith("English Explanation:"):
                        current_question["explanation_en"] = sub_line.replace("English Explanation:", "").strip()
                    elif sub_line.startswith("单选题ID:") or sub_line.startswith("是非题ID:") or sub_line.startswith("论述题ID:") or sub_line.startswith("计算题ID:"):
                        # 遇到下一个题目，结束当前题目解析
                        break
                    else:
                        if not current_question["chinese_stem"]:
                            current_question["chinese_stem"] = sub_line
                        elif not current_question["english_stem"]:
                            current_question["english_stem"] = sub_line
                    i += 1

                # 检查必填字段是否缺失
                missing_parts = []
                if not current_question["chinese_stem"]:
                    missing_parts.append("中文题干")
                if not current_question["english_stem"]:
                    missing_parts.append("英文题干")
                if not current_question["correct_answer"]:
                    missing_parts.append("正确答案")

                if missing_parts:
                    warning_msg = f"警告：是非题 {current_question['id']} 的 {', '.join(missing_parts)} 部分缺失，将显示空白"
                    self.warnings.append(warning_msg)

                questions.append(current_question)
                current_question = {}

            elif line.startswith("论述题ID:"):
                current_type = "essay"
                current_question = {
                    "type": current_type,
                    "id": line.replace("论述题ID:", "").strip(),
                    "chinese_stem": "",
                    "english_stem": "",
                    "scoring_guide": [],
                    "scoring_guide_en": [],
                    "knowledge_point": "待设置",
                    "tags": ["待分类"]
                }

                # 查找题干和评分标准
                i += 1
                is_scoring_guide = False
                while i < len(text_lines):
                    sub_line = text_lines[i].strip()
                    if sub_line.startswith("知识点:"):
                        current_question["knowledge_point"] = sub_line.replace("知识点:", "").strip()
                    elif sub_line.startswith("标签:"):
                        tags_str = sub_line.replace("标签:", "").strip()
                        current_question["tags"] = [tag.strip() for tag in tags_str.split(',') if tag.strip()]
                    elif sub_line.startswith("评分标准:"):
                        is_scoring_guide = True
                        guide_text = sub_line.replace("评分标准:", "").strip()
                        if guide_text:
                            current_question["scoring_guide"].append(guide_text)
                    elif sub_line.startswith("English Scoring Guide:"):
                        guide_text = sub_line.replace("English Scoring Guide:", "").strip()
                        if guide_text:
                            current_question["scoring_guide_en"].append(guide_text)
                    elif sub_line.startswith("单选题ID:") or sub_line.startswith("是非题ID:") or sub_line.startswith("论述题ID:") or sub_line.startswith("计算题ID:"):
                        # 遇到下一个题目，结束当前题目解析
                        break
                    else:
                        if not is_scoring_guide and not current_question["chinese_stem"]:
                            current_question["chinese_stem"] = sub_line
                        elif not is_scoring_guide and not current_question["english_stem"]:
                            current_question["english_stem"] = sub_line
                    i += 1

                # 检查必填字段是否缺失
                missing_parts = []
                if not current_question["chinese_stem"]:
                    missing_parts.append("中文题干")
                if not current_question["english_stem"]:
                    missing_parts.append("英文题干")

                if missing_parts:
                    warning_msg = f"警告：论述题 {current_question['id']} 的 {', '.join(missing_parts)} 部分缺失，将显示空白"
                    self.warnings.append(warning_msg)

                questions.append(current_question)
                current_question = {}

            elif line.startswith("计算题ID:"):
                current_type = "calculation"
                current_question = {
                    "type": current_type,
                    "id": line.replace("计算题ID:", "").strip(),
                    "context": {"chinese": "", "english": ""},
                    "alternatives": [],
                    "parameters": {},
                    "requirements": {"chinese": [], "english": []},
                    "knowledge_point": "待设置",
                    "tags": ["待分类"]
                }

                # 查找计算题内容
                i += 1
                while i < len(text_lines):
                    sub_line = text_lines[i].strip()
                    if sub_line.startswith("知识点:"):
                        current_question["knowledge_point"] = sub_line.replace("知识点:", "").strip()
                    elif sub_line.startswith("标签:"):
                        tags_str = sub_line.replace("标签:", "").strip()
                        current_question["tags"] = [tag.strip() for tag in tags_str.split(',') if tag.strip()]
                    elif sub_line.startswith("参数:"):
                        # 解析参数
                        params_str = sub_line.replace("参数:", "").strip()
                        param_pairs = params_str.split(';')
                        for pair in param_pairs:
                            if ':' in pair:
                                key, value = pair.split(':', 1)
                                current_question["parameters"][key.strip()] = value.strip()
                    elif sub_line.startswith("要求:"):
                        req_text = sub_line.replace("要求:", "").strip()
                        if req_text:
                            current_question["requirements"]["chinese"].append(req_text)
                    elif sub_line.startswith("English Requirements:"):
                        req_text = sub_line.replace("English Requirements:", "").strip()
                        if req_text:
                            current_question["requirements"]["english"].append(req_text)
                    elif sub_line.startswith("单选题ID:") or sub_line.startswith("是非题ID:") or sub_line.startswith("论述题ID:") or sub_line.startswith("计算题ID:"):
                        # 遇到下一个题目，结束当前题目解析
                        break
                    else:
                        # 处理背景信息
                        if not current_question["context"]["chinese"]:
                            current_question["context"]["chinese"] = sub_line
                        elif not current_question["context"]["english"]:
                            current_question["context"]["english"] = sub_line
                    i += 1

                # 检查必填字段是否缺失
                missing_parts = []
                if not current_question["context"]["chinese"]:
                    missing_parts.append("中文背景")
                if not current_question["context"]["english"]:
                    missing_parts.append("英文背景")

                if missing_parts:
                    warning_msg = f"警告：计算题 {current_question['id']} 的 {', '.join(missing_parts)} 部分缺失，将显示空白"
                    self.warnings.append(warning_msg)

                questions.append(current_question)
                current_question = {}
            else:
                i += 1

        return questions
    
    def parse_single_choice(self, text_lines, start_index):
        """解析单选题"""
        # 这是一个简化示例，实际实现会更复杂
        return {}
    
    def parse_true_false(self, text_lines, start_index):
        """解析是非题"""
        return {}
    
    def parse_essay(self, text_lines, start_index):
        """解析论述题"""
        return {}
    
    def parse_calculation(self, text_lines, start_index):
        """解析计算题"""
        return {}
    
    def save_questions_to_csv(self, questions, output_csv_path):
        """
        将题目保存为CSV格式

        Args:
            questions: 题目列表
            output_csv_path: 输出路径
        """
        with open(output_csv_path, 'w', newline='', encoding='utf-8-sig') as csvfile:
            if not questions:
                return

            # 确定字段名
            fieldnames = set()
            for q in questions:
                fieldnames.update(q.keys())
                # 对于嵌套字段，展开它们
                if 'options' in q:
                    for opt_key in q['options'].keys():
                        fieldnames.add(f'option_{opt_key}_ch' if opt_key in ['A', 'B', 'C', 'D'] else f'option_{opt_key}')
                if 'parameters' in q:
                    for param_key in q['parameters'].keys():
                        fieldnames.add(f'parameter_{param_key}')
                if 'requirements' in q:
                    if isinstance(q['requirements'].get('chinese'), list):
                        fieldnames.add('requirements_chinese')
                    if isinstance(q['requirements'].get('english'), list):
                        fieldnames.add('requirements_english')
                if 'scoring_guide' in q:
                    if isinstance(q['scoring_guide'], list):
                        fieldnames.add('scoring_guide')
                    if isinstance(q['scoring_guide_en'], list):
                        fieldnames.add('scoring_guide_en')

            # 定义标准字段
            standard_fieldnames = [
                'id', 'type', 'difficulty', 'chinese_stem', 'english_stem',
                'option_A_ch', 'option_A_en', 'option_B_ch', 'option_B_en',
                'option_C_ch', 'option_C_en', 'option_D_ch', 'option_D_en',
                'correct_answer', 'knowledge_point', 'tags',
                'explanation', 'explanation_en', 'scoring_guide', 'scoring_guide_en',
                'context_chinese', 'context_english', 'parameters', 'requirements_chinese', 'requirements_english'
            ]

            # 合并字段
            all_fieldnames = list(set(standard_fieldnames + list(fieldnames)))

            writer = csv.DictWriter(csvfile, fieldnames=all_fieldnames)
            writer.writeheader()

            for question in questions:
                row = {}

                # 基本字段
                row['id'] = question.get('id', '')
                row['type'] = question.get('type', '')
                row['difficulty'] = question.get('difficulty', 'medium')
                row['chinese_stem'] = question.get('chinese_stem', '')
                row['english_stem'] = question.get('english_stem', '')
                row['knowledge_point'] = question.get('knowledge_point', '待设置')
                row['tags'] = ','.join(question.get('tags', []))

                # 根据题型处理特定字段
                if question.get('type') == 'single_choice' and 'options' in question:
                    row['option_A_ch'] = question['options'].get('A', '')
                    row['option_A_en'] = question['options'].get('A_en', '')
                    row['option_B_ch'] = question['options'].get('B', '')
                    row['option_B_en'] = question['options'].get('B_en', '')
                    row['option_C_ch'] = question['options'].get('C', '')
                    row['option_C_en'] = question['options'].get('C_en', '')
                    row['option_D_ch'] = question['options'].get('D', '')
                    row['option_D_en'] = question['options'].get('D_en', '')
                    row['correct_answer'] = question.get('correct_answer', '')

                elif question.get('type') == 'true_false':
                    row['correct_answer'] = question.get('correct_answer', '')
                    row['explanation'] = question.get('explanation', '')
                    row['explanation_en'] = question.get('explanation_en', '')

                elif question.get('type') == 'essay':
                    row['scoring_guide'] = ';'.join(question.get('scoring_guide', []))
                    row['scoring_guide_en'] = ';'.join(question.get('scoring_guide_en', []))

                elif question.get('type') == 'calculation':
                    row['context_chinese'] = question['context'].get('chinese', '')
                    row['context_english'] = question['context'].get('english', '')
                    # 将参数字典转换为字符串
                    params_str = ';'.join([f"{k}:{v}" for k, v in question.get('parameters', {}).items()])
                    row['parameters'] = params_str
                    # 将要求列表转换为字符串
                    req_ch_str = ';'.join(question['requirements'].get('chinese', []))
                    req_en_str = ';'.join(question['requirements'].get('english', []))
                    row['requirements_chinese'] = req_ch_str
                    row['requirements_english'] = req_en_str

                # 确保所有字段都有值
                for field in all_fieldnames:
                    if field not in row:
                        row[field] = ''

                writer.writerow(row)
    
    def create_word_template(self, template_type, output_path):
        """
        创建Word模板文档
        
        Args:
            template_type: 模板类型 ('single_choice', 'true_false', 'essay', 'calculation')
            output_path: 输出路径
        """
        doc = Document()
        
        # 根据题型创建模板内容
        if template_type == 'single_choice':
            doc.add_heading('林业经济学（双语）试题模板 - 单选题', 0)
            doc.add_paragraph('请按照以下格式填写单选题：')
            doc.add_paragraph('')
            doc.add_paragraph('单选题ID: sc_001')
            doc.add_paragraph('难度: medium')
            doc.add_paragraph('知识点: Merit goods definition')
            doc.add_paragraph('标签: basic,definition,public goods')
            doc.add_paragraph('中文题干: Merit goods在林业经济学中指的是什么？（）')
            doc.add_paragraph('英文题干: What does "Merit goods" refer to in forest economics? ()')
            doc.add_paragraph('A. 社会价值低于私人消费者价值的商品')
            doc.add_paragraph('A_en. Merchandise with social value lower than private consumer value')
            doc.add_paragraph('B. 社会价值超过私人消费者价值的商品')
            doc.add_paragraph('B_en. Merchandise with social value exceeding private consumer value')
            doc.add_paragraph('C. 仅具有私人价值的商品')
            doc.add_paragraph('C_en. Merchandise with private value only')
            doc.add_paragraph('D. 仅具有社会价值的商品')
            doc.add_paragraph('D_en. Merchandise with social value only')
            doc.add_paragraph('正确答案: B')
            
        elif template_type == 'true_false':
            doc.add_heading('林业经济学（双语）试题模板 - 是非题', 0)
            doc.add_paragraph('请按照以下格式填写是非题：')
            doc.add_paragraph('')
            doc.add_paragraph('是非题ID: tf_001')
            doc.add_paragraph('知识点: CVM methodology')
            doc.add_paragraph('标签: valuation,method,preference')
            doc.add_paragraph('中文题干: "Contingent valuation method"是一种了解消费者偏好的显示型（Revealed preference)方法。')
            doc.add_paragraph('英文题干: The Contingent Valuation Method (CVM) is a revealed preference approach for eliciting consumer preferences.')
            doc.add_paragraph('答案: F')
            doc.add_paragraph('解释: CVM是陈述型（Stated preference）方法，而非显示型（Revealed preference）方法。')
            doc.add_paragraph('English Explanation: CVM is a stated preference approach, not a revealed preference approach.')
            
        elif template_type == 'essay':
            doc.add_heading('林业经济学（双语）试题模板 - 论述题', 0)
            doc.add_paragraph('请按照以下格式填写论述题：')
            doc.add_paragraph('')
            doc.add_paragraph('论述题ID: es_001')
            doc.add_paragraph('知识点: Types of externalities')
            doc.add_paragraph('标签: externality,classification,policy')
            doc.add_paragraph('中文题干: 有哪些不同类型的"外部性"？请列举出至少三种，并举例分析每种外部性分别是如何产生的？有何应对措施？')
            doc.add_paragraph('英文题干: What are the different types of externalities? Please identify at least three distinct categories and provide illustrative examples demonstrating how each type of externality arises. What policy interventions or remedial measures can be implemented to address them?')
            doc.add_paragraph('评分标准: 识别三种以上外部性类型 (5分); 每种类型提供具体例子 (5分); 分析产生机制 (3分); 提出应对措施 (2分)')
            doc.add_paragraph('English Scoring Guide: Identify three or more types of externalities (5 points); Provide specific examples for each type (5 points); Analyze generation mechanisms (3 points); Propose countermeasures (2 points)')
            
        elif template_type == 'calculation':
            doc.add_heading('林业经济学（双语）试题模板 - 计算题', 0)
            doc.add_paragraph('请按照以下格式填写计算题：')
            doc.add_paragraph('')
            doc.add_paragraph('计算题ID: calc_001')
            doc.add_paragraph('知识点: NPV calculation')
            doc.add_paragraph('标签: NPV,investment,decision')
            doc.add_paragraph('中文背景: 某村庄有一片荒地，如果自然更新，每30年可以产出价值10.9万元的木材，采伐成本为1万元。村委会现在要从以下投资决策中做出选择：')
            doc.add_paragraph('英文背景: A village possesses a tract of wasteland. Under natural regeneration, it would yield timber valued at ¥109,000 every 30 years, with harvesting costs of ¥10,000. The village committee must now choose among the following investment alternatives:')
            doc.add_paragraph('投资方案1: 投资20万元种植松树，每20年产出木材价值70万元的木材，采伐成本为6万元，每年的管理成本为1000元。')
            doc.add_paragraph('Investment Option 1: Invest ¥200,000 to establish a pine plantation, generating timber valued at ¥700,000 every 20 years, with harvesting costs of ¥60,000 and annual management costs of ¥1,000.')
            doc.add_paragraph('参数: 折现率:0.05; 时间系数:5年=1.3,10年=1.6,20年=2.6,30年=4.3; 公式:F=A*[(1+i)^n-1]/i')
            doc.add_paragraph('要求: 请计算不同投资决策的净现值，并根据净现值考虑选择哪种投资决策？#假如你是村委会的决策人，在投资决策时，除了净现值以外，你还将考虑哪些因素？分析在考虑这些因素后，你会如何改变决策？')
            doc.add_paragraph('English Requirements: Calculate the net present value (NPV) of each investment alternative and determine which option should be selected based on NPV analysis.#If you were the decision-maker for the village committee, what additional factors beyond NPV would you consider in the investment decision? Analyze how consideration of these factors might alter your decision.')
        
        doc.save(output_path)


    def parse_questions_from_word_doc(self, word_file_path):
        """
        从Word文档直接解析题目

        Args:
            word_file_path: Word文档路径

        Returns:
            解析的题目列表
        """
        try:
            # 重置警告列表
            self.warnings = []

            # 读取Word文档
            doc = Document(word_file_path)

            # 提取所有文本
            all_text = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    all_text.append(text)

            # 解析题目
            questions = self.parse_questions_from_text(all_text)

            return questions

        except Exception as e:
            print(f"解析Word文档失败: {str(e)}")
            return []


def main():
    converter = WordToCsvConverter()

    # 示例用法
    print("Word到CSV转换器")
    print("1. 转换Word文档为CSV")
    print("2. 创建Word模板")
    print("3. 解析Word文档为题目")

    choice = input("请选择操作 (1/2/3): ")

    if choice == "1":
        word_path = input("请输入Word文档路径: ")
        csv_path = input("请输入输出CSV路径: ")

        success, message = converter.convert_word_to_csv(word_path, csv_path)
        print(message)

    elif choice == "2":
        print("可选模板类型: single_choice, true_false, essay, calculation")
        template_type = input("请输入模板类型: ")
        output_path = input("请输入输出路径: ")

        converter.create_word_template(template_type, output_path)
        print(f"模板已创建: {output_path}")

    elif choice == "3":
        word_path = input("请输入Word文档路径: ")
        questions = converter.parse_questions_from_word_doc(word_path)
        print(f"解析到 {len(questions)} 道题目:")
        for i, q in enumerate(questions, 1):
            print(f"{i}. {q.get('type', '未知类型')} - {q.get('chinese_stem', '无题干')}")

    else:
        print("无效选择")


if __name__ == "__main__":
    main()